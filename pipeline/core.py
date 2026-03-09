#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate a full DOCX medical-topic market report.
Outputs are written to:
  autofile/<医学主题>/
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple
from urllib.parse import quote_plus, urlparse
from urllib.request import urlopen

import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.patches import FancyBboxPatch
from docx import Document
from openpyxl import load_workbook
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


matplotlib.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
matplotlib.rcParams["axes.unicode_minus"] = False


TOPIC_LABEL = "医学主题"
DATA_DIR = Path("data")
DISEASE_NAME = "示例医学主题"
def report_title_for_topic(topic_name: str) -> str:
    topic_name = topic_name.strip()
    if topic_name.endswith("市场分析报告"):
        return f"《{topic_name}》"
    if topic_name.endswith("市场分析"):
        return f"《{topic_name}报告》"
    if topic_name.endswith("市场"):
        return f"《{topic_name}分析报告》"
    return f"《{topic_name}市场分析报告》"


REPORT_TITLE = report_title_for_topic(DISEASE_NAME)
EXCEL_PATH = DATA_DIR / f"{DISEASE_NAME}.xlsx"
TEMPLATE_PATH = Path("template.docx")
OUT_ROOT = Path("autofile") / DISEASE_NAME
FIG_DIR = OUT_ROOT / "figures"
FINAL_DOCX = OUT_ROOT / f"{REPORT_TITLE}_final.docx"
LITE_OUTPUT = False
WORKFLOW_MODE = "assistant"

LEGACY_DISEASE_TOKENS = [
    "儿童止咳祛痰",
    "儿童咳嗽",
    "颈椎病",
    "NAFLD",
    "非酒精性脂肪性肝病(NAFLD)",
]

QUARTER_RE = re.compile(r"^(20\d{2})Q([1-4])$")
FIG_TITLE_SERIAL_RE = re.compile(r"^\s*图表\d+-\d+[：:]\s*")
FIG23_CODEX_SPEC_NAME = "fig23_codex_spec.json"
FIG23_CODEX_SPEC_TEMPLATE_NAME = "fig23_codex_spec_template.json"
FIG23_CODEX_PROMPT_NAME = "fig23_codex_prompt.txt"
CODEX_CONTENT_BLUEPRINT_NAME = "codex_content_blueprint.txt"
FIGURE_SPECS_CODEX_TEMPLATE_NAME = "figure_specs_codex_template.json"
FIGURE_SPECS_CODEX_PROMPT_NAME = "figure_specs_codex_prompt.txt"
SEMANTIC_REVIEW_PROMPT_NAME = "semantic_review_prompt.txt"
CODEX_GAP_PANEL_NAME = "codex_gap_panel.txt"
CHAPTER_PRECHECK_NAME = "chapter_precheck.txt"
CH4_NARRATIVE_BRIEF_NAME = "ch04_narrative_brief.txt"
CHAPTER_MIN_CHARS = {
    1: 3000,
    2: 3500,
    3: 4800,
    4: 3000,
    5: 4800,
    6: 4800,
    7: 4800,
}
CHAPTER_CHAR_TOLERANCE = 100
LOW_CONFIDENCE_PROFILE_KEYWORDS = {
    "疼痛",
    "炎症",
    "感染",
    "综合征",
    "障碍",
    "疾病",
    "病症",
    "骨科",
}
PROFILE_CONFIG_PATH = Path(__file__).with_name("disease_profiles.json")
_PROFILE_CONFIG_CACHE: Dict[str, object] | None = None
_ACTIVE_PROFILE_CACHE: Tuple[str, str] | None = None


def chapter_char_shortfall(chapter: int, chars: int) -> int:
    floor = int(CHAPTER_MIN_CHARS.get(chapter, 0))
    return max(0, floor - int(chars))


def chapter_char_gate_ok(chapter: int, chars: int) -> bool:
    return chapter_char_shortfall(chapter, chars) <= CHAPTER_CHAR_TOLERANCE


def qkey(q: str) -> int:
    m = QUARTER_RE.match(str(q).strip())
    if not m:
        return -1
    return int(m.group(1)) * 10 + int(m.group(2))


def year_of_quarter(q: str) -> int:
    return int(str(q)[:4])


def default_excel_path(topic_name: str) -> Path:
    return DATA_DIR / f"{topic_name.strip()}.xlsx"


def disease_text_replacements() -> Dict[str, str]:
    replacements = {
        "<<<在此填写疾病名>>>": DISEASE_NAME,
        "<<<疾病名>>>": DISEASE_NAME,
    }
    for token in LEGACY_DISEASE_TOKENS:
        replacements[token] = DISEASE_NAME
    return replacements


def normalize_disease_text(text: str) -> str:
    out = str(text)
    for old, new in disease_text_replacements().items():
        out = out.replace(old, new)
    return out


def strip_figure_serial_prefix(text: str) -> str:
    raw = normalize_disease_text(str(text)).strip()
    return FIG_TITLE_SERIAL_RE.sub("", raw, count=1).strip()


def compose_figure_title(serial: str, title_fragment: str) -> str:
    fragment = strip_figure_serial_prefix(title_fragment)
    if not fragment:
        return f"图表{serial}"
    return normalize_disease_text(f"图表{serial}：{fragment}")


def normalize_reference_line(line: str) -> str:
    out = normalize_disease_text(str(line))
    for token in ["\u00A0", "\u2002", "\u2003", "\u2009", "\u202F", "\u3000", "\t"]:
        out = out.replace(token, " ")
    out = re.sub(r" {2,}", " ", out)
    return out.strip()


def _is_codex_authored_marker(value: object) -> bool:
    marker = str(value).strip().lower()
    return marker in {"codex", "openai-codex", "codex-cli"}


def read_fig23_codex_spec() -> Dict[str, object]:
    path = OUT_ROOT / FIG23_CODEX_SPEC_NAME
    if not path.exists():
        return {}
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    return raw if isinstance(raw, dict) else {}


def fig23_codex_authored_ok() -> bool:
    raw = read_fig23_codex_spec()
    if not raw:
        return False
    marker = raw.get("authored_by", raw.get("generated_by", ""))
    return _is_codex_authored_marker(marker)


def fig23_spec_origin(profile_id: str | None = None) -> str:
    if read_fig23_codex_spec():
        return "codex_spec"

    file_path = OUT_ROOT / "figure_specs.json"
    if file_path.exists():
        try:
            raw = json.loads(file_path.read_text(encoding="utf-8"))
        except Exception:
            raw = {}
        if isinstance(raw, dict) and isinstance(raw.get("fig_2_3"), dict):
            return "figure_specs"

    profile = get_profile_data(profile_id)
    figure_specs = profile.get("figure_specs", {})
    if isinstance(figure_specs, dict) and isinstance(figure_specs.get("fig_2_3"), dict):
        return "profile"
    return "profile"


def normalize_disease_value(value):
    if isinstance(value, str):
        return normalize_disease_text(value)
    return value


def load_disease_profiles() -> Dict[str, object]:
    global _PROFILE_CONFIG_CACHE
    if _PROFILE_CONFIG_CACHE is not None:
        return _PROFILE_CONFIG_CACHE

    fallback: Dict[str, object] = {
        "default_profile": "generic",
        "profiles": [
            {"id": "cervical", "priority": 100, "keywords": ["颈椎", "腰椎", "脊柱", "关节", "骨", "肌", "疼痛", "骨科", "椎间盘"], "query_aliases": {"颈椎病": "cervical spondylosis"}, "fig23_layout_mode": "dual_panel", "fig23_require_core_node": True, "fig23_core_label_template": "{disease}\n核心病理负担", "fig23_disallow_nodes": ["代谢系统", "营养状态"], "fig23_top_to_core": ["神经系统", "肌肉骨骼系统", "血管系统"], "fig23_core_to_bottom": ["睡眠系统", "心理行为", "内分泌代谢"], "fig23_dual_panel": {}, "fig23_semantic": {"expected_title_template": "图表2-3：{disease}相关系统关系图（分层布局）", "forbidden_title_templates": []}},
            {"id": "gastritis", "priority": 110, "keywords": ["慢性胃炎", "胃炎", "萎缩性胃炎", "胃黏膜", "幽门螺杆菌", "消化不良", "胃病", "反流性胃炎"], "query_aliases": {"慢性胃炎": "chronic gastritis", "萎缩性胃炎": "atrophic gastritis"}, "fig23_layout_mode": "causal_chain", "fig23_require_core_node": True, "fig23_core_label_template": "{disease}\n(胃黏膜炎症)", "fig23_disallow_nodes": ["代谢系统", "营养状态"], "fig23_top_to_core": [], "fig23_core_to_bottom": [], "fig23_semantic": {"expected_title_template": "图表2-3：{disease}病因-病理-临床后果关系图（分层布局）", "forbidden_title_templates": ["图表2-3：{disease}相关系统关系图（分层布局）"]}},
            {"id": "pharyngitis", "priority": 108, "keywords": ["慢性咽炎", "咽炎", "咽喉炎", "咽喉", "咽部", "咽痛", "咽异物感"], "query_aliases": {"慢性咽炎": "chronic pharyngitis", "慢性咽喉炎": "chronic pharyngolaryngitis"}, "fig23_layout_mode": "systems_map", "fig23_require_core_node": True, "fig23_core_label_template": "{disease}\n(咽黏膜慢性炎症)", "fig23_disallow_nodes": ["代谢系统", "营养状态"], "fig23_top_to_core": ["免疫系统", "神经系统", "上气道微生态"], "fig23_core_to_bottom": ["耳鼻喉症状", "睡眠系统", "心理行为"], "fig23_semantic": {"expected_title_template": "图表2-3：{disease}相关系统关系图（分层布局）", "forbidden_title_templates": []}},
            {"id": "osteoarthritis", "priority": 112, "keywords": ["骨关节炎", "膝骨关节炎", "髋骨关节炎", "退行性关节炎", "osteoarthritis", "OA"], "query_aliases": {"骨关节炎": "osteoarthritis", "膝骨关节炎": "knee osteoarthritis"}, "fig23_layout_mode": "systems_map", "fig23_require_core_node": True, "fig23_core_label_template": "{disease}\n(关节软骨退变)", "fig23_disallow_nodes": ["代谢系统", "营养状态"], "fig23_top_to_core": ["免疫系统", "神经系统", "肌肉骨骼系统"], "fig23_core_to_bottom": ["运动功能", "睡眠系统", "心理行为"], "fig23_semantic": {"expected_title_template": "图表2-3：{disease}相关系统关系图（分层布局）", "forbidden_title_templates": []}},
            {"id": "respiratory", "priority": 90, "keywords": ["咳", "痰", "呼吸", "肺", "哮喘", "支气管", "气道", "感冒", "鼻炎"], "query_aliases": {"儿童止咳祛痰": "pediatric cough expectorant", "儿童咳嗽": "pediatric chronic cough"}, "fig23_layout_mode": "systems_map", "fig23_require_core_node": True, "fig23_core_label_template": "{disease}\n核心病理负担", "fig23_disallow_nodes": ["代谢系统", "营养状态"], "fig23_top_to_core": ["神经系统", "内分泌系统", "免疫系统"], "fig23_core_to_bottom": ["呼吸系统", "消化系统", "肌肉骨骼系统"], "fig23_semantic": {"expected_title_template": "图表2-3：{disease}相关系统关系图（分层布局）", "forbidden_title_templates": []}},
            {"id": "generic", "priority": 0, "keywords": [], "query_aliases": {}, "fig23_layout_mode": "systems_map", "fig23_require_core_node": True, "fig23_core_label_template": "{disease}\n核心病理负担", "fig23_disallow_nodes": ["代谢系统", "营养状态"], "fig23_top_to_core": ["神经系统", "内分泌系统", "免疫系统"], "fig23_core_to_bottom": ["心血管系统", "消化系统", "肾脏系统"], "fig23_semantic": {"expected_title_template": "图表2-3：{disease}相关系统关系图（分层布局）", "forbidden_title_templates": []}},
        ],
    }

    if not PROFILE_CONFIG_PATH.exists():
        _PROFILE_CONFIG_CACHE = fallback
        return _PROFILE_CONFIG_CACHE

    try:
        raw = json.loads(PROFILE_CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception:
        _PROFILE_CONFIG_CACHE = fallback
        return _PROFILE_CONFIG_CACHE

    profiles = raw.get("profiles", [])
    if not isinstance(profiles, list) or not profiles:
        _PROFILE_CONFIG_CACHE = fallback
        return _PROFILE_CONFIG_CACHE

    default_profile = str(raw.get("default_profile", "generic")).strip() or "generic"
    normalized_profiles: List[Dict[str, object]] = []
    for p in profiles:
        if not isinstance(p, dict):
            continue
        pid = str(p.get("id", "")).strip()
        if not pid:
            continue
        normalized_profiles.append(
            {
                "id": pid,
                "priority": int(p.get("priority", 0)),
                "keywords": [str(x).strip() for x in (p.get("keywords") or []) if str(x).strip()],
                "query_aliases": {str(k).strip(): str(v).strip() for k, v in (p.get("query_aliases") or {}).items() if str(k).strip() and str(v).strip()},
                "fig23_layout_mode": str(p.get("fig23_layout_mode", "systems_map")).strip() or "systems_map",
                "fig23_require_core_node": bool(p.get("fig23_require_core_node", False)),
                "fig23_core_label_template": str(p.get("fig23_core_label_template", "{disease}\n核心病理负担")).strip() or "{disease}\n核心病理负担",
                "fig23_disallow_nodes": [str(x).strip() for x in (p.get("fig23_disallow_nodes") or []) if str(x).strip()],
                "fig23_top_to_core": [str(x).strip() for x in (p.get("fig23_top_to_core") or []) if str(x).strip()],
                "fig23_core_to_bottom": [str(x).strip() for x in (p.get("fig23_core_to_bottom") or []) if str(x).strip()],
                "fig23_dual_panel": p.get("fig23_dual_panel") if isinstance(p.get("fig23_dual_panel"), dict) else {},
                "fig23_semantic": p.get("fig23_semantic") or {},
                "figure_specs": p.get("figure_specs") if isinstance(p.get("figure_specs"), dict) else {},
            }
        )
    if not normalized_profiles:
        _PROFILE_CONFIG_CACHE = fallback
        return _PROFILE_CONFIG_CACHE

    profile_ids = {str(x.get("id", "")) for x in normalized_profiles}
    if default_profile not in profile_ids:
        default_profile = "generic" if "generic" in profile_ids else str(normalized_profiles[-1].get("id", "generic"))
    _PROFILE_CONFIG_CACHE = {"default_profile": default_profile, "profiles": normalized_profiles}
    return _PROFILE_CONFIG_CACHE


def sorted_profiles() -> List[Dict[str, object]]:
    cfg = load_disease_profiles()
    profiles = list(cfg.get("profiles", []))
    profiles.sort(key=lambda x: int(x.get("priority", 0)), reverse=True)
    return profiles


def _normalize_profile_match_text(value: str) -> str:
    return re.sub(r"[\s\-_()（）·]+", "", str(value).strip().lower())


def _is_low_confidence_keyword(keyword: str) -> bool:
    kw = str(keyword).strip().lower()
    if not kw:
        return True
    if len(kw) <= 1:
        return True
    return kw in LOW_CONFIDENCE_PROFILE_KEYWORDS


def resolve_profile_id(disease_name: str) -> str:
    name = disease_name.strip()
    cfg = load_disease_profiles()
    default_id = str(cfg.get("default_profile", "generic"))
    if not name:
        return default_id

    normalized_name = _normalize_profile_match_text(name)

    # 1) Exact match first: query_alias key / profile keyword.
    for p in sorted_profiles():
        pid = str(p.get("id", "")).strip()
        aliases = p.get("query_aliases", {}) or {}
        if isinstance(aliases, dict):
            for alias in aliases.keys():
                if _normalize_profile_match_text(str(alias)) == normalized_name:
                    return pid
        for kw in p.get("keywords", []):
            if _normalize_profile_match_text(str(kw)) == normalized_name:
                return pid

    # 2) Fuzzy match with confidence scoring. A profile must have at least one
    # high-confidence hit; low-confidence generic tokens (e.g. "疼痛") cannot
    # decide profile alone.
    best_pid = default_id
    best_score = 0.0
    best_priority = -10**9
    for p in sorted_profiles():
        pid = str(p.get("id", "")).strip()
        score = 0.0
        has_high_conf_hit = False
        for kw in p.get("keywords", []):
            sk = str(kw).strip()
            if not sk or sk not in name:
                continue
            if _is_low_confidence_keyword(sk):
                score += 0.25
            else:
                has_high_conf_hit = True
                score += 1.0 + min(len(sk), 8) * 0.15
        if (not has_high_conf_hit) or score <= 0:
            continue
        pri = int(p.get("priority", 0))
        if (score > best_score) or (math.isclose(score, best_score) and pri > best_priority):
            best_score = score
            best_priority = pri
            best_pid = pid

    if best_score > 0:
        return best_pid
    return default_id


def active_profile_id() -> str:
    global _ACTIVE_PROFILE_CACHE
    name = DISEASE_NAME.strip()
    if _ACTIVE_PROFILE_CACHE and _ACTIVE_PROFILE_CACHE[0] == name:
        return _ACTIVE_PROFILE_CACHE[1]
    pid = resolve_profile_id(name)
    _ACTIVE_PROFILE_CACHE = (name, pid)
    return pid


def get_profile_data(profile_id: str | None = None) -> Dict[str, object]:
    pid = profile_id or active_profile_id()
    for p in sorted_profiles():
        if str(p.get("id", "")).strip() == pid:
            return p
    cfg = load_disease_profiles()
    fallback_id = str(cfg.get("default_profile", "generic"))
    for p in sorted_profiles():
        if str(p.get("id", "")).strip() == fallback_id:
            return p
    return {
        "id": "generic",
        "fig23_layout_mode": "systems_map",
        "fig23_require_core_node": True,
        "fig23_core_label_template": "{disease}\n核心病理负担",
        "fig23_disallow_nodes": ["代谢系统", "营养状态"],
        "fig23_top_to_core": ["神经系统", "内分泌系统", "免疫系统"],
        "fig23_core_to_bottom": ["心血管系统", "消化系统", "肾脏系统"],
        "fig23_dual_panel": {},
        "fig23_semantic": {},
        "figure_specs": {},
    }


def query_alias_map() -> Dict[str, str]:
    merged: Dict[str, str] = {}
    for p in sorted_profiles():
        aliases = p.get("query_aliases", {}) or {}
        if isinstance(aliases, dict):
            for k, v in aliases.items():
                sk, sv = str(k).strip(), str(v).strip()
                if sk and sv:
                    merged[sk] = sv
    return merged


def disease_query_term(disease_name: str) -> str:
    name = disease_name.strip()
    alias_map = query_alias_map()
    if name in alias_map:
        return alias_map[name]
    if any("\u4e00" <= ch <= "\u9fff" for ch in name):
        return f"{name} treatment guideline"
    return name


def configure_output_mode(lite_output: bool = False) -> None:
    global LITE_OUTPUT
    LITE_OUTPUT = bool(lite_output)


def render_disease_template(template: str, default: str = "") -> str:
    raw = str(template).strip() or str(default).strip()
    # Support escaped newlines from JSON config (e.g. "\\n").
    raw = raw.replace("\\n", "\n").replace("\\t", "\t")
    try:
        rendered = raw.format(disease=DISEASE_NAME)
    except Exception:
        rendered = raw.replace("{disease}", DISEASE_NAME)
    return normalize_disease_text(rendered)


def render_profile_title_template(template: str) -> str:
    return render_disease_template(template, default="图表2-3：{disease}相关系统关系图（分层布局）")


def fig23_layout_mode(profile_id: str | None = None) -> str:
    runtime_spec = load_figure_specs(profile_id).get("fig_2_3", {})
    if isinstance(runtime_spec, dict):
        mode = str(runtime_spec.get("layout_mode", "")).strip()
        if mode:
            return mode
    profile = get_profile_data(profile_id)
    mode = str(profile.get("fig23_layout_mode", "systems_map")).strip()
    return mode or "systems_map"


def fig23_expected_caption(profile_id: str | None = None) -> str:
    runtime_spec = load_figure_specs(profile_id).get("fig_2_3", {})
    if isinstance(runtime_spec, dict):
        raw_caption = str(runtime_spec.get("caption") or runtime_spec.get("title") or "").strip()
        if raw_caption:
            return compose_figure_title("2-3", render_disease_template(raw_caption, default=raw_caption))
    profile = get_profile_data(profile_id)
    sem = profile.get("fig23_semantic", {})
    sem_dict = sem if isinstance(sem, dict) else {}
    template = str(sem_dict.get("expected_title_template", "图表2-3：{disease}相关系统关系图（分层布局）")).strip()
    return render_profile_title_template(template)


def fig23_forbidden_captions(profile_id: str | None = None) -> List[str]:
    profile = get_profile_data(profile_id)
    sem = profile.get("fig23_semantic", {})
    sem_dict = sem if isinstance(sem, dict) else {}
    templates = sem_dict.get("forbidden_title_templates", []) or []
    out: List[str] = []
    for t in templates:
        s = render_profile_title_template(str(t))
        if s:
            out.append(s)
    return out


def fig23_require_core_node(profile_id: str | None = None) -> bool:
    profile = get_profile_data(profile_id)
    return bool(profile.get("fig23_require_core_node", False))


def fig23_core_label(profile_id: str | None = None) -> str:
    profile = get_profile_data(profile_id)
    template = str(profile.get("fig23_core_label_template", "{disease}\n核心病理负担")).strip() or "{disease}\n核心病理负担"
    return render_disease_template(template)


def fig23_disallow_nodes(profile_id: str | None = None) -> List[str]:
    profile = get_profile_data(profile_id)
    out = [str(x).strip() for x in (profile.get("fig23_disallow_nodes") or []) if str(x).strip()]
    runtime_spec = load_figure_specs(profile_id).get("fig_2_3", {})
    if isinstance(runtime_spec, dict):
        extra = runtime_spec.get("disallow_nodes")
        if isinstance(extra, list):
            for item in extra:
                token = str(item).strip()
                if token and token not in out:
                    out.append(token)
    return out


def fig23_top_to_core_nodes(profile_id: str | None = None) -> List[str]:
    profile = get_profile_data(profile_id)
    return [str(x).strip() for x in (profile.get("fig23_top_to_core") or []) if str(x).strip()]


def fig23_core_to_bottom_nodes(profile_id: str | None = None) -> List[str]:
    profile = get_profile_data(profile_id)
    return [str(x).strip() for x in (profile.get("fig23_core_to_bottom") or []) if str(x).strip()]


def fig23_dual_panel_config(profile_id: str | None = None) -> Dict[str, object]:
    runtime_spec = load_figure_specs(profile_id).get("fig_2_3", {})
    if isinstance(runtime_spec, dict):
        cfg = runtime_spec.get("dual_panel", runtime_spec.get("panels", {}))
        if isinstance(cfg, dict) and cfg:
            return cfg
    profile = get_profile_data(profile_id)
    cfg = profile.get("fig23_dual_panel", {})
    return cfg if isinstance(cfg, dict) else {}


def fig23_layered_path_config(profile_id: str | None = None) -> Dict[str, object]:
    runtime_spec = load_figure_specs(profile_id).get("fig_2_3", {})
    if isinstance(runtime_spec, dict):
        cfg = runtime_spec.get("layered_path", runtime_spec.get("pathway", {}))
        if isinstance(cfg, dict) and cfg:
            return cfg
    return {}


def _fig23_anchor_point(x: float, y: float, width: float, height: float, anchor: str) -> Tuple[float, float]:
    m = {
        "center": (x, y),
        "north": (x, y + height / 2.0),
        "south": (x, y - height / 2.0),
        "west": (x - width / 2.0, y),
        "east": (x + width / 2.0, y),
    }
    return m.get(anchor, m["center"])


def _fig23_label_role(label: str) -> str:
    t = normalize_disease_text(str(label)).replace("\n", "").strip()
    if not t:
        return "other"
    # Driver-like labels (etiology/pathophysiology upstream).
    if any(k in t for k in ["病因", "驱动", "炎症", "免疫", "神经", "内分泌", "HPA", "回路", "微生态", "供血", "受压"]):
        if "核心" in t and any(k in t for k in ["症状", "后果", "风险", "负担"]):
            return "outcome"
        return "driver"
    # Outcome-like labels (clinical burden/downstream consequences).
    if any(k in t for k in ["症状", "后果", "风险", "负担", "回避", "功能", "并发", "睡眠", "心血管", "消化", "行为"]):
        return "outcome"
    if "系统" in t:
        return "system"
    return "other"


def _fig23_is_system_semantic(label: str) -> bool:
    t = normalize_disease_text(str(label)).replace("\n", "").strip()
    if not t:
        return False
    if "系统" in t:
        return True
    # Keep this list focused on domain labels that are routinely used as
    # subsystem names even without an explicit "系统" suffix.
    implicit_system_terms = [
        "神经",
        "免疫",
        "内分泌",
        "血管",
        "肌肉骨骼",
        "呼吸",
        "消化",
        "肾脏",
        "睡眠",
        "心理行为",
        "上气道微生态",
    ]
    return any(k in t for k in implicit_system_terms)


def validate_fig23_structural_rules(profile_id: str | None = None) -> Dict[str, object]:
    """
    Structural QA for fig_2_3. The checks are config-driven and run before/without image OCR:
      1) causal direction consistency (semantic heuristic)
      2) same-track overlap (line collisions on same vertical track)
      3) bidirectional readability (A->B and B->A must be geometrically distinguishable)
      4) layer consistency (same row should not mix system/non-system semantics)
    """

    out: Dict[str, object] = {
        "layout_mode": fig23_layout_mode(profile_id),
        "causal_direction_issues": [],
        "same_track_overlap_issues": [],
        "bidirectional_readability_issues": [],
        "layer_consistency_issues": [],
        "node_spacing_issues": [],
    }
    if str(out["layout_mode"]) == "layered_path":
        cfg = fig23_layered_path_config(profile_id)
        if not isinstance(cfg, dict):
            return out
        left_nodes = [str(x).strip() for x in (cfg.get("left_nodes") or []) if str(x).strip()]
        right_nodes = [str(x).strip() for x in (cfg.get("right_nodes") or []) if str(x).strip()]
        if len(left_nodes) < 2 or len(right_nodes) < 2:
            out["node_spacing_issues"].append("layered_path:左右两侧节点数量不足，无法表达核心关系")
        return out
    if str(out["layout_mode"]) != "dual_panel":
        return out

    panel_cfg = fig23_dual_panel_config(profile_id)
    if not isinstance(panel_cfg, dict):
        return out

    for panel_name in ("left", "right"):
        cfg = panel_cfg.get(panel_name)
        if not isinstance(cfg, dict):
            continue

        panel_title = render_disease_template(str(cfg.get("title", "")).strip())
        nodes_raw = cfg.get("nodes", [])
        edges_raw = cfg.get("edges", [])

        nodes: Dict[str, Dict[str, object]] = {}
        if isinstance(nodes_raw, list):
            for node in nodes_raw:
                if not isinstance(node, dict):
                    continue
                node_id = str(node.get("id", "")).strip()
                if not node_id:
                    continue
                x = float(node.get("x", 0.5))
                y = float(node.get("y", 0.5))
                w = float(node.get("width", node.get("w", 0.30)))
                h = float(node.get("height", node.get("h", 0.12)))
                label_tpl = str(node.get("label", node_id)).strip().replace("\\n", "\n")
                label = render_disease_template(label_tpl, default=node_id)
                nodes[node_id] = {
                    "id": node_id,
                    "label": label,
                    "x": x,
                    "y": y,
                    "w": w,
                    "h": h,
                }

        node_ids = sorted(list(nodes.keys()))
        node_spacing_seen: set[str] = set()
        boundary_seen: set[str] = set()
        for node_id in node_ids:
            node = nodes[node_id]
            left = float(node["x"]) - float(node["w"]) / 2.0
            right = float(node["x"]) + float(node["w"]) / 2.0
            bottom = float(node["y"]) - float(node["h"]) / 2.0
            top = float(node["y"]) + float(node["h"]) / 2.0
            if left < 0.02 or right > 0.98 or bottom < 0.04 or top > 0.96:
                issue_key = f"{panel_name}:{node_id}:boundary"
                if issue_key not in boundary_seen:
                    boundary_seen.add(issue_key)
                    out["node_spacing_issues"].append(f"{panel_name}:{node_id} 过于贴近画布边界（L={left:.3f},R={right:.3f},B={bottom:.3f},T={top:.3f}）")
        for a_idx in range(len(node_ids)):
            for b_idx in range(a_idx + 1, len(node_ids)):
                a = node_ids[a_idx]
                b = node_ids[b_idx]
                na = nodes[a]
                nb = nodes[b]
                ax_gap = abs(float(na["x"]) - float(nb["x"])) - (float(na["w"]) + float(nb["w"])) / 2.0
                ay_gap = abs(float(na["y"]) - float(nb["y"])) - (float(na["h"]) + float(nb["h"])) / 2.0
                issue_key = f"{panel_name}:{a}|{b}"
                if issue_key in node_spacing_seen:
                    continue
                if ax_gap < -0.002 and ay_gap < -0.002:
                    node_spacing_seen.add(issue_key)
                    out["node_spacing_issues"].append(f"{panel_name}:{a} 与 {b} 节点框发生重叠")
                    continue
                same_row = abs(float(na["y"]) - float(nb["y"])) <= max(float(na["h"]), float(nb["h"])) * 0.35
                same_col = abs(float(na["x"]) - float(nb["x"])) <= max(float(na["w"]), float(nb["w"])) * 0.35
                if same_row and ax_gap < 0.03:
                    node_spacing_seen.add(issue_key)
                    out["node_spacing_issues"].append(f"{panel_name}:{a} 与 {b} 同层间距不足（gap={ax_gap:.3f}）")
                    continue
                if same_col and ay_gap < 0.04:
                    node_spacing_seen.add(issue_key)
                    out["node_spacing_issues"].append(f"{panel_name}:{a} 与 {b} 纵向间距不足（gap={ay_gap:.3f}）")

        edges: List[Dict[str, object]] = []
        if isinstance(edges_raw, list):
            for idx, edge in enumerate(edges_raw):
                if not isinstance(edge, dict):
                    continue
                src = str(edge.get("from", "")).strip()
                dst = str(edge.get("to", "")).strip()
                if src not in nodes or dst not in nodes:
                    continue
                from_anchor = str(edge.get("from_anchor", "center")).strip() or "center"
                to_anchor = str(edge.get("to_anchor", "center")).strip() or "center"
                src_node = nodes[src]
                dst_node = nodes[dst]
                src_pt = _fig23_anchor_point(float(src_node["x"]), float(src_node["y"]), float(src_node["w"]), float(src_node["h"]), from_anchor)
                dst_pt = _fig23_anchor_point(float(dst_node["x"]), float(dst_node["y"]), float(dst_node["w"]), float(dst_node["h"]), to_anchor)

                via_pts: List[Tuple[float, float]] = []
                via_raw = edge.get("via", [])
                if isinstance(via_raw, list):
                    for p in via_raw:
                        if isinstance(p, (list, tuple)) and len(p) >= 2:
                            try:
                                via_pts.append((float(p[0]), float(p[1])))
                            except Exception:
                                continue
                points = [src_pt] + via_pts + [dst_pt]
                edges.append(
                    {
                        "idx": idx,
                        "src": src,
                        "dst": dst,
                        "from_anchor": from_anchor,
                        "to_anchor": to_anchor,
                        "dashed": bool(edge.get("dashed", False)),
                        "via": via_pts,
                        "points": points,
                    }
                )

        # 1) Causal direction semantic check.
        for e in edges:
            if bool(e.get("dashed", False)):
                continue
            src = str(e["src"])
            dst = str(e["dst"])
            src_label = str(nodes[src]["label"])
            dst_label = str(nodes[dst]["label"])
            src_role = _fig23_label_role(src_label)
            dst_role = _fig23_label_role(dst_label)
            src_y = float(nodes[src]["y"])
            dst_y = float(nodes[dst]["y"])

            if src_role == "outcome" and dst_role == "driver":
                out["causal_direction_issues"].append(
                    f"{panel_name}:{src}->{dst} 可能因果反向（下游结果指向上游驱动）"
                )
            if src_role in {"driver", "system"} and dst_role == "outcome" and src_y <= (dst_y + 0.02):
                out["causal_direction_issues"].append(
                    f"{panel_name}:{src}->{dst} 方向可疑（驱动到后果未形成下行传导）"
                )

        # 2) Same-track overlap check.
        vertical_segments: List[Dict[str, object]] = []
        for e in edges:
            pts = e.get("points", [])
            if not isinstance(pts, list) or len(pts) < 2:
                continue
            for i in range(len(pts) - 1):
                p1 = pts[i]
                p2 = pts[i + 1]
                if not isinstance(p1, tuple) or not isinstance(p2, tuple):
                    continue
                x1, y1 = float(p1[0]), float(p1[1])
                x2, y2 = float(p2[0]), float(p2[1])
                if abs(x1 - x2) <= 0.015:
                    y_low = min(y1, y2)
                    y_high = max(y1, y2)
                    if (y_high - y_low) >= 0.08:
                        vertical_segments.append(
                            {
                                "edge_idx": int(e["idx"]),
                                "src": str(e["src"]),
                                "dst": str(e["dst"]),
                                "x": (x1 + x2) / 2.0,
                                "y_low": y_low,
                                "y_high": y_high,
                            }
                        )
        same_track_seen: set[str] = set()
        for i in range(len(vertical_segments)):
            for j in range(i + 1, len(vertical_segments)):
                a = vertical_segments[i]
                b = vertical_segments[j]
                if int(a["edge_idx"]) == int(b["edge_idx"]):
                    continue
                if abs(float(a["x"]) - float(b["x"])) > 0.015:
                    continue
                overlap_low = max(float(a["y_low"]), float(b["y_low"]))
                overlap_high = min(float(a["y_high"]), float(b["y_high"]))
                overlap_len = overlap_high - overlap_low
                if overlap_len < 0.12:
                    continue
                pair_key = "|".join(sorted([f"{a['src']}->{a['dst']}", f"{b['src']}->{b['dst']}"]))
                if pair_key in same_track_seen:
                    continue
                same_track_seen.add(pair_key)
                out["same_track_overlap_issues"].append(
                    f"{panel_name}:{a['src']}->{a['dst']} 与 {b['src']}->{b['dst']} 存在同轨重叠"
                )

        # 3) Bidirectional readability check.
        by_dir: Dict[Tuple[str, str], List[Dict[str, object]]] = {}
        for e in edges:
            k = (str(e["src"]), str(e["dst"]))
            by_dir.setdefault(k, []).append(e)
        for a_idx in range(len(node_ids)):
            for b_idx in range(a_idx + 1, len(node_ids)):
                a = node_ids[a_idx]
                b = node_ids[b_idx]
                ab = by_dir.get((a, b), [])
                ba = by_dir.get((b, a), [])
                if not ab or not ba:
                    continue
                readable = False
                for e1 in ab:
                    for e2 in ba:
                        via1 = e1.get("via", [])
                        via2 = e2.get("via", [])
                        has_detour = bool(via1) or bool(via2)
                        anchors_same_track = (e1.get("from_anchor"), e1.get("to_anchor")) == (
                            e2.get("to_anchor"),
                            e2.get("from_anchor"),
                        )
                        if has_detour or (not anchors_same_track):
                            readable = True
                            break
                    if readable:
                        break
                if not readable:
                    out["bidirectional_readability_issues"].append(
                        f"{panel_name}:{a}<->{b} 双向关系未做几何避让（可读性不足）"
                    )

        # 4) Layer consistency check (same row semantic mixing).
        row_buckets: Dict[int, List[str]] = {}
        for node_id, node in nodes.items():
            y = float(node["y"])
            # 0.08 row height bucket for "same layer" judgment.
            bucket = int(round(y / 0.08))
            row_buckets.setdefault(bucket, []).append(node_id)
        for _, ids in row_buckets.items():
            if len(ids) < 2:
                continue
            labels = [str(nodes[n]["label"]) for n in ids]
            has_system = any(_fig23_is_system_semantic(lb) for lb in labels)
            has_non_system = any((not _fig23_is_system_semantic(lb)) for lb in labels)
            if has_system and has_non_system:
                out["layer_consistency_issues"].append(
                    f"{panel_name}:同层节点混合“系统”与“非系统”语义（{', '.join(ids)}）"
                )

        # Panel semantic hint checks.
        if "后果" in panel_title:
            rev = 0
            for e in edges:
                if bool(e.get("dashed", False)):
                    continue
                src_role = _fig23_label_role(str(nodes[str(e["src"])]["label"]))
                dst_role = _fig23_label_role(str(nodes[str(e["dst"])]["label"]))
                if src_role == "outcome" and dst_role in {"driver", "system"}:
                    rev += 1
            if rev > 0:
                out["causal_direction_issues"].append(
                    f"{panel_name}:标题为“后果层”但存在{rev}条后果反指驱动/系统的实线关系"
                )

    return out


def load_figure_specs(profile_id: str | None = None) -> Dict[str, Dict[str, object]]:
    """
    Figure generation config, merged in this order:
      1) profile.figure_specs
      2) autofile/<disease>/figure_specs.json
      3) autofile/<disease>/fig23_codex_spec.json (fig_2_3 only)
    Later source has higher priority.
    """
    merged: Dict[str, Dict[str, object]] = {}

    profile = get_profile_data(profile_id)
    p_specs = profile.get("figure_specs", {})
    if isinstance(p_specs, dict):
        for k, v in p_specs.items():
            sid = str(k).strip()
            if not sid or not isinstance(v, dict):
                continue
            merged[sid] = dict(v)

    file_path = OUT_ROOT / "figure_specs.json"
    if file_path.exists():
        try:
            raw = json.loads(file_path.read_text(encoding="utf-8"))
        except Exception:
            raw = {}
        if isinstance(raw, dict):
            for k, v in raw.items():
                sid = str(k).strip()
                if not sid or not isinstance(v, dict):
                    continue
                base = merged.get(sid, {})
                nxt = dict(base)
                nxt.update(v)
                merged[sid] = nxt

    fig23_codex = read_fig23_codex_spec()
    if fig23_codex:
        base = merged.get("fig_2_3", {})
        nxt = dict(base) if isinstance(base, dict) else {}
        nxt.update(fig23_codex)
        merged["fig_2_3"] = nxt
    return merged


def _safe_http_json(url: str, timeout: int = 10) -> dict:
    with urlopen(url, timeout=timeout) as resp:
        raw = resp.read().decode("utf-8", errors="ignore")
    return json.loads(raw)


def fetch_pubmed_evidence(disease_name: str, max_items: int = 8) -> List[Tuple[str, str, str, str, str]]:
    """
    Returns list of tuples:
      (title, source_org, year, keypoint, url)
    """
    term = disease_query_term(disease_name)
    queries = [
        f'("{term}"[Title/Abstract]) AND (guideline OR consensus OR recommendation)',
        f'("{term}"[Title/Abstract]) AND (diagnosis OR treatment)',
    ]
    pmids: List[str] = []
    seen: set[str] = set()
    try:
        for q in queries:
            esearch_url = (
                "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
                f"?db=pubmed&retmode=json&retmax=12&sort=relevance&mindate=2016&datetype=pdat&term={quote_plus(q)}"
            )
            data = _safe_http_json(esearch_url, timeout=10)
            ids = data.get("esearchresult", {}).get("idlist", []) or []
            for pid in ids:
                if pid not in seen:
                    seen.add(pid)
                    pmids.append(pid)
                if len(pmids) >= max_items * 2:
                    break
            if len(pmids) >= max_items * 2:
                break
    except Exception:
        return []

    if not pmids:
        return []

    summary_items: List[Tuple[str, str, str, str, str]] = []
    try:
        id_str = ",".join(pmids[: max_items * 2])
        esummary_url = (
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
            f"?db=pubmed&retmode=json&id={quote_plus(id_str)}"
        )
        data = _safe_http_json(esummary_url, timeout=10)
        result = data.get("result", {})
        for pid in pmids:
            item = result.get(pid, {})
            if not item:
                continue
            title = str(item.get("title", "")).strip().rstrip(".")
            source = str(item.get("fulljournalname", "")).strip() or str(item.get("source", "")).strip() or "PubMed"
            pubdate = str(item.get("pubdate", "")).strip()
            year_m = re.search(r"(19|20)\d{2}", pubdate)
            year = year_m.group(0) if year_m else "2024"
            if not title:
                continue
            keypoint = "提供诊断、治疗或随访证据，可用于人群分层与处方路径构建"
            url = f"https://pubmed.ncbi.nlm.nih.gov/{pid}/"
            summary_items.append((title, source, year, keypoint, url))
            if len(summary_items) >= max_items:
                break
    except Exception:
        return []

    return summary_items


def resolve_topic_name(
    topic: str | None = None,
    disease: str | None = None,
    from_readme: bool = False,
    readme_path: str | Path = "README.md",
) -> str:
    if topic and topic.strip():
        return topic.strip()
    if disease and disease.strip():
        return disease.strip()

    if not from_readme:
        raise ValueError("缺少医学主题。请传入 --topic，或使用 --from-readme 从 README 读取。")

    path = Path(readme_path)
    if not path.exists():
        raise FileNotFoundError(f"README file not found: {path}")

    text = path.read_text(encoding="utf-8")
    m = re.search(r"^\s*医学主题\s*[：:]\s*(.+?)\s*$", text, re.M)
    if not m:
        m = re.search(r"^\s*疾病名\s*[：:]\s*(.+?)\s*$", text, re.M)
    if not m:
        raise ValueError(f"在 {path} 中未找到“医学主题：”或“疾病名：”配置行。")

    val = m.group(1).strip()
    invalid_tokens = ["<<<", ">>>", "在此填写", "<疾病名>", "疾病名占位符", "<医学主题>", "医学主题占位符"]
    if (not val) or any(tok in val for tok in invalid_tokens):
        raise ValueError(
            f"{path} 中的医学主题仍是占位符（当前值：{val}）。请先改成真实医学主题，或直接用 --topic。"
        )
    return val


def resolve_disease_name(
    disease: str | None = None,
    from_readme: bool = False,
    readme_path: str | Path = "README.md",
) -> str:
    return resolve_topic_name(topic=None, disease=disease, from_readme=from_readme, readme_path=readme_path)


def configure_runtime(
    disease_name: str,
    excel_path: Path | None = None,
    template_path: Path | None = None,
    out_base: Path | None = None,
) -> None:
    """Configure global runtime paths so the pipeline can run for any medical topic."""
    global DISEASE_NAME, REPORT_TITLE, EXCEL_PATH, TEMPLATE_PATH, OUT_ROOT, FIG_DIR, FINAL_DOCX, _ACTIVE_PROFILE_CACHE

    DISEASE_NAME = disease_name.strip()
    _ACTIVE_PROFILE_CACHE = None
    REPORT_TITLE = report_title_for_topic(DISEASE_NAME)
    EXCEL_PATH = Path(excel_path) if excel_path is not None else default_excel_path(DISEASE_NAME)
    TEMPLATE_PATH = Path(template_path) if template_path is not None else Path("template.docx")
    base = Path(out_base) if out_base is not None else Path("autofile")
    OUT_ROOT = base / DISEASE_NAME
    FIG_DIR = OUT_ROOT / "figures"
    FINAL_DOCX = OUT_ROOT / f"{REPORT_TITLE}_final.docx"


def is_cervical_profile() -> bool:
    """Return True when the active disease should use musculoskeletal profile."""
    return active_profile_id() == "cervical"


def is_respiratory_profile() -> bool:
    return active_profile_id() == "respiratory"


def is_gastritis_profile() -> bool:
    return active_profile_id() == "gastritis"


def is_pharyngitis_profile() -> bool:
    return active_profile_id() == "pharyngitis"


def is_sciatica_profile() -> bool:
    return active_profile_id() == "sciatica"


def ensure_runtime_dirs() -> None:
    OUT_ROOT.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def cleanup_stale_final_docx() -> None:
    for old_path in OUT_ROOT.glob("*_final.docx"):
        if old_path == FINAL_DOCX:
            continue
        backup_if_exists(old_path)
        old_path.unlink(missing_ok=True)


def backup_if_exists(path: Path) -> None:
    if path.exists():
        backup_dir = OUT_ROOT / "backup"
        backup_dir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = backup_dir / f"{path.name}.{stamp}.bak"
        shutil.copy2(path, backup_path)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    backup_if_exists(path)
    path.write_text(normalize_disease_text(text), encoding="utf-8")


def write_csv(path: Path, rows: List[Dict[str, str]], headers: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    backup_if_exists(path)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        normalized_rows: List[Dict[str, str]] = []
        for row in rows:
            normalized_rows.append({k: normalize_disease_value(v) for k, v in row.items()})
        writer.writerows(normalized_rows)


def get_workbook_sheet_names(xlsx: Path) -> List[str]:
    try:
        return list(pd.ExcelFile(xlsx).sheet_names)
    except Exception:
        try:
            wb = load_workbook(xlsx, read_only=True, data_only=False)
            return [ws.title for ws in wb.worksheets]
        except Exception:
            return []


def parse_category_sheet(xlsx: Path, sheet: str) -> pd.DataFrame:
    raw = pd.read_excel(xlsx, sheet_name=sheet, header=None)
    records: List[Tuple[str, float]] = []
    for i in range(raw.shape[0]):
        q = str(raw.iat[i, 0]).strip()
        if QUARTER_RE.match(q):
            val = pd.to_numeric(raw.iat[i, 1], errors="coerce")
            if pd.notna(val):
                records.append((q, float(val)))
    df = pd.DataFrame(records, columns=["quarter", "sales"])
    if df.empty:
        raise ValueError(f"Sheet {sheet} has no quarter-sales rows.")
    df = df.drop_duplicates(subset=["quarter"]).copy()
    df["qk"] = df["quarter"].apply(qkey)
    df = df.sort_values("qk").drop(columns=["qk"]).reset_index(drop=True)
    return df


def missing_category_sheet_frame(quarters: List[str], column: str) -> pd.DataFrame:
    if not quarters:
        return pd.DataFrame(columns=["quarter", column])
    return pd.DataFrame({"quarter": list(quarters), column: [0.0] * len(quarters)})


def collapse_duplicate_top_rows(latest_q: str, latest_df: pd.DataFrame, full_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    if full_df.empty:
        return latest_df, full_df

    quarter_cols = [c for c in full_df.columns if c not in {"rank", "name"} and QUARTER_RE.match(str(c))]
    if not quarter_cols:
        return latest_df, full_df

    merged = full_df.copy()
    merged["name"] = merged["name"].astype(str).str.strip()
    merged = merged[merged["name"] != ""].copy()
    for col in quarter_cols:
        merged[col] = pd.to_numeric(merged[col], errors="coerce").fillna(0.0)
    merged = merged.groupby("name", as_index=False)[quarter_cols].sum()
    sort_col = latest_q if latest_q in merged.columns else quarter_cols[-1]
    merged = merged.sort_values([sort_col, "name"], ascending=[False, True]).reset_index(drop=True)
    merged.insert(0, "rank", range(1, len(merged) + 1))

    latest = merged[["rank", "name", sort_col]].rename(columns={sort_col: "sales"})
    latest = latest[latest["sales"] > 0].head(10).reset_index(drop=True)
    return latest, merged


def parse_top_sheet(xlsx: Path, sheet: str) -> Tuple[str, pd.DataFrame, pd.DataFrame]:
    """
    Returns:
      latest_quarter,
      latest_top10_df(columns=rank,name,sales),
      full_top_df(columns=rank,name,<quarters...>)
    """
    raw = pd.read_excel(xlsx, sheet_name=sheet, header=None)
    rank_row = None
    # Prefer exact header row "排名" (avoid matching title row like "品种排名及筛选...")
    for i in range(min(raw.shape[0], 16)):
        c0 = str(raw.iat[i, 0]).strip()
        if c0 == "排名" or c0.startswith("排名"):
            rank_row = i
            break
    if rank_row is None:
        for i in range(min(raw.shape[0], 16)):
            c0 = str(raw.iat[i, 0]).strip()
            if "排名" in c0:
                rank_row = i
                break
    if rank_row is None:
        rank_row = 1
    quarter_row = rank_row + 1

    # Top sheet usually has multiple quarter blocks:
    # sales / growth / share / company count / dosage forms ...
    # We must lock to "销售额（万元）" block only.
    sales_start = None
    for j in range(raw.shape[1]):
        h = str(raw.iat[rank_row, j]).strip()
        if "销售额" in h and "万" in h:
            sales_start = j
            break
    if sales_start is None:
        # fallback to first quarter-like region
        sales_start = 2

    sales_end = raw.shape[1]
    for j in range(sales_start + 1, raw.shape[1]):
        h = str(raw.iat[rank_row, j]).strip()
        if h and h != "nan":
            sales_end = j
            break

    quarter_cols: List[Tuple[int, str]] = []
    for j in range(sales_start, sales_end):
        val = str(raw.iat[quarter_row, j]).strip()
        if QUARTER_RE.match(val):
            quarter_cols.append((j, val))

    if not quarter_cols:
        raise ValueError(f"Sheet {sheet} has no quarter columns.")

    quarter_cols = sorted(quarter_cols, key=lambda x: qkey(x[1]))
    latest_q = quarter_cols[-1][1]

    records = []
    seen_ranks: set[int] = set()
    for i in range(quarter_row + 1, raw.shape[0]):
        rank = pd.to_numeric(raw.iat[i, 0], errors="coerce")
        if pd.isna(rank):
            continue
        rank_i = int(rank)
        if rank_i < 1 or rank_i > 50:
            continue
        # Many sheets contain repeated rank blocks (e.g., 通用名TOP20 + 剂型TOP20).
        # Keep the first complete rank block only to avoid double counting in CR5.
        if rank_i in seen_ranks:
            if len(seen_ranks) >= 20:
                break
            continue
        name = str(raw.iat[i, 1]).strip()
        if not name or name == "nan":
            continue
        rec: Dict[str, float | str | int] = {"rank": rank_i, "name": name}
        valid_count = 0
        for col, q in quarter_cols:
            val = pd.to_numeric(raw.iat[i, col], errors="coerce")
            if pd.notna(val):
                rec[q] = float(val)
                valid_count += 1
            else:
                rec[q] = np.nan
        if valid_count > 0:
            records.append(rec)
            seen_ranks.add(rank_i)

    full = pd.DataFrame(records).sort_values("rank").reset_index(drop=True)
    latest = full[["rank", "name", latest_q]].rename(columns={latest_q: "sales"})
    latest = latest.dropna(subset=["sales"]).head(10).copy()
    return latest_q, latest, full


@dataclass
class Ch4Data:
    quarterly: pd.DataFrame
    annual: pd.DataFrame
    latest_quarter: str
    latest_share: pd.DataFrame
    yoy_latest: pd.DataFrame
    top_hospital: pd.DataFrame
    top_drugstore: pd.DataFrame
    top_online: pd.DataFrame
    cr5_latest: pd.DataFrame
    cr5_trend: pd.DataFrame
    top_latest_name: Dict[str, str]


CH4_CHANNELS = ["医院端", "药店端", "线上端"]
CH4_CHANNEL_ORDER = {name: idx for idx, name in enumerate(CH4_CHANNELS)}
CH4_CHANNEL_ALIASES = {
    "医院端": "医院端",
    "医院": "医院端",
    "hospital": "医院端",
    "hospital channel": "医院端",
    "药店端": "药店端",
    "药店": "药店端",
    "零售药店": "药店端",
    "drugstore": "药店端",
    "pharmacy": "药店端",
    "retail pharmacy": "药店端",
    "线上端": "线上端",
    "线上": "线上端",
    "网上药店": "线上端",
    "online": "线上端",
    "e-commerce": "线上端",
    "ecommerce": "线上端",
    "电商": "线上端",
}
CH4_PLACEHOLDER_NAME_TERMS = ("示例", "sample", "example", "待替换", "todo")


def normalize_ch4_channel(value: object) -> str:
    text = str(value).strip()
    if not text:
        return ""
    return CH4_CHANNEL_ALIASES.get(text.lower(), CH4_CHANNEL_ALIASES.get(text, text))


def build_ch4_data_from_legacy_parser(xlsx: Path) -> Ch4Data:
    available_sheets = set(get_workbook_sheet_names(xlsx))
    channel_specs: List[Tuple[str, str, str]] = [
        ("医院品类", "医院端", "hospital"),
        ("药店品类", "药店端", "drugstore"),
        ("线上品类", "线上端", "online"),
    ]
    parsed_frames: Dict[str, pd.DataFrame] = {}
    discovered_quarters: List[str] = []
    for sheet_name, channel_label, column in channel_specs:
        if sheet_name not in available_sheets:
            print(f"警告：缺少{channel_label}季度sheet，后续将按 0 补齐。")
            continue
        df = parse_category_sheet(xlsx, sheet_name).rename(columns={"sales": column})
        parsed_frames[column] = df
        discovered_quarters.extend(df["quarter"].astype(str).tolist())

    if not parsed_frames:
        raise ValueError(f"{xlsx.name} 未找到任何渠道季度销售额sheet（医院品类/药店品类/线上品类）。")

    quarter_order = sorted({q for q in discovered_quarters if QUARTER_RE.match(str(q))}, key=qkey)
    for _, _, column in channel_specs:
        if column not in parsed_frames:
            parsed_frames[column] = missing_category_sheet_frame(quarter_order, column)

    quarterly = parsed_frames["hospital"].merge(parsed_frames["drugstore"], on="quarter", how="outer").merge(parsed_frames["online"], on="quarter", how="outer")
    for col in ["hospital", "drugstore", "online"]:
        quarterly[col] = pd.to_numeric(quarterly[col], errors="coerce").fillna(0.0)
    quarterly["total"] = quarterly["hospital"] + quarterly["drugstore"] + quarterly["online"]
    quarterly["year"] = quarterly["quarter"].apply(year_of_quarter)
    quarterly["qk"] = quarterly["quarter"].apply(qkey)
    quarterly = quarterly.sort_values("qk").reset_index(drop=True)

    annual = (
        quarterly.groupby("year", as_index=False)[["hospital", "drugstore", "online", "total"]]
        .sum()
        .sort_values("year")
        .reset_index(drop=True)
    )

    latest_q = quarterly.iloc[-1]["quarter"]
    latest_row = quarterly.iloc[-1]
    latest_share = pd.DataFrame(
        [
            {"channel": "医院端", "sales": latest_row["hospital"], "share_pct": latest_row["hospital"] / latest_row["total"] * 100},
            {"channel": "药店端", "sales": latest_row["drugstore"], "share_pct": latest_row["drugstore"] / latest_row["total"] * 100},
            {"channel": "线上端", "sales": latest_row["online"], "share_pct": latest_row["online"] / latest_row["total"] * 100},
        ]
    )

    prev_q = f"{int(str(latest_q)[:4]) - 1}Q{str(latest_q)[-1]}"
    prev_row = quarterly[quarterly["quarter"] == prev_q]
    if prev_row.empty:
        yoy_latest = pd.DataFrame(
            [
                {"channel": "医院端", "yoy_pct": np.nan},
                {"channel": "药店端", "yoy_pct": np.nan},
                {"channel": "线上端", "yoy_pct": np.nan},
            ]
        )
    else:
        p = prev_row.iloc[0]
        yoy_latest = pd.DataFrame(
            [
                {"channel": "医院端", "yoy_pct": (latest_row["hospital"] - p["hospital"]) / p["hospital"] * 100 if p["hospital"] else np.nan},
                {"channel": "药店端", "yoy_pct": (latest_row["drugstore"] - p["drugstore"]) / p["drugstore"] * 100 if p["drugstore"] else np.nan},
                {"channel": "线上端", "yoy_pct": (latest_row["online"] - p["online"]) / p["online"] * 100 if p["online"] else np.nan},
            ]
        )

    def fallback_top(sheet: str, channel_label: str, reason: str | None = None) -> Tuple[pd.DataFrame, pd.DataFrame]:
        top10 = pd.DataFrame(columns=["rank", "name", "sales"])
        full = pd.DataFrame(columns=["rank", "name"])
        suffix = f"（{reason}）" if reason else ""
        print(f"警告：{sheet}解析失败或源表为空，已标记为缺失数据{suffix}。")
        return top10, full

    def safe_parse_top(sheet: str, channel_label: str) -> Tuple[str, pd.DataFrame, pd.DataFrame]:
        if sheet not in available_sheets:
            top10, full = fallback_top(sheet, channel_label, reason="缺少sheet")
            return str(latest_q), top10, full
        try:
            parsed_q, parsed_top10, parsed_full = parse_top_sheet(xlsx, sheet)
            parsed_top10, parsed_full = collapse_duplicate_top_rows(parsed_q, parsed_top10, parsed_full)
            return parsed_q, parsed_top10, parsed_full
        except Exception:
            top10, full = fallback_top(sheet, channel_label)
            return str(latest_q), top10, full

    h_q, h_top10, h_full = safe_parse_top("医院top", "医院端")
    d_q, d_top10, d_full = safe_parse_top("药店top", "药店端")
    o_q, o_top10, o_full = safe_parse_top("线上top", "线上端")
    if h_q != latest_q or d_q != latest_q or o_q != latest_q:
        pass

    def cr5_value(full_df: pd.DataFrame, channel_total_latest: float, q: str) -> float:
        if q not in full_df.columns or channel_total_latest <= 0:
            return np.nan
        top5_sum = full_df[full_df["rank"] <= 5][q].sum(skipna=True)
        return float(top5_sum / channel_total_latest * 100)

    cr5_latest = pd.DataFrame(
        [
            {"channel": "医院端", "cr5_pct": cr5_value(h_full, float(latest_row["hospital"]), latest_q)},
            {"channel": "药店端", "cr5_pct": cr5_value(d_full, float(latest_row["drugstore"]), latest_q)},
            {"channel": "线上端", "cr5_pct": cr5_value(o_full, float(latest_row["online"]), latest_q)},
        ]
    )

    trend_records = []
    for _, row in quarterly.iterrows():
        q = row["quarter"]
        for channel, full_df, denom_col in [
            ("医院端", h_full, "hospital"),
            ("药店端", d_full, "drugstore"),
            ("线上端", o_full, "online"),
        ]:
            if q in full_df.columns and row[denom_col] > 0:
                top5 = full_df[full_df["rank"] <= 5][q].sum(skipna=True)
                trend_records.append({"quarter": q, "channel": channel, "cr5_pct": float(top5 / row[denom_col] * 100)})
    cr5_trend = pd.DataFrame(trend_records)
    if cr5_trend.empty:
        cr5_trend = pd.DataFrame(columns=["quarter", "channel", "cr5_pct"])
    else:
        cr5_trend["qk"] = cr5_trend["quarter"].apply(qkey)
        cr5_trend = cr5_trend.sort_values(["qk", "channel"]).drop(columns=["qk"]).reset_index(drop=True)

    top_latest_name = {
        "医院端": str(h_top10.iloc[0]["name"]) if not h_top10.empty else "N/A",
        "药店端": str(d_top10.iloc[0]["name"]) if not d_top10.empty else "N/A",
        "线上端": str(o_top10.iloc[0]["name"]) if not o_top10.empty else "N/A",
    }

    return Ch4Data(
        quarterly=quarterly,
        annual=annual,
        latest_quarter=str(latest_q),
        latest_share=latest_share,
        yoy_latest=yoy_latest,
        top_hospital=h_top10,
        top_drugstore=d_top10,
        top_online=o_top10,
        cr5_latest=cr5_latest,
        cr5_trend=cr5_trend,
        top_latest_name=top_latest_name,
    )


def write_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    backup_if_exists(path)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def build_ch4_workbook_preview_lines(xlsx: Path, max_preview_rows: int = 4, max_preview_cols: int = 12) -> List[str]:
    wb = load_workbook(xlsx, read_only=True, data_only=False)
    lines = [
        "【第四章Excel工作簿预览】",
        f"文件：{xlsx.name}",
    ]
    for ws in wb.worksheets:
        lines.append(f"- Sheet={ws.title} | max_row={ws.max_row} | max_col={ws.max_column}")
        shown = 0
        for ridx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            vals = list(row[:max_preview_cols])
            if not any(v not in (None, "") for v in vals):
                continue
            preview = " | ".join([str(v).strip() if v is not None else "" for v in vals])
            lines.append(f"  row{ridx}: {preview}")
            shown += 1
            if shown >= max_preview_rows:
                break
        if shown == 0:
            lines.append("  非空预览：空sheet或未保存到单元格")
    return lines


def build_ch4_codex_extract_template(xlsx: Path) -> Dict[str, object]:
    wb = load_workbook(xlsx, read_only=True, data_only=False)
    sheet_names = [ws.title for ws in wb.worksheets]
    return {
        "schema_version": "ch4_codex_extract_v1",
        "disease": DISEASE_NAME,
        "source_workbook": xlsx.name,
        "available_sheets": sheet_names,
        "latest_quarter": "",
        "sheet_mapping": {
            "hospital_category": {"sheet": "医院品类", "status": "待确认", "header_rows": "", "note": ""},
            "hospital_top": {"sheet": "医院top", "status": "待确认", "header_rows": "", "note": ""},
            "drugstore_category": {"sheet": "药店品类", "status": "待确认", "header_rows": "", "note": ""},
            "drugstore_top": {"sheet": "药店top", "status": "待确认", "header_rows": "", "note": ""},
            "online_category": {"sheet": "线上品类", "status": "待确认", "header_rows": "", "note": ""},
            "online_top": {"sheet": "线上top", "status": "待确认", "header_rows": "", "note": ""},
        },
        "tables": {
            "quarterly_channel": [
                {"quarter": "2025Q3", "hospital": 0, "drugstore": 0, "online": 0}
            ],
            "top10_hospital": [
                {"rank": 1, "name": "示例通用名", "sales": 0}
            ],
            "top10_drugstore": [],
            "top10_online": [],
            "cr5_latest": [
                {"channel": "医院端", "cr5_pct": 0},
                {"channel": "药店端", "cr5_pct": 0},
                {"channel": "线上端", "cr5_pct": 0}
            ],
            "cr5_trend": [
                {"quarter": "2025Q3", "channel": "医院端", "cr5_pct": 0}
            ]
        },
        "notes": [
            "由当前 Codex 会话读取第四章 Excel 后填充。",
            "禁止杜撰数值；所有数字必须可回溯到工作簿单元格区域。",
            "latest_quarter 必填，且必须等于 quarterly_channel 中最后一个季度。",
            "quarterly_channel 必须一季度一行、季度唯一，不得重复。",
            "top10_* 仅保留最新季度 Top10，rank 必须唯一且限定为 1-10，不得保留示例占位值。",
            "cr5_latest/cr5_trend 的 channel 固定使用 医院端/药店端/线上端。",
            "若某 top sheet 为空或未保存到单元格，请保留空数组，并在 sheet_mapping 与 notes 中明确说明。"
        ]
    }


def build_ch4_codex_prompt(xlsx: Path) -> str:
    preview_lines = build_ch4_workbook_preview_lines(xlsx)
    extract_path = OUT_ROOT / "ch04_codex_extract.json"
    template_path = OUT_ROOT / "ch04_codex_extract_template.json"
    lines = [
        "【第四章 Codex 数据提取任务】",
        f"医学主题：{DISEASE_NAME}",
        f"源文件：{xlsx.name}",
        f"目标输出：{extract_path}",
        f"参考模板：{template_path}",
        "",
        "请由当前 Codex 会话读取第四章 Excel，并生成结构化 JSON，供脚本后续画图与装配使用。",
        "",
        "强制要求：",
        "1) 只允许基于工作簿中的真实数据填充，禁止推测、补造或平滑。",
        "2) 必须先识别 sheet 对应关系与表头层级，再输出标准化结果。",
        "3) latest_quarter 必填，且必须等于 quarterly_channel 中最后一个季度。",
        "4) quarterly_channel 必须覆盖医院/药店/线上三端季度销售额；一季度仅一行，季度值唯一，不得重复。",
        "5) top10_hospital / top10_drugstore / top10_online 仅保留最新季度 Top10；rank 必须唯一且限制在 1-10；若源 sheet 为空，则保留空数组并写明原因。",
        "6) cr5_latest / cr5_trend 的 channel 固定使用“医院端 / 药店端 / 线上端”；cr5_trend 中 quarter+channel 组合不得重复。",
        "7) 不得保留模板示例值或占位名（如“示例通用名”）。",
        "8) 输出 JSON 必须符合模板字段，不要新增随意字段名。",
        "",
        "建议检查项：",
        "- 哪一行是排名行，哪一行是季度行",
        "- 销售额块与增长率/份额块的列边界是否分离",
        "- 最新季度是否与季度趋势表最后一期一致",
        "- Top10 数值是否取自最新季度列，而不是整表首列或年汇总列",
        "- 若存在重复季度、重复 rank、重复 quarter+channel，说明提取有误，需要重做",
        "- 若 Excel 可视界面有数据但 openpyxl 预览为空，请在 notes 说明“工作簿未保存到单元格”或“需另存后重试”",
        "",
        "输出完成后，再运行：",
        f"python scripts/run_pipeline.py --topic \"{DISEASE_NAME}\"",
        "",
        *preview_lines,
    ]
    return "\n".join(lines)


def write_ch4_codex_helper_files(xlsx: Path) -> None:
    template = build_ch4_codex_extract_template(xlsx)
    write_json(OUT_ROOT / "ch04_codex_extract_template.json", template)
    write_text(OUT_ROOT / "ch04_codex_prompt.txt", build_ch4_codex_prompt(xlsx) + "\n")
    write_text(OUT_ROOT / "ch04_workbook_preview.txt", "\n".join(build_ch4_workbook_preview_lines(xlsx)) + "\n")


def write_ch4_codex_review_files(raw: Dict[str, object], latest_quarter: str) -> None:
    mapping = raw.get("sheet_mapping", {})
    notes = raw.get("notes", [])
    map_lines = [
        "【第四章 Sheet 映射】",
        f"文件：{str(raw.get('source_workbook', EXCEL_PATH.name)).strip() or EXCEL_PATH.name}",
        f"最新季度：{latest_quarter}",
        "",
    ]
    if isinstance(mapping, dict) and mapping:
        for key, value in mapping.items():
            if isinstance(value, dict):
                map_lines.append(
                    f"- {key}: sheet={value.get('sheet', '')} | status={value.get('status', '')} | header_rows={value.get('header_rows', '')} | note={value.get('note', '')}"
                )
            else:
                map_lines.append(f"- {key}: {value}")
    else:
        map_lines.append("- 未提供 sheet_mapping")
    write_text(OUT_ROOT / "ch04_sheet_map.txt", "\n".join(map_lines) + "\n")

    review_lines = [
        "【第四章 Codex 提取审阅】",
        f"schema_version：{raw.get('schema_version', 'N/A')}",
        f"source_workbook：{raw.get('source_workbook', EXCEL_PATH.name)}",
        f"latest_quarter：{latest_quarter}",
        "",
        "【提取备注】",
    ]
    if isinstance(notes, list) and notes:
        for note in notes:
            review_lines.append(f"- {note}")
    else:
        review_lines.append("- 无")
    write_text(OUT_ROOT / "ch04_codex_review.txt", "\n".join(review_lines) + "\n")


def _extract_table(tables: Dict[str, object], key: str, columns: List[str], allow_empty: bool = True) -> pd.DataFrame:
    raw = tables.get(key, [])
    if raw is None:
        raw = []
    if not isinstance(raw, list):
        raise ValueError(f"ch04_codex_extract.json 中的 {key} 必须为数组。")
    if not raw:
        if allow_empty:
            return pd.DataFrame(columns=columns)
        raise ValueError(f"ch04_codex_extract.json 缺少必填表：{key}。")
    df = pd.DataFrame(raw)
    missing = [c for c in columns if c not in df.columns]
    if missing:
        raise ValueError(f"ch04_codex_extract.json 中的 {key} 缺少字段：{', '.join(missing)}。")
    return df[columns].copy()


def build_ch4_data_from_codex_extract(extract_path: Path) -> Ch4Data:
    raw = json.loads(extract_path.read_text(encoding="utf-8"))
    if not isinstance(raw, dict):
        raise ValueError("ch04_codex_extract.json 顶层必须为对象。")
    schema_version = str(raw.get("schema_version", "")).strip()
    if schema_version != "ch4_codex_extract_v1":
        raise ValueError(f"schema_version 不正确：{schema_version or '空'}")
    topic_name = str(raw.get("topic", "")).strip() or str(raw.get("disease", "")).strip()
    if topic_name != DISEASE_NAME:
        raise ValueError(f"topic 不匹配：当前运行主题为 {DISEASE_NAME}，JSON 中为 {topic_name or '空'}")
    source_workbook = str(raw.get("source_workbook", "")).strip()
    if source_workbook != EXCEL_PATH.name:
        raise ValueError(f"source_workbook 不匹配：当前 Excel 为 {EXCEL_PATH.name}，JSON 中为 {source_workbook or '空'}")
    tables = raw.get("tables", {})
    if not isinstance(tables, dict):
        raise ValueError("ch04_codex_extract.json 缺少 tables 对象。")

    quarterly = _extract_table(tables, "quarterly_channel", ["quarter", "hospital", "drugstore", "online"], allow_empty=False)
    quarterly["quarter"] = quarterly["quarter"].astype(str).str.strip()
    bad_quarters = [q for q in quarterly["quarter"].tolist() if not QUARTER_RE.match(q)]
    if bad_quarters:
        raise ValueError("quarterly_channel 存在非法季度值：" + ", ".join(sorted(set(bad_quarters))))
    for col in ["hospital", "drugstore", "online"]:
        quarterly[col] = pd.to_numeric(quarterly[col], errors="coerce").fillna(0.0)
    dup_quarters = quarterly.loc[quarterly["quarter"].duplicated(keep=False), "quarter"].astype(str).tolist()
    if dup_quarters:
        raise ValueError("quarterly_channel 存在重复季度：" + ", ".join(sorted(set(dup_quarters))))
    quarterly["qk"] = quarterly["quarter"].apply(qkey)
    quarterly = quarterly.sort_values("qk").reset_index(drop=True)
    quarterly["total"] = quarterly["hospital"] + quarterly["drugstore"] + quarterly["online"]
    quarterly["year"] = quarterly["quarter"].apply(year_of_quarter)

    latest_q_raw = str(raw.get("latest_quarter", "")).strip()
    if not latest_q_raw:
        raise ValueError("latest_quarter 不能为空。")
    if latest_q_raw not in set(quarterly["quarter"].tolist()):
        raise ValueError(f"latest_quarter={latest_q_raw} 未出现在 quarterly_channel 中。")
    latest_q = str(quarterly.iloc[-1]["quarter"])
    if latest_q_raw != latest_q:
        raise ValueError(f"latest_quarter={latest_q_raw} 与 quarterly_channel 最后一个季度 {latest_q} 不一致。")
    latest_row = quarterly[quarterly["quarter"] == latest_q].iloc[-1]

    annual = (
        quarterly.groupby("year", as_index=False)[["hospital", "drugstore", "online", "total"]]
        .sum()
        .sort_values("year")
        .reset_index(drop=True)
    )

    latest_share = pd.DataFrame(
        [
            {"channel": "医院端", "sales": latest_row["hospital"], "share_pct": latest_row["hospital"] / latest_row["total"] * 100 if latest_row["total"] else np.nan},
            {"channel": "药店端", "sales": latest_row["drugstore"], "share_pct": latest_row["drugstore"] / latest_row["total"] * 100 if latest_row["total"] else np.nan},
            {"channel": "线上端", "sales": latest_row["online"], "share_pct": latest_row["online"] / latest_row["total"] * 100 if latest_row["total"] else np.nan},
        ]
    )

    prev_q = f"{int(str(latest_q)[:4]) - 1}Q{str(latest_q)[-1]}"
    prev_row = quarterly[quarterly["quarter"] == prev_q]
    if prev_row.empty:
        yoy_latest = pd.DataFrame(
            [
                {"channel": "医院端", "yoy_pct": np.nan},
                {"channel": "药店端", "yoy_pct": np.nan},
                {"channel": "线上端", "yoy_pct": np.nan},
            ]
        )
    else:
        p = prev_row.iloc[0]
        yoy_latest = pd.DataFrame(
            [
                {"channel": "医院端", "yoy_pct": (latest_row["hospital"] - p["hospital"]) / p["hospital"] * 100 if p["hospital"] else np.nan},
                {"channel": "药店端", "yoy_pct": (latest_row["drugstore"] - p["drugstore"]) / p["drugstore"] * 100 if p["drugstore"] else np.nan},
                {"channel": "线上端", "yoy_pct": (latest_row["online"] - p["online"]) / p["online"] * 100 if p["online"] else np.nan},
            ]
        )

    def _top10(key: str) -> pd.DataFrame:
        df = _extract_table(tables, key, ["rank", "name", "sales"], allow_empty=True)
        if df.empty:
            return pd.DataFrame(columns=["rank", "name", "sales"])
        df["rank"] = pd.to_numeric(df["rank"], errors="coerce")
        df["sales"] = pd.to_numeric(df["sales"], errors="coerce")
        df["name"] = df["name"].astype(str).str.strip()
        df = df.dropna(subset=["rank", "sales"])
        df = df[df["name"] != ""]
        if df.empty:
            return pd.DataFrame(columns=["rank", "name", "sales"])
        bad_rank_rows = df[(df["rank"] < 1) | (df["rank"] > 10) | ((df["rank"] % 1) != 0)]
        if not bad_rank_rows.empty:
            raise ValueError(f"{key} 存在非法 rank，仅允许 1-10 的整数。")
        df["rank"] = df["rank"].astype(int)
        dup_ranks = df.loc[df["rank"].duplicated(keep=False), "rank"].astype(int).tolist()
        if dup_ranks:
            raise ValueError(f"{key} 存在重复 rank：" + ", ".join(str(x) for x in sorted(set(dup_ranks))))
        dup_names = df.loc[df["name"].duplicated(keep=False), "name"].astype(str).tolist()
        if dup_names:
            raise ValueError(f"{key} 存在重复通用名：" + "、".join(sorted(set(dup_names))))
        placeholder_names = [
            name for name in df["name"].astype(str).tolist()
            if any(token in name.lower() for token in CH4_PLACEHOLDER_NAME_TERMS)
        ]
        if placeholder_names:
            raise ValueError(f"{key} 仍包含模板占位名：" + "、".join(sorted(set(placeholder_names))))
        if len(df) > 10:
            raise ValueError(f"{key} 超过 10 行，请仅保留最新季度 Top10。")
        return df.sort_values("rank").reset_index(drop=True)

    top_hospital = _top10("top10_hospital")
    top_drugstore = _top10("top10_drugstore")
    top_online = _top10("top10_online")

    cr5_latest = _extract_table(tables, "cr5_latest", ["channel", "cr5_pct"], allow_empty=True)
    if cr5_latest.empty:
        def _cr5_from_top(df: pd.DataFrame, total_value: float) -> float:
            if df.empty or total_value <= 0:
                return np.nan
            top5_sum = df[df["rank"] <= 5]["sales"].sum(skipna=True)
            return float(top5_sum / total_value * 100)

        cr5_latest = pd.DataFrame(
            [
                {"channel": "医院端", "cr5_pct": _cr5_from_top(top_hospital, float(latest_row["hospital"]))},
                {"channel": "药店端", "cr5_pct": _cr5_from_top(top_drugstore, float(latest_row["drugstore"]))},
                {"channel": "线上端", "cr5_pct": _cr5_from_top(top_online, float(latest_row["online"]))},
            ]
        )
    else:
        cr5_latest["channel"] = cr5_latest["channel"].apply(normalize_ch4_channel)
        cr5_latest["cr5_pct"] = pd.to_numeric(cr5_latest["cr5_pct"], errors="coerce")
        bad_channels = [c for c in cr5_latest["channel"].astype(str).tolist() if c not in CH4_CHANNEL_ORDER]
        if bad_channels:
            raise ValueError("cr5_latest 存在非法 channel：" + "、".join(sorted(set(bad_channels))))
        dup_channels = cr5_latest.loc[cr5_latest["channel"].duplicated(keep=False), "channel"].astype(str).tolist()
        if dup_channels:
            raise ValueError("cr5_latest 存在重复 channel：" + "、".join(sorted(set(dup_channels))))
        cr5_latest_map = dict(zip(cr5_latest["channel"], cr5_latest["cr5_pct"]))
        cr5_latest = pd.DataFrame(
            [{"channel": channel, "cr5_pct": cr5_latest_map.get(channel, np.nan)} for channel in CH4_CHANNELS]
        )

    cr5_trend = _extract_table(tables, "cr5_trend", ["quarter", "channel", "cr5_pct"], allow_empty=True)
    if cr5_trend.empty:
        cr5_trend = pd.DataFrame(columns=["quarter", "channel", "cr5_pct"])
    else:
        cr5_trend["quarter"] = cr5_trend["quarter"].astype(str).str.strip()
        cr5_trend["channel"] = cr5_trend["channel"].apply(normalize_ch4_channel)
        cr5_trend["cr5_pct"] = pd.to_numeric(cr5_trend["cr5_pct"], errors="coerce")
        cr5_trend = cr5_trend.dropna(subset=["quarter"]).copy()
        cr5_trend = cr5_trend[cr5_trend["quarter"].apply(lambda x: bool(QUARTER_RE.match(x)))]
        if not cr5_trend.empty:
            bad_trend_channels = [c for c in cr5_trend["channel"].astype(str).tolist() if c not in CH4_CHANNEL_ORDER]
            if bad_trend_channels:
                raise ValueError("cr5_trend 存在非法 channel：" + "、".join(sorted(set(bad_trend_channels))))
            dup_pairs = cr5_trend.loc[
                cr5_trend.duplicated(subset=["quarter", "channel"], keep=False),
                ["quarter", "channel"],
            ]
            if not dup_pairs.empty:
                dup_text = [f"{row.quarter}/{row.channel}" for row in dup_pairs.drop_duplicates().itertuples(index=False)]
                raise ValueError("cr5_trend 存在重复 quarter+channel：" + "、".join(dup_text))
            cr5_trend["qk"] = cr5_trend["quarter"].apply(qkey)
            cr5_trend["channel_order"] = cr5_trend["channel"].map(CH4_CHANNEL_ORDER)
            cr5_trend = cr5_trend.sort_values(["qk", "channel_order"]).drop(columns=["qk", "channel_order"]).reset_index(drop=True)

    def _cr5_from_top(df: pd.DataFrame, total_value: float) -> float:
        if df.empty or total_value <= 0:
            return np.nan
        top5_sum = df[df["rank"] <= 5]["sales"].sum(skipna=True)
        return float(top5_sum / total_value * 100)

    for channel, top_df, total_value in [
        ("医院端", top_hospital, float(latest_row["hospital"])),
        ("药店端", top_drugstore, float(latest_row["drugstore"])),
        ("线上端", top_online, float(latest_row["online"])),
    ]:
        expected_cr5 = _cr5_from_top(top_df, total_value)
        provided_row = cr5_latest[cr5_latest["channel"] == channel]
        provided_cr5 = float(provided_row.iloc[0]["cr5_pct"]) if not provided_row.empty and pd.notna(provided_row.iloc[0]["cr5_pct"]) else np.nan
        if pd.notna(expected_cr5) and pd.notna(provided_cr5) and abs(provided_cr5 - expected_cr5) > 0.5:
            raise ValueError(
                f"cr5_latest 中 {channel} 与 Top5/季度总额计算结果不一致：填报={provided_cr5:.2f}，计算={expected_cr5:.2f}"
            )

    if not cr5_trend.empty:
        latest_trend = cr5_trend[cr5_trend["quarter"] == latest_q]
        latest_trend_map = dict(zip(latest_trend["channel"], latest_trend["cr5_pct"]))
        for channel in CH4_CHANNELS:
            trend_value = latest_trend_map.get(channel, np.nan)
            latest_value = cr5_latest.loc[cr5_latest["channel"] == channel, "cr5_pct"].iloc[0]
            if pd.notna(trend_value) and pd.notna(latest_value) and abs(float(trend_value) - float(latest_value)) > 0.5:
                raise ValueError(
                    f"cr5_trend 中最新季度 {latest_q} 的 {channel} 与 cr5_latest 不一致：趋势={float(trend_value):.2f}，最新值={float(latest_value):.2f}"
                )

    top_latest_name = {
        "医院端": str(top_hospital.iloc[0]["name"]) if not top_hospital.empty else "N/A",
        "药店端": str(top_drugstore.iloc[0]["name"]) if not top_drugstore.empty else "N/A",
        "线上端": str(top_online.iloc[0]["name"]) if not top_online.empty else "N/A",
    }

    write_ch4_codex_review_files(raw, latest_q)

    return Ch4Data(
        quarterly=quarterly,
        annual=annual,
        latest_quarter=str(latest_q),
        latest_share=latest_share,
        yoy_latest=yoy_latest,
        top_hospital=top_hospital,
        top_drugstore=top_drugstore,
        top_online=top_online,
        cr5_latest=cr5_latest,
        cr5_trend=cr5_trend,
        top_latest_name=top_latest_name,
    )


def build_ch4_data(xlsx: Path) -> Ch4Data:
    write_ch4_codex_helper_files(xlsx)
    extract_path = OUT_ROOT / "ch04_codex_extract.json"
    if not extract_path.exists():
        raise RuntimeError(
            "未检测到第四章结构化提取结果 ch04_codex_extract.json。"
            "请先由当前 Codex 会话读取第四章 Excel，参考 ch04_codex_prompt.txt 与 ch04_codex_extract_template.json 生成该文件，再重新执行。"
        )
    try:
        return build_ch4_data_from_codex_extract(extract_path)
    except Exception as exc:
        raise RuntimeError(
            "第四章结构化提取结果无效。请先检查 ch04_codex_extract.json、ch04_sheet_map.txt 与 ch04_codex_review.txt 后再重试。"
        ) from exc


def write_ch4_profile_files(ch4: Ch4Data) -> None:
    missing_top_channels = []
    if ch4.top_hospital.empty:
        missing_top_channels.append("医院端")
    if ch4.top_drugstore.empty:
        missing_top_channels.append("药店端")
    if ch4.top_online.empty:
        missing_top_channels.append("线上端")
    competition_support = "可直接支持（top通用名、CR5）"
    top_support = "可直接支持（最新季度Top10）"
    if missing_top_channels:
        channels = "、".join(missing_top_channels)
        competition_support = f"部分支持（{channels}TOP源表为空，相关Top/CR5仅对有数据渠道输出）"
        top_support = f"部分支持（{channels}未提供最新季度Top表）"

    profile_lines = [
        "【第四章Excel剖面】",
        f"文件：{EXCEL_PATH.name}",
        "第四章标准化来源：ch04_codex_extract.json（由 Codex 读取 Excel 后生成）",
        "Sheet映射：详见 ch04_sheet_map.txt",
        f"季度范围：{ch4.quarterly['quarter'].iloc[0]} - {ch4.quarterly['quarter'].iloc[-1]}",
        f"记录条数（渠道季度）：{len(ch4.quarterly)}",
        "粒度：季度，金额单位：万元，口径：米内网终端销售额",
        "缺失与异常：详见 ch04_codex_review.txt；脚本仅消费结构化提取结果，不再直接猜测 Excel top 表结构。",
        "",
        "【可支撑分析】",
        "1) 规模趋势：可直接支持（季度+年度）",
        "2) 渠道结构：可直接支持（医院/药店/线上）",
        f"3) 竞争格局：{competition_support}",
        f"4) 重点品种：{top_support}",
        "5) 区域分析：原始表未提供地区维度，采用渠道与品种结构替代说明。",
    ]
    write_text(OUT_ROOT / "ch04_excel_profile.txt", "\n".join(profile_lines))

    dict_lines = [
        "【第四章数据字典】",
        "quarter：季度（YYYYQn），用于趋势和同比分析",
        "hospital/drugstore/online：三端销售额（万元）",
        "total：三端合计销售额（万元）",
        "share_pct：渠道份额（%）",
        "yoy_pct：同季度同比增速（%）",
        "rank：通用名排名（top表）",
        "name：通用名",
        "sales：最新季度销售额（万元）",
        "cr5_pct：前五通用名销售额占该渠道总额比例（%）",
        "可支撑图表：规模趋势、结构占比、同比对比、top品种、CR5对比/趋势",
    ]
    write_text(OUT_ROOT / "ch04_data_dictionary.txt", "\n".join(dict_lines))

    agg_path = OUT_ROOT / "ch04_agg_tables.xlsx"
    backup_if_exists(agg_path)
    with pd.ExcelWriter(agg_path, engine="openpyxl") as writer:
        ch4.quarterly.to_excel(writer, index=False, sheet_name="quarterly_channel")
        ch4.annual.to_excel(writer, index=False, sheet_name="annual_channel")
        ch4.latest_share.to_excel(writer, index=False, sheet_name="latest_share")
        ch4.yoy_latest.to_excel(writer, index=False, sheet_name="latest_yoy")
        ch4.top_hospital.to_excel(writer, index=False, sheet_name="top10_hospital")
        ch4.top_drugstore.to_excel(writer, index=False, sheet_name="top10_drugstore")
        ch4.top_online.to_excel(writer, index=False, sheet_name="top10_online")
        ch4.cr5_latest.to_excel(writer, index=False, sheet_name="cr5_latest")
        ch4.cr5_trend.to_excel(writer, index=False, sheet_name="cr5_trend")


def build_ch4_narrative_brief(ch4: Ch4Data) -> str:
    latest_row = ch4.quarterly.iloc[-1]
    latest_quarter = str(latest_row["quarter"])
    latest_total = float(latest_row["hospital"]) + float(latest_row["drugstore"]) + float(latest_row["online"])

    def pct(part: float, total: float) -> float:
        if total <= 0:
            return 0.0
        return part / total * 100.0

    latest_hospital = float(latest_row["hospital"])
    latest_drugstore = float(latest_row["drugstore"])
    latest_online = float(latest_row["online"])

    latest_share = {
        "hospital": pct(latest_hospital, latest_total),
        "drugstore": pct(latest_drugstore, latest_total),
        "online": pct(latest_online, latest_total),
    }

    latest_year = int(str(latest_quarter)[:4])
    prev_same_quarter = f"{latest_year - 1}{str(latest_quarter)[4:]}"
    quarterly_lookup = {str(row["quarter"]): row for _, row in ch4.quarterly.iterrows()}
    prev_row = quarterly_lookup.get(prev_same_quarter)
    yoy_lines: List[str] = []
    if prev_row is not None:
        prev_total = float(prev_row["hospital"]) + float(prev_row["drugstore"]) + float(prev_row["online"])

        def yoy(now: float, then: float) -> float:
            if then == 0:
                return 0.0
            return (now - then) / then * 100.0

        yoy_lines = [
            f"- Latest quarter YoY total: {yoy(latest_total, prev_total):.2f}%",
            f"- Latest quarter YoY hospital: {yoy(latest_hospital, float(prev_row['hospital'])):.2f}%",
            f"- Latest quarter YoY drugstore: {yoy(latest_drugstore, float(prev_row['drugstore'])):.2f}%",
            f"- Latest quarter YoY online: {yoy(latest_online, float(prev_row['online'])):.2f}%",
        ]

    first_row = ch4.quarterly.iloc[0]
    first_total = float(first_row["hospital"]) + float(first_row["drugstore"]) + float(first_row["online"])
    quarter_span = max(len(ch4.quarterly) - 1, 0)
    years = quarter_span / 4.0 if quarter_span > 0 else 0.0
    total_cagr = (((latest_total / first_total) ** (1.0 / years)) - 1.0) * 100.0 if (first_total > 0 and years > 0) else 0.0

    annual_lines: List[str] = []
    if not ch4.annual.empty:
        for _, row in ch4.annual.iterrows():
            year_value = str(row["year"])
            hospital = float(row["hospital"])
            drugstore = float(row["drugstore"])
            online = float(row["online"])
            total = hospital + drugstore + online
            annual_lines.append(
                f"- {year_value}: total={total:.1f}, hospital_share={pct(hospital, total):.1f}%, drugstore_share={pct(drugstore, total):.1f}%, online_share={pct(online, total):.1f}%"
            )

    def top_lines(label: str, df: pd.DataFrame) -> List[str]:
        if df.empty:
            return [f"- {label}: no top data"]
        rows = []
        for _, row in df.head(3).iterrows():
            rows.append(f"{int(row['rank'])}.{str(row['name']).strip()} ({float(row['sales']):.1f})")
        return [f"- {label}: " + "; ".join(rows)]

    cr5_lines: List[str] = []
    if not ch4.cr5_latest.empty:
        for _, row in ch4.cr5_latest.iterrows():
            cr5_lines.append(f"- {str(row['channel']).strip()}: CR5={float(row['cr5_pct']):.2f}%")

    lines = [
        "[Chapter 4 Narrative Brief]",
        f"Topic: {DISEASE_NAME}",
        f"Latest quarter: {latest_quarter}",
        f"Latest quarter total sales (万元): {latest_total:.1f}",
        f"- Hospital: {latest_hospital:.1f} ({latest_share['hospital']:.2f}%)",
        f"- Drugstore: {latest_drugstore:.1f} ({latest_share['drugstore']:.2f}%)",
        f"- Online: {latest_online:.1f} ({latest_share['online']:.2f}%)",
        f"- Long-window quarterly CAGR: {total_cagr:.2f}%",
        "",
        "[Latest YoY]",
    ]
    lines.extend(yoy_lines if yoy_lines else ["- Same quarter last year not available"])
    lines.extend(
        [
            "",
            "[Annual Structure]",
        ]
    )
    lines.extend(annual_lines if annual_lines else ["- Annual data not available"])
    lines.extend(
        [
            "",
            "[Top Products]",
        ]
    )
    lines.extend(top_lines("Hospital", ch4.top_hospital))
    lines.extend(top_lines("Drugstore", ch4.top_drugstore))
    lines.extend(top_lines("Online", ch4.top_online))
    lines.extend(
        [
            "",
            "[CR5]",
        ]
    )
    lines.extend(cr5_lines if cr5_lines else ["- CR5 data not available"])
    lines.extend(
        [
            "",
            "[Suggested Writing Angles]",
            "- Treat the market as hospital-led, with drugstore/online acting as follow-up and continuation channels.",
            "- Explain short-term volatility with channel structure shifts instead of only using total sales direction.",
            "- Use top-product and CR5 structure to discuss concentration, not just scale.",
            "- Do not invent region-level conclusions because the workbook does not provide region data.",
            "- Keep every chapter-4 claim inside Excel-derived scope only.",
        ]
    )
    return "\n".join(lines)


def load_existing_text_bundle_partial(specs: List[BlockSpec]) -> Tuple[Dict[str, str], str]:
    pattern = re.compile(
        r"\[\[BLOCK_ID=(?P<id>[^\]]+)\]\]\n(?P<title>[^\n]*)\n(?P<body>.*?)\n\[\[END_BLOCK_ID=(?P=id)\]\]",
        re.S,
    )
    out = {s.block_id: "" for s in specs}
    for ch in range(1, 8):
        path = OUT_ROOT / f"ch0{ch}.txt"
        if not path.exists():
            continue
        try:
            raw = path.read_text(encoding="utf-8")
        except Exception:
            continue
        for match in pattern.finditer(raw):
            block_id = str(match.group("id")).strip()
            if block_id in out:
                out[block_id] = match.group("body").strip()
    summary_path = OUT_ROOT / "summary.txt"
    summary_text = summary_path.read_text(encoding="utf-8").strip() if summary_path.exists() else ""
    return out, summary_text


def build_codex_gap_panel(specs: List[BlockSpec], block_text: Dict[str, str], summary_text: str) -> str:
    chapter_to_specs: Dict[int, List[BlockSpec]] = {}
    for spec in specs:
        chapter_to_specs.setdefault(spec.chapter, []).append(spec)

    current_total = sum(len(re.sub(r"\s+", "", str(block_text.get(s.block_id, "")))) for s in specs) + len(re.sub(r"\s+", "", summary_text))
    summary_chars = len(re.sub(r"\s+", "", summary_text))
    block_rows: List[Tuple[int, str, int, int, int]] = []
    lines = [
        "[Codex Gap Panel]",
        f"Topic: {DISEASE_NAME}",
        f"Current total chars (chapters + summary): {current_total}",
        f"Target total chars: 30000-34000",
        f"Current summary chars: {summary_chars}",
        "",
        "[Chapter / Block Gaps]",
    ]
    for chapter in range(1, 8):
        ch_specs = chapter_to_specs.get(chapter, [])
        ch_chars = sum(len(re.sub(r"\s+", "", str(block_text.get(s.block_id, "")))) for s in ch_specs)
        ch_floor = int(CHAPTER_MIN_CHARS.get(chapter, 0))
        strict_gap = chapter_char_shortfall(chapter, ch_chars)
        rewrite_gap = max(0, strict_gap - CHAPTER_CHAR_TOLERANCE)
        lines.append(
            f"- Chapter {chapter}: current={ch_chars}, floor={ch_floor}, tolerance={CHAPTER_CHAR_TOLERANCE}, strict_gap={strict_gap}, rewrite_gap={rewrite_gap}"
        )
        for spec in ch_specs:
            current_chars = len(re.sub(r"\s+", "", str(block_text.get(spec.block_id, ""))))
            block_gap = max(0, int(spec.target_chars) - current_chars)
            status = "missing" if current_chars == 0 else ("short" if block_gap > 0 else "ok")
            lines.append(
                f"  - {spec.block_id} | current={current_chars} | target={spec.target_chars} | gap={block_gap} | status={status} | figures={spec.fig_ids if spec.fig_ids else 'none'}"
            )
            block_rows.append((block_gap, spec.block_id, current_chars, int(spec.target_chars), spec.chapter))
    lines.extend(
        [
            "",
            "[Priority Blocks]",
        ]
    )
    ranked = [row for row in sorted(block_rows, key=lambda item: item[0], reverse=True) if row[0] > 0]
    if ranked:
        for gap, block_id, current_chars, target_chars, chapter in ranked[:10]:
            lines.append(
                f"- Chapter {chapter} / {block_id}: gap={gap}, current={current_chars}, target={target_chars}"
            )
    else:
        lines.append("- No block gaps detected")
    lines.extend(
        [
            "",
            "[Summary]",
            f"- summary current={summary_chars}, recommended=1200-1500, gap_to_1200={max(0, 1200 - summary_chars)}",
        ]
    )
    return "\n".join(lines)


def build_chapter_precheck(specs: List[BlockSpec], block_text: Dict[str, str], summary_text: str) -> str:
    chapter_to_specs: Dict[int, List[BlockSpec]] = {}
    for spec in specs:
        chapter_to_specs.setdefault(spec.chapter, []).append(spec)

    lines = [
        "[Chapter Precheck]",
        f"Topic: {DISEASE_NAME}",
        "This is a lightweight writing-stage check, not the final QA gate.",
        "",
    ]
    for chapter in range(1, 8):
        ch_specs = chapter_to_specs.get(chapter, [])
        missing_blocks = [s.block_id for s in ch_specs if not str(block_text.get(s.block_id, "")).strip()]
        ch_text = "\n".join(str(block_text.get(s.block_id, "")) for s in ch_specs)
        ch_paras: List[str] = []
        for spec in ch_specs:
            ch_paras.extend(split_paragraphs(str(block_text.get(spec.block_id, ""))))
        ch_chars = len(re.sub(r"\s+", "", ch_text))
        ch_floor = int(CHAPTER_MIN_CHARS.get(chapter, 0))
        cite_cnt = len(re.findall(r"\[\d+\]", ch_text))
        anchored = sum(1 for paragraph in ch_paras if paragraph_has_anchor(paragraph))
        anchor_cov = anchored / len(ch_paras) if ch_paras else 0.0
        ch_dup, _ = sentence_repeat_stats(ch_text)
        medical_hits = sum(1 for pattern in MEDICAL_PATTERNS.values() if re.search(pattern, ch_text)) if chapter <= 3 else -1

        fail_reasons: List[str] = []
        warn_reasons: List[str] = []
        if missing_blocks:
            fail_reasons.append("missing blocks=" + ",".join(missing_blocks))
        strict_gap = chapter_char_shortfall(chapter, ch_chars)
        if strict_gap > CHAPTER_CHAR_TOLERANCE:
            fail_reasons.append(f"chars={ch_chars} below floor by {strict_gap} (> tolerance {CHAPTER_CHAR_TOLERANCE})")
        elif strict_gap > 0:
            warn_reasons.append(f"chars below floor by {strict_gap}, but within tolerance")
        elif ch_chars < ch_floor + 150:
            warn_reasons.append(f"chars only slightly above floor ({ch_chars})")
        if cite_cnt == 0:
            fail_reasons.append("no citations")
        if anchor_cov < 0.70:
            fail_reasons.append(f"anchor coverage too low ({anchor_cov*100:.1f}%)")
        elif anchor_cov < 0.85:
            warn_reasons.append(f"anchor coverage marginal ({anchor_cov*100:.1f}%)")
        if ch_dup >= 4:
            fail_reasons.append(f"sentence duplication too high ({ch_dup})")
        if chapter <= 3 and medical_hits < 2:
            fail_reasons.append(f"medical pattern hits too low ({medical_hits})")

        status = "FAIL" if fail_reasons else ("WARN" if warn_reasons else "PASS")
        lines.append(
            f"- Chapter {chapter} | status={status} | chars={ch_chars} | paragraphs={len(ch_paras)} | citations={cite_cnt} | anchor_cov={anchor_cov*100:.1f}%"
        )
        if chapter <= 3:
            lines.append(f"  medical_pattern_hits={medical_hits}")
        if fail_reasons:
            lines.append("  fail_reasons=" + "; ".join(fail_reasons))
        if warn_reasons:
            lines.append("  warn_reasons=" + "; ".join(warn_reasons))

    summary_chars = len(re.sub(r"\s+", "", summary_text))
    summary_paras = split_paragraphs(summary_text)
    summary_cites = len(re.findall(r"\[\d+\]", summary_text))
    summary_status = "FAIL" if summary_chars == 0 else ("WARN" if summary_chars < 1200 else "PASS")
    lines.extend(
        [
            "",
            f"- Summary | status={summary_status} | chars={summary_chars} | paragraphs={len(summary_paras)} | citations={summary_cites}",
        ]
    )
    return "\n".join(lines)


def write_codex_progress_assets(
    specs: List[BlockSpec],
    block_text: Dict[str, str] | None = None,
    summary_text: str | None = None,
) -> None:
    if block_text is None or summary_text is None:
        partial_block_text, partial_summary_text = load_existing_text_bundle_partial(specs)
        if block_text is None:
            block_text = partial_block_text
        if summary_text is None:
            summary_text = partial_summary_text
    normalized_block_text = {s.block_id: str((block_text or {}).get(s.block_id, "")).strip() for s in specs}
    normalized_summary = str(summary_text or "").strip()
    write_text(OUT_ROOT / CODEX_GAP_PANEL_NAME, build_codex_gap_panel(specs, normalized_block_text, normalized_summary) + "\n")
    write_text(OUT_ROOT / CHAPTER_PRECHECK_NAME, build_chapter_precheck(specs, normalized_block_text, normalized_summary) + "\n")


def build_evidence_and_refs() -> Tuple[str, str]:
    def to_ref(idx: int, title: str, org: str, year: str, url: str, dtype: str = "EB/OL") -> str:
        y = year if re.search(r"(19|20)\d{2}", str(year)) else "2024"
        return f"[{idx}] {org}. {title}[{dtype}]. {y}. {url}"

    if is_cervical_profile():
        evidence = [
            ("E01", "Musculoskeletal conditions", "World Health Organization", "2022", "肌肉骨骼疾病负担及残疾影响", "https://www.who.int/news-room/fact-sheets/detail/musculoskeletal-conditions"),
            ("E02", "颈椎病康复诊疗专家共识", "中国康复医学会", "2021", "分层康复路径、功能评估与复评窗口", "https://guide.medlive.cn/guideline/27562"),
            ("E03", "颈椎病诊疗与手术指征专家共识", "中华医学会骨科相关学组", "2020", "保守治疗、介入与手术分层决策", "https://guide.medlive.cn/guideline/24306"),
            ("E04", "国家基本医疗保险、工伤保险和生育保险药品目录（2024年）", "国家医疗保障局", "2024", "支付规则影响药物可及性与处方结构", "https://www.nhsa.gov.cn/art/2024/11/28/art_53_14887.html"),
            ("E05", "国家医疗质量安全改进目标（骨科相关）", "国家卫生健康委员会", "2024", "医疗质量指标与分级诊疗执行要求", "https://www.nhc.gov.cn/yzygj/s7659/"),
            ("E06", "Diagnosis and Treatment of Cervical Radiculopathy", "North American Spine Society", "2020", "神经根型颈椎病的诊断评估与治疗建议", "https://www.spine.org/ResearchClinicalCare/QualityImprovement/ClinicalGuidelines"),
            ("E07", "Cervical Myelopathy", "American Association of Neurological Surgeons", "2024", "脊髓型颈椎病红旗征与转诊建议", "https://www.aans.org/Patients/Neurosurgical-Conditions-and-Treatments/Cervical-Myelopathy"),
            ("E08", "国家统计数据发布平台（人口与社会）", "国家统计局", "2024", "老龄化、职业久坐与颈椎疾病需求变化", "https://www.stats.gov.cn/"),
            ("E09", "药品说明书修订公告汇总", "中国食品药品检定研究院", "2024", "镇痛、肌松及神经营养类药品标签更新", "https://www.cpi.ac.cn/tggg/ypsmsxdgg/"),
            ("E10", "中国卫生健康统计年鉴及国家统计数据库", "国家卫生健康委员会/国家统计局", "2024", "门急诊就诊量与骨科康复服务供给变化", "https://www.stats.gov.cn/"),
            ("E11", "颈椎病外科治疗与围手术期管理共识", "中华外科相关学组", "2021", "手术适应证、并发症管理与随访标准", "https://guide.medlive.cn/guideline/27112"),
            ("E12", "米内网终端数据口径说明与原始数据文件", "米内网/项目数据", "2025", "医院/药店/线上三端同口径比较基础", f"{EXCEL_PATH.name}"),
            ("E13", "颈肩痛诊疗与慢性疼痛管理共识", "中国疼痛医学相关学会", "2022", "疼痛分层、功能结局与长期管理", "https://guide.medlive.cn/guideline/28605"),
            ("E14", "中医骨伤科颈椎病诊疗指南", "中国中医药相关学会", "2021", "中医辨证分型与针灸推拿干预建议", "https://guide.medlive.cn/guideline/26067"),
            ("E15", "国家药监局法规与政策文件索引", "国家药品监督管理局", "2024", "全生命周期合规要求与监管边界", "https://www.nmpa.gov.cn/xxgk/fgwj/"),
        ]
    elif is_sciatica_profile():
        evidence = [
            ("E01", "Musculoskeletal conditions", "World Health Organization", "2022", "肌肉骨骼疾病负担持续上升，腰背痛与神经根痛管理需求显著", "https://www.who.int/news-room/fact-sheets/detail/musculoskeletal-conditions"),
            ("E02", "Low back pain and sciatica in over 16s: assessment and management", "National Institute for Health and Care Excellence", "2020", "坐骨神经痛诊断、影像时机、保守治疗与转诊路径建议", "https://www.nice.org.uk/guidance/ng59"),
            ("E03", "Lumbar Disc Herniation with Radiculopathy", "North American Spine Society", "2012", "腰椎间盘突出伴神经根病的评估、保守治疗、介入与手术证据", "https://www.spine.org/ResearchClinicalCare/QualityImprovement/ClinicalGuidelines"),
            ("E04", "2025年版国家基本医疗保险、生育保险和工伤保险药品目录和商业健康保险创新药品目录", "国家医疗保障局/人力资源社会保障部", "2025", "支付范围与目录动态调整影响镇痛、辅助治疗和院外可及性", "https://www.nhsa.gov.cn/art/2025/12/7/art_14_18972.html"),
            ("E05", "2025年国家医疗质量安全改进目标", "国家卫生健康委员会", "2025", "强化医疗质量目标管理，推动骨科、康复和疼痛相关质量改进", "https://www.nhc.gov.cn/yzygj/c100068/202503/ad63fb8ce9e24013a68db52049ecc524.shtml"),
            ("E06", "Herniated Disc", "American Association of Neurological Surgeons", "2024", "多数神经根痛可先行保守治疗，但持续症状和神经缺损需升级评估", "https://www.aans.org/patients/conditions-treatments/herniated-disc/"),
            ("E07", "Cauda Equina Syndrome", "American Association of Neurological Surgeons", "2024", "马尾综合征红旗征包括尿潴留、鞍区麻木和进行性肌力下降", "https://www.aans.org/patients/conditions-treatments/cauda-equina-syndrome/"),
            ("E08", "2025年年末人口数及其构成", "国家统计局", "2026", "老龄化、城镇就业与职业负荷结构变化影响腰腿痛与康复需求", "https://www.stats.gov.cn/zt_18555/zthd/lhfw/2026lhzt/2026hgjj/202602/t20260228_1962667.html"),
            ("E09", "药品说明书修订公告专栏", "中国食品药品检定研究院/国家药监局", "2025", "说明书修订与风险提示持续影响镇痛药和辅助治疗药物的合规使用", "https://www.cpi.ac.cn/tggg/ypsmsxdgg/"),
            ("E10", "2024年我国卫生健康事业发展统计公报新闻解读稿", "国家卫生健康委员会", "2025", "门诊、住院和基层服务量变化为骨科疼痛与康复资源配置提供背景", "https://www.nhc.gov.cn/guihuaxxs/c100132/202512/9db8db5488a748c4b1e1068a4a8c4455.shtml"),
            ("E11", "专家解读《中国人群身体活动指南（2021）》", "国家卫生健康委员会", "2021", "身体活动与分层运动处方是慢性腰腿痛长期管理的重要基础", "https://www.nhc.gov.cn/wjw/ftsp/202112/44e8325ad5934eb0b42b2987e3148315.shtml"),
            ("E12", "米内网终端数据口径说明与原始数据文件", "米内网/项目数据", "2025", "医院、药店、线上三端季度销售额与Top品种比较基础", f"{EXCEL_PATH.name}"),
            ("E13", "How effective are physiotherapy interventions in treating people with sciatica? A systematic review and meta-analysis", "PubMed", "2023", "物理治疗是首线保守治疗的重要组成，但疗效受方案与人群分层影响", "https://pubmed.ncbi.nlm.nih.gov/36580149/"),
            ("E14", "Acupuncture vs Sham Acupuncture for Chronic Sciatica From Herniated Disk: A Randomized Clinical Trial", "PubMed", "2024", "针刺对慢性坐骨神经痛疼痛和功能改善具有循证支持", "https://pubmed.ncbi.nlm.nih.gov/39401008/"),
            ("E15", "法规文件", "国家药品监督管理局", "2025", "药品全生命周期监管、说明书合规和院外传播边界持续收紧", "https://www.nmpa.gov.cn/xxgk/fgwj/"),
        ]
    elif is_gastritis_profile():
        evidence = [
            ("E01", "中国慢性胃炎诊治指南（2022年，上海）", "中华消化杂志", "2023", "慢性胃炎定义、分型、诊断与治疗路径", "https://rs.yiigle.com/cmaid/1473570"),
            ("E02", "第五次全国幽门螺杆菌感染处理共识报告", "中华消化杂志", "2022", "Hp检测方法、根除方案与复查窗口", "https://rs.yiigle.com/cmaid/1413767"),
            ("E03", "Kyoto global consensus report on gastritis", "Gut", "2015", "胃炎病因学、分类与胃癌风险链路", "https://gut.bmj.com/content/64/9/1353"),
            ("E04", "Maastricht VI/Florence consensus report", "Gut", "2022", "幽门螺杆菌感染管理国际推荐", "https://gut.bmj.com/content/71/9/1724"),
            ("E05", "ACG Guideline on Treatment of Helicobacter pylori Infection", "American Journal of Gastroenterology", "2024", "Hp根除治疗推荐与耐药背景下方案选择", "https://journals.lww.com/ajg/abstract/2024/10000/acg_clinical_guideline__treatment_of_helicobacter.13.aspx"),
            ("E06", "Global Cancer Observatory: Stomach cancer fact sheet", "IARC/WHO", "2024", "胃癌疾病负担与流行病学对比", "https://gco.iarc.who.int/media/globocan/factsheets/cancers/7-stomach-fact-sheet.pdf"),
            ("E07", "Diagnosis and treatment protocol for chronic gastritis (China)", "中国临床相关指南解读", "2022", "慢性胃炎临床分层与治疗终点", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9602100/"),
            ("E08", "国家统计数据发布平台（人口与社会）", "国家统计局", "2024", "人口结构与就医需求变化", "https://www.stats.gov.cn/"),
            ("E09", "国家基本医疗保险药品目录（2024年）", "国家医疗保障局", "2024", "支付规则影响抑酸/胃黏膜保护等用药可及性", "https://www.nhsa.gov.cn/art/2024/11/28/art_53_14887.html"),
            ("E10", "国家医保局2024年药品目录调整新闻发布会", "国家医疗保障局", "2024", "目录准入价值导向与支付口径", "https://www.nhsa.gov.cn/art/2024/11/28/art_52_14890.html"),
            ("E11", "中国卫生健康统计年鉴及国家统计数据库", "国家卫生健康委员会/国家统计局", "2024", "门急诊及消化系统疾病服务供给变化", "https://www.stats.gov.cn/"),
            ("E12", "米内网终端数据口径说明与原始数据文件", "米内网/项目数据", "2025", "医院/药店/线上三端同口径比较基础", f"{EXCEL_PATH.name}"),
            ("E13", "Management of epithelial precancerous conditions and lesions in the stomach", "MAPS II Guideline", "2019", "胃癌前病变随访分层与内镜病理管理", "https://pubmed.ncbi.nlm.nih.gov/30841008/"),
            ("E14", "Functional dyspepsia and overlap management evidence", "Rome Foundation/临床综述", "2023", "功能性消化不良与胃炎相关症状管理边界", "https://www.theromefoundation.org"),
            ("E15", "国家药监局法规与政策文件索引", "国家药品监督管理局", "2024", "药品全生命周期合规要求与监管边界", "https://www.nmpa.gov.cn/xxgk/fgwj/"),
        ]
    elif is_respiratory_profile():
        evidence = [
            ("E01", "Pneumonia (Fact sheet)", "World Health Organization", "2024", "儿童呼吸道疾病负担、死亡风险与干预重点", "https://www.who.int/en/news-room/fact-sheets/detail/pneumonia"),
            ("E02", "儿童社区获得性肺炎诊疗规范（2019年版）", "国家中医药管理局/国家卫健委", "2019", "分级诊疗、病情评估与规范用药路径", "https://www.natcm.gov.cn/bangongshi/zhengcewenjian/2019-02-13/9022.html"),
            ("E03", "儿童腺病毒肺炎诊疗规范（2019年版）", "国家卫生健康委员会", "2019", "重症识别、监测指标与动态评估", "https://www.nhc.gov.cn/ylyjs/zcwj/201906/6ef61f110a5548a489b6be3fc1674bc7.shtml"),
            ("E04", "国家基本医疗保险、工伤保险和生育保险药品目录（2024年）", "国家医疗保障局", "2024", "支付规则影响可及性与处方结构", "https://www.nhsa.gov.cn/art/2024/11/28/art_53_14887.html"),
            ("E05", "国家医保局2024年药品目录调整新闻发布会", "国家医疗保障局", "2024", "目录调整原则与临床价值导向", "https://www.nhsa.gov.cn/art/2024/11/28/art_52_14890.html"),
            ("E06", "PRAC recommends restrictions on use of codeine for cough and cold in children", "European Medicines Agency", "2015", "儿童镇咳成分年龄与安全边界", "https://www.ema.europa.eu/en/news/prac-recommends-restrictions-use-codeine-cough-cold-children"),
            ("E07", "FDA Drug Safety Communication: Codeine and Tramadol in Children", "U.S. Food and Drug Administration", "2017", "儿童使用可待因/曲马多的禁忌与风险警示", "https://www.fda.gov/drugs/drug-safety-and-availability"),
            ("E08", "国家统计数据发布平台（人口与社会）", "国家统计局", "2024", "儿童人群规模与区域分布变化", "https://www.stats.gov.cn/"),
            ("E09", "药品说明书修订公告汇总", "中国食品药品检定研究院", "2024", "说明书标签与警示信息更新", "https://www.cpi.ac.cn/tggg/ypsmsxdgg/"),
            ("E10", "中国统计年鉴及国家统计数据库", "国家统计局", "2024", "长期趋势对终端需求的结构性影响", "https://www.stats.gov.cn/"),
            ("E11", "儿童肺炎链球菌疾病诊疗和预防专家共识", "中华医学会儿科学分会等", "2020", "临床分层、预防与治疗要点", "https://rs.yiigle.com/cmaid/1297364"),
            ("E12", "米内网终端数据口径说明与原始数据文件", "米内网/项目数据", "2025", "医院/药店/线上三端同口径比较基础", f"{EXCEL_PATH.name}"),
            ("E13", "中国儿童慢性咳嗽诊断与治疗指南（2021）", "中华儿科杂志", "2021", "病程分类、红旗征与诊疗路径", "https://guide.medlive.cn/guideline/26694"),
            ("E14", "儿童咳嗽中西医结合诊治专家共识", "中西医结合相关学会", "2021", "中西医协同分层与处置建议", "https://guide.medlive.cn/guideline/23735"),
            ("E15", "国家药监局法规与政策文件索引", "国家药品监督管理局", "2024", "全生命周期合规要求与监管边界", "https://www.nmpa.gov.cn/xxgk/fgwj/"),
        ]
    else:
        dynamic = fetch_pubmed_evidence(DISEASE_NAME, max_items=8)
        evidence = []
        eid = 1
        for title, org, year, keypoint, url in dynamic:
            evidence.append((f"E{eid:02d}", title, org, year, keypoint, url))
            eid += 1
            if eid > 8:
                break
        fallback_dynamic = [
            (f"{DISEASE_NAME}诊疗指南与共识检索（PubMed）", "PubMed", "2024", "用于补充疾病特异性指南/共识证据", f"https://pubmed.ncbi.nlm.nih.gov/?term={quote_plus(disease_query_term(DISEASE_NAME) + ' guideline consensus')}"),
            (f"{DISEASE_NAME}系统综述检索（PubMed）", "PubMed", "2024", "用于补充治疗终点与人群分层证据", f"https://pubmed.ncbi.nlm.nih.gov/?term={quote_plus(disease_query_term(DISEASE_NAME) + ' systematic review treatment')}"),
        ]
        while len(evidence) < 8:
            title, org, year, keypoint, url = fallback_dynamic[(len(evidence) - len(dynamic)) % len(fallback_dynamic)]
            evidence.append((f"E{len(evidence)+1:02d}", title, org, year, keypoint, url))
        evidence.extend(
            [
                ("E09", "国家基本医疗保险、工伤保险和生育保险药品目录（2024年）", "国家医疗保障局", "2024", "支付规则影响可及性与处方结构", "https://www.nhsa.gov.cn/art/2024/11/28/art_53_14887.html"),
                ("E10", "国家医保局2024年药品目录调整新闻发布会", "国家医疗保障局", "2024", "目录调整原则与临床价值导向", "https://www.nhsa.gov.cn/art/2024/11/28/art_52_14890.html"),
                ("E11", "国家统计数据发布平台（人口与社会）", "国家统计局", "2024", "人口结构与就医需求变化", "https://www.stats.gov.cn/"),
                ("E12", "米内网终端数据口径说明与原始数据文件", "米内网/项目数据", "2025", "医院/药店/线上三端同口径比较基础", f"{EXCEL_PATH.name}"),
                ("E13", "中国卫生健康统计年鉴及国家统计数据库", "国家卫生健康委员会/国家统计局", "2024", "长期趋势对终端需求的结构性影响", "https://www.stats.gov.cn/"),
                ("E14", "药品说明书修订公告汇总", "中国食品药品检定研究院", "2024", "说明书标签与警示信息更新", "https://www.cpi.ac.cn/tggg/ypsmsxdgg/"),
                ("E15", "国家药监局法规与政策文件索引", "国家药品监督管理局", "2024", "全生命周期合规要求与监管边界", "https://www.nmpa.gov.cn/xxgk/fgwj/"),
            ]
        )

    evidence_lines = ["证据ID|标题|机构/作者|年份|要点|可追溯来源"]
    for e in evidence:
        evidence_lines.append("|".join(e))

    if is_cervical_profile():
        refs = [
            "[1] World Health Organization. Musculoskeletal conditions[EB/OL]. 2022. https://www.who.int/news-room/fact-sheets/detail/musculoskeletal-conditions",
            "[2] 中国康复医学会. 颈椎病康复诊疗专家共识[S/OL]. 2021. https://guide.medlive.cn/guideline/27562",
            "[3] 中华医学会骨科相关学组. 颈椎病诊疗与手术指征专家共识[S/OL]. 2020. https://guide.medlive.cn/guideline/24306",
            "[4] 国家医疗保障局. 国家基本医疗保险、工伤保险和生育保险药品目录（2024年）[S/OL]. 2024. https://www.nhsa.gov.cn/art/2024/11/28/art_53_14887.html",
            "[5] 国家卫生健康委员会. 国家医疗质量安全改进目标（骨科相关）[S/OL]. 2024. https://www.nhc.gov.cn/yzygj/s7659/",
            "[6] North American Spine Society. Diagnosis and Treatment of Cervical Radiculopathy[EB/OL]. 2020. https://www.spine.org/ResearchClinicalCare/QualityImprovement/ClinicalGuidelines",
            "[7] American Association of Neurological Surgeons. Cervical Myelopathy[EB/OL]. 2024. https://www.aans.org/Patients/Neurosurgical-Conditions-and-Treatments/Cervical-Myelopathy",
            "[8] 国家统计局. 国家统计数据发布平台（人口与社会）[DB/OL]. 2024. https://www.stats.gov.cn/",
            "[9] 中国食品药品检定研究院. 药品说明书修订公告汇总[EB/OL]. 2024. https://www.cpi.ac.cn/tggg/ypsmsxdgg/",
            "[10] 国家卫生健康委员会/国家统计局. 中国卫生健康统计年鉴及国家统计数据库[DB/OL]. 2024. https://www.stats.gov.cn/",
            "[11] 中华外科相关学组. 颈椎病外科治疗与围手术期管理共识[J/OL]. 2021. https://guide.medlive.cn/guideline/27112",
            f"[12] 米内网/项目数据. {EXCEL_PATH.name}（医院/药店/线上口径）[DB]. 2025.",
            "[13] 中国疼痛医学相关学会. 颈肩痛诊疗与慢性疼痛管理共识[S/OL]. 2022. https://guide.medlive.cn/guideline/28605",
            "[14] 中国中医药相关学会. 中医骨伤科颈椎病诊疗指南[S/OL]. 2021. https://guide.medlive.cn/guideline/26067",
            "[15] 国家药品监督管理局. 法规文件索引与政策发布[EB/OL]. 2024. https://www.nmpa.gov.cn/xxgk/fgwj/",
        ]
    elif is_gastritis_profile():
        refs = [
            "[1] 中华消化杂志. 中国慢性胃炎诊治指南（2022年，上海）[S/OL]. 2023. https://rs.yiigle.com/cmaid/1473570",
            "[2] 中华消化杂志. 第五次全国幽门螺杆菌感染处理共识报告[S/OL]. 2022. https://rs.yiigle.com/cmaid/1413767",
            "[3] Gut. Kyoto global consensus report on gastritis[EB/OL]. 2015. https://gut.bmj.com/content/64/9/1353",
            "[4] Gut. Maastricht VI/Florence consensus report[EB/OL]. 2022. https://gut.bmj.com/content/71/9/1724",
            "[5] American Journal of Gastroenterology. ACG clinical guideline: treatment of Helicobacter pylori infection[EB/OL]. 2024. https://journals.lww.com/ajg/abstract/2024/10000/acg_clinical_guideline__treatment_of_helicobacter.13.aspx",
            "[6] IARC/WHO. Global Cancer Observatory: Stomach cancer fact sheet[EB/OL]. 2024. https://gco.iarc.who.int/media/globocan/factsheets/cancers/7-stomach-fact-sheet.pdf",
            "[7] PMC. Diagnosis and treatment protocol for chronic gastritis in China[EB/OL]. 2022. https://pmc.ncbi.nlm.nih.gov/articles/PMC9602100/",
            "[8] 国家统计局. 国家统计数据发布平台（人口与社会）[DB/OL]. 2024. https://www.stats.gov.cn/",
            "[9] 国家医疗保障局. 国家基本医疗保险、工伤保险和生育保险药品目录（2024年）[S/OL]. 2024. https://www.nhsa.gov.cn/art/2024/11/28/art_53_14887.html",
            "[10] 国家医疗保障局. 2024年药品目录调整新闻发布会[EB/OL]. 2024. https://www.nhsa.gov.cn/art/2024/11/28/art_52_14890.html",
            "[11] 国家卫生健康委员会/国家统计局. 中国卫生健康统计年鉴及国家统计数据库[DB/OL]. 2024. https://www.stats.gov.cn/",
            f"[12] 米内网/项目数据. {EXCEL_PATH.name}（医院/药店/线上口径）[DB]. 2025.",
            "[13] PubMed. Management of epithelial precancerous conditions and lesions in the stomach (MAPS II)[EB/OL]. 2019. https://pubmed.ncbi.nlm.nih.gov/30841008/",
            "[14] Rome Foundation. Functional dyspepsia and overlap management evidence[EB/OL]. 2023. https://www.theromefoundation.org/",
            "[15] 国家药品监督管理局. 法规文件索引与政策发布[EB/OL]. 2024. https://www.nmpa.gov.cn/xxgk/fgwj/",
        ]
    elif is_respiratory_profile():
        refs = [
            "[1] World Health Organization. Pneumonia (Fact sheet)[EB/OL]. 2024. https://www.who.int/en/news-room/fact-sheets/detail/pneumonia",
            "[2] 国家中医药管理局. 儿童社区获得性肺炎诊疗规范（2019年版）[S/OL]. 2019. https://www.natcm.gov.cn/bangongshi/zhengcewenjian/2019-02-13/9022.html",
            "[3] 国家卫生健康委员会. 儿童腺病毒肺炎诊疗规范（2019年版）[S/OL]. 2019. https://www.nhc.gov.cn/ylyjs/zcwj/201906/6ef61f110a5548a489b6be3fc1674bc7.shtml",
            "[4] 国家医疗保障局. 国家基本医疗保险、工伤保险和生育保险药品目录（2024年）[S/OL]. 2024. https://www.nhsa.gov.cn/art/2024/11/28/art_53_14887.html",
            "[5] 国家医疗保障局. 2024年药品目录调整新闻发布会[EB/OL]. 2024. https://www.nhsa.gov.cn/art/2024/11/28/art_52_14890.html",
            "[6] European Medicines Agency. PRAC recommends restrictions on use of codeine for cough and cold in children[EB/OL]. 2015. https://www.ema.europa.eu/en/news/prac-recommends-restrictions-use-codeine-cough-cold-children",
            "[7] U.S. Food and Drug Administration. Drug Safety Communication: Use of Codeine and Tramadol in Children[EB/OL]. 2017. https://www.fda.gov/drugs/drug-safety-and-availability",
            "[8] 国家统计局. 国家统计数据发布平台（人口与社会）[DB/OL]. 2024. https://www.stats.gov.cn/",
            "[9] 中国食品药品检定研究院. 药品说明书修订公告汇总[EB/OL]. 2024. https://www.cpi.ac.cn/tggg/ypsmsxdgg/",
            "[10] 国家统计局. 中国统计年鉴及国家统计数据库[DB/OL]. 2024. https://www.stats.gov.cn/",
            "[11] 中华医学会儿科学分会等. 儿童肺炎链球菌疾病诊疗和预防专家共识[J/OL]. 2020. https://rs.yiigle.com/cmaid/1297364",
            f"[12] 米内网/项目数据. {EXCEL_PATH.name}（医院/药店/线上口径）[DB]. 2025.",
            "[13] 中华儿科杂志. 中国儿童慢性咳嗽诊断与治疗指南（2021）[S/OL]. 2021. https://guide.medlive.cn/guideline/26694",
            "[14] 儿童咳嗽中西医结合诊治专家共识[S/OL]. 2021. https://guide.medlive.cn/guideline/23735",
            "[15] 国家药品监督管理局. 法规文件索引与政策发布[EB/OL]. 2024. https://www.nmpa.gov.cn/xxgk/fgwj/",
        ]
    else:
        refs = []
        for i, e in enumerate(evidence, start=1):
            _, title, org, year, _, url = e
            dtype = "DB" if (EXCEL_PATH.name in url) else "EB/OL"
            refs.append(to_ref(i, title, org, year, url, dtype=dtype))

    return "\n".join(evidence_lines), "\n".join(refs)


@dataclass
class BlockSpec:
    block_id: str
    chapter: int
    subtitle: str
    target_chars: int
    topics: List[str]
    evidence_ids: str
    fig_ids: str


def build_block_specs() -> List[BlockSpec]:
    if is_cervical_profile():
        return [
            BlockSpec("1.1", 1, "1.1 疾病定义与分型（中西医角度）", 1250, ["定义边界", "分型标准", "影像分级", "中西医术语映射", "病程分期"], "E01|E02|E03", "fig_1_1"),
            BlockSpec("1.2", 1, "1.2 发病机制与病理生理", 1250, ["椎间盘退变", "骨赘形成", "神经根受压", "脊髓受压", "椎动脉供血"], "E01|E03|E13", "fig_1_2|fig_1_3"),
            BlockSpec("1.3", 1, "1.3 本章小结", 900, ["认知框架", "风险分层", "证据导向", "康复协同"], "E01|E02|E13", "fig_1_4"),
            BlockSpec("2.1", 2, "2.1 与神经、肌肉骨骼和血管系统的联系", 1500, ["神经功能", "肌肉骨骼代偿", "血管供血", "姿势负荷", "系统交互"], "E02|E03|E11", "fig_2_1"),
            BlockSpec("2.2", 2, "2.2 常见并发症与合并症", 1500, ["神经根型并发", "脊髓型风险", "椎动脉症状", "慢性疼痛失眠", "焦虑抑郁"], "E03|E11|E13", "fig_2_2|fig_2_3"),
            BlockSpec("2.3", 2, "2.3 本章小结", 1000, ["系统管理", "风险前移", "连续干预", "依从提升"], "E03|E11|E12", ""),
            BlockSpec("3.1", 3, "3.1 临床诊断标准与检查手段", 1600, ["病史采集", "体格检查", "红旗征识别", "影像路径", "功能量表"], "E02|E03|E11", "fig_3_1"),
            BlockSpec("3.2", 3, "3.2 西医治疗体系（药物、介入、手术、康复）", 1600, ["保守治疗", "药物镇痛", "介入治疗", "手术适应证", "安全监测"], "E03|E06|E13", "fig_3_2"),
            BlockSpec("3.3", 3, "3.3 中医辨证体系与常用方药", 1400, ["辨证分型", "治则治法", "中成药应用", "针灸推拿", "中西协同"], "E14|E13|E11", "fig_3_3"),
            BlockSpec("3.4", 3, "3.4 本章小结", 950, ["规范路径", "证据整合", "风险平衡", "长期管理"], "E13|E14|E15", ""),
            BlockSpec("4.1", 4, "4.1 治疗药物市场概况", 900, ["渠道规模", "季度趋势", "结构占比", "增长驱动"], "E09|E10|E14", "fig_4_1|fig_4_2"),
            BlockSpec("4.2", 4, "4.2 主要治疗药物分析", 900, ["头部通用名", "渠道差异", "品种生命周期", "结构优化"], "E08|E14|E15", "fig_4_3|fig_4_4"),
            BlockSpec("4.3", 4, "4.3 市场格局与竞争态势", 900, ["集中度", "竞争壁垒", "挑战者路径", "效率竞争"], "E08|E14|E15", "fig_4_5|fig_4_8"),
            BlockSpec("4.4", 4, "4.4 本章小结", 800, ["数据闭环", "经营动作", "跨部门协同", "季度复盘"], "E09|E10|E14", "fig_4_6|fig_4_7"),
            BlockSpec("5.1", 5, "5.1 患者群体结构与画像（性别/年龄/地域/职业负荷/用药偏好）", 1600, ["年龄结构", "就诊场景", "地区差异", "职业负荷", "负担分层"], "E04|E08|E13", "fig_5_1"),
            BlockSpec("5.2", 5, "5.2 医生处方偏好与诊疗习惯（路径差异/未满足需求）", 1600, ["处方偏好", "科室差异", "证据偏好", "未满足需求", "康复路径"], "E03|E05|E08", "fig_5_2|fig_5_4"),
            BlockSpec("5.3", 5, "5.3 患者依从性与长期管理（复评管理/全周期流程）", 1400, ["依从瓶颈", "居家训练", "复评管理", "随访机制", "长期控制"], "E04|E05|E12", "fig_5_3"),
            BlockSpec("5.4", 5, "5.4 本章小结", 900, ["患者中心", "医生协同", "连续管理", "服务化能力"], "E04|E05|E15", ""),
            BlockSpec("6.1", 6, "6.1 疾病政策环境（6.1.1全球；6.1.2中国）", 1760, ["政策主线", "骨科分级诊疗", "准入与支付", "质量体系", "监管协同"], "E09|E10|E11", "fig_6_1"),
            BlockSpec("6.2", 6, "6.2 疾病监管趋势（审评审批/质量控制/医保支付/行业监管影响）", 1760, ["审评提速", "质量标准", "医保支付", "合规传播", "全链条监管"], "E09|E10|E11", "fig_6_2"),
            BlockSpec("6.3", 6, "6.3 本章小结", 1080, ["政策约束", "策略边界", "合规经营", "长期确定性"], "E09|E10|E15", ""),
            BlockSpec("7.1", 7, "7.1 未来市场预测（负担/市场规模/治疗方案演进）", 1760, ["需求预测", "规模外推", "渠道重构", "证据升级", "方案演进"], "E08|E13|E14", "fig_7_1"),
            BlockSpec("7.2", 7, "7.2 战略建议（市场部/战略部：定位、证据、准入、渠道、生命周期）", 1760, ["定位策略", "证据策略", "准入策略", "渠道策略", "生命周期管理"], "E08|E10|E15", "fig_7_2"),
            BlockSpec("7.3", 7, "7.3 本章小结", 1080, ["前瞻判断", "执行闭环", "组织协同", "风险应对"], "E10|E14|E15", ""),
        ]
    if is_sciatica_profile():
        return [
            BlockSpec("1.1", 1, "1.1 疾病定义与病因分层（椎间盘突出/椎管狭窄/梨状肌相关）", 1250, ["定义边界", "病因分层", "神经根定位", "病程阶段", "高危场景"], "E01|E02|E03", "fig_1_1"),
            BlockSpec("1.2", 1, "1.2 发病机制与病理生理（神经根受压-炎症放大-疼痛敏化）", 1250, ["神经根受压", "炎症放大", "疼痛敏化", "功能受限", "风险触发因素"], "E01|E03|E13", "fig_1_2|fig_1_3"),
            BlockSpec("1.3", 1, "1.3 本章小结", 900, ["诊断框架", "病因分层", "风险识别", "康复主线"], "E01|E02|E13", "fig_1_4"),
            BlockSpec("2.1", 2, "2.1 与神经、肌肉骨骼及睡眠心理系统的关联", 1500, ["神经功能", "肌肉代偿", "步态改变", "睡眠受损", "系统交互"], "E02|E03|E11", "fig_2_1"),
            BlockSpec("2.2", 2, "2.2 常见并发问题与风险管理", 1500, ["慢性神经病理性疼痛", "运动功能下降", "睡眠障碍", "情绪负担", "红旗征"], "E03|E07|E13", "fig_2_2|fig_2_3"),
            BlockSpec("2.3", 2, "2.3 本章小结", 1000, ["风险前移", "分层复评", "功能管理", "长期控制"], "E03|E11|E12", ""),
            BlockSpec("3.1", 3, "3.1 临床诊断标准与检查路径", 1600, ["病史采集", "神经体征", "红旗征识别", "影像时机", "功能量表"], "E02|E03|E07", "fig_3_1"),
            BlockSpec("3.2", 3, "3.2 西医治疗体系（药物、介入、手术、康复）", 1600, ["起始治疗", "镇痛策略", "介入治疗", "手术适应证", "安全监测"], "E03|E06|E13", "fig_3_2"),
            BlockSpec("3.3", 3, "3.3 中医辨证路径与中西协同管理", 1400, ["辨证分型", "针灸应用", "中成药", "运动康复", "疗效终点"], "E11|E13|E14", "fig_3_3"),
            BlockSpec("3.4", 3, "3.4 本章小结", 950, ["分层诊疗", "复评管理", "证据闭环", "长期获益"], "E13|E14|E15", ""),
            BlockSpec("4.1", 4, "4.1 治疗药物市场概况", 900, ["渠道规模", "季度趋势", "结构占比", "增长驱动"], "E09|E10|E12", "fig_4_1|fig_4_2"),
            BlockSpec("4.2", 4, "4.2 主要治疗药物分析", 900, ["头部通用名", "渠道差异", "品类结构", "增速分化"], "E08|E10|E12", "fig_4_3|fig_4_4"),
            BlockSpec("4.3", 4, "4.3 市场格局与竞争态势", 900, ["集中度", "核心品种", "竞争壁垒", "替代路径"], "E08|E12|E15", "fig_4_5|fig_4_8"),
            BlockSpec("4.4", 4, "4.4 本章小结", 800, ["口径一致性", "数据复算", "主要结论", "风险提示"], "E09|E10|E12", "fig_4_6|fig_4_7"),
            BlockSpec("5.1", 5, "5.1 患者群体结构与画像（年龄/职业负荷/疼痛分层/用药偏好）", 1600, ["年龄结构", "职业负荷", "就诊场景", "疼痛分层", "需求差异"], "E04|E08|E13", "fig_5_1"),
            BlockSpec("5.2", 5, "5.2 医生处方偏好与诊疗习惯（骨科/疼痛科/康复科差异）", 1600, ["处方偏好", "科室差异", "证据偏好", "未满足需求", "转诊路径"], "E03|E05|E08", "fig_5_2|fig_5_4"),
            BlockSpec("5.3", 5, "5.3 患者依从性与长期管理（运动处方/复评窗口/全周期管理）", 1400, ["依从瓶颈", "运动处方", "复评窗口", "随访机制", "长期控制"], "E04|E11|E12", "fig_5_3"),
            BlockSpec("5.4", 5, "5.4 本章小结", 900, ["患者中心", "路径协同", "连续管理", "服务化能力"], "E04|E05|E15", ""),
            BlockSpec("6.1", 6, "6.1 疾病政策环境（全球与中国）", 1700, ["指南更新", "医保支付", "质量管理", "分级诊疗", "监管协同"], "E04|E05|E15", "fig_6_1"),
            BlockSpec("6.2", 6, "6.2 监管趋势对品类与渠道的影响", 1700, ["说明书边界", "支付规则", "合规传播", "院内准入", "院外规范"], "E04|E09|E15", "fig_6_2"),
            BlockSpec("6.3", 6, "6.3 本章小结", 1000, ["政策边界", "支付约束", "合规经营", "长期确定性"], "E04|E05|E15", ""),
            BlockSpec("7.1", 7, "7.1 市场预测（方法-假设-情景）", 1700, ["需求预测", "规模外推", "渠道重构", "治疗演进", "敏感性"], "E08|E10|E12", "fig_7_1"),
            BlockSpec("7.2", 7, "7.2 战略建议（产品/证据/准入/渠道/生命周期）", 1700, ["产品定位", "证据策略", "准入策略", "渠道组合", "生命周期管理"], "E04|E10|E15", "fig_7_2"),
            BlockSpec("7.3", 7, "7.3 本章小结", 1000, ["增长判断", "资源配置", "执行优先级", "风险对冲"], "E10|E12|E15", ""),
        ]
    if is_gastritis_profile():
        return [
            BlockSpec("1.1", 1, "1.1 疾病定义与病因分层（Hp相关/非Hp相关）", 1250, ["定义边界", "病因结构", "病理分型", "高危人群", "临床分层"], "E01|E02|E03", "fig_1_1"),
            BlockSpec("1.2", 1, "1.2 发病机制与病理演进（炎症-萎缩-肠化）", 1250, ["炎症链路", "萎缩进展", "肠化风险", "异型增生", "进展触发因素"], "E01|E03|E13", "fig_1_2|fig_1_3"),
            BlockSpec("1.3", 1, "1.3 本章小结", 900, ["病因分层", "风险链路", "证据等级", "诊疗边界"], "E01|E02|E13", "fig_1_4"),
            BlockSpec("2.1", 2, "2.1 与免疫、代谢及消化动力系统的关联", 1500, ["免疫炎症", "代谢影响", "胃酸分泌", "胃动力", "系统交互"], "E02|E03|E11", "fig_2_1"),
            BlockSpec("2.2", 2, "2.2 并发症与癌前病变风险管理", 1500, ["糜烂风险", "溃疡风险", "贫血风险", "癌前病变", "随访窗口"], "E03|E11|E13", "fig_2_2|fig_2_3"),
            BlockSpec("2.3", 2, "2.3 本章小结", 1000, ["分层随访", "风险前移", "并发管理", "可及性约束"], "E03|E11|E12", ""),
            BlockSpec("3.1", 3, "3.1 临床诊断标准与检查路径（内镜+病理+Hp检测）", 1600, ["症状筛查", "Hp检测", "内镜检查", "病理分级", "复查路径"], "E01|E02|E13", "fig_3_1"),
            BlockSpec("3.2", 3, "3.2 西医治疗体系（抑酸/黏膜保护/Hp根除/促动力）", 1600, ["抑酸方案", "黏膜保护", "根除治疗", "促动力/消化酶", "不良反应监测"], "E02|E04|E05", "fig_3_2"),
            BlockSpec("3.3", 3, "3.3 中医辨证路径与中西协同治疗", 1400, ["辨证分型", "治则治法", "中成药应用", "中西协同", "疗效终点"], "E01|E14|E15", "fig_3_3"),
            BlockSpec("3.4", 3, "3.4 本章小结", 950, ["分层诊疗", "随访管理", "安全边界", "证据闭环"], "E05|E13|E15", ""),
            BlockSpec("4.1", 4, "4.1 市场口径定义与三端规模趋势", 900, ["口径定义", "品类范围", "季度趋势", "结构占比"], "E09|E10|E12", "fig_4_1|fig_4_2"),
            BlockSpec("4.2", 4, "4.2 核心治疗品类与重点通用名", 900, ["品类结构", "头部通用名", "渠道差异", "价格带"], "E08|E10|E12", "fig_4_3|fig_4_4"),
            BlockSpec("4.3", 4, "4.3 竞争格局与玩家地图（CR5/TOP10）", 900, ["CR5", "TOP10覆盖", "玩家分层", "竞争壁垒"], "E08|E12|E15", "fig_4_5|fig_4_8"),
            BlockSpec("4.4", 4, "4.4 本章小结", 800, ["口径一致性", "复算路径", "主要结论", "风险提示"], "E09|E10|E12", "fig_4_6|fig_4_7"),
            BlockSpec("5.1", 5, "5.1 患者画像与诊疗分层（年龄/病因/病理风险）", 1600, ["年龄结构", "病因分层", "病理风险", "就医路径", "用药偏好"], "E01|E08|E13", "fig_5_1"),
            BlockSpec("5.2", 5, "5.2 处方路径与未满足临床需求", 1600, ["科室处方差异", "治疗目标差异", "依从挑战", "未满足需求", "证据缺口"], "E02|E05|E11", "fig_5_2|fig_5_4"),
            BlockSpec("5.3", 5, "5.3 长期管理与复查依从（根除后与癌前病变）", 1400, ["根除后管理", "复查窗口", "依从性因素", "院内外协同", "长期风险"], "E04|E05|E13", "fig_5_3"),
            BlockSpec("5.4", 5, "5.4 本章小结", 900, ["人群分层", "处方策略", "复查路径", "证据升级"], "E04|E05|E15", ""),
            BlockSpec("6.1", 6, "6.1 政策环境（全球与中国）", 1700, ["指南更新", "医保支付", "审评监管", "质量标准", "合规传播"], "E09|E10|E11", "fig_6_1"),
            BlockSpec("6.2", 6, "6.2 监管趋势对品类与渠道的影响", 1700, ["审评审批", "药品说明书", "支付规则", "院内准入", "互联网规范"], "E09|E10|E15", "fig_6_2"),
            BlockSpec("6.3", 6, "6.3 本章小结", 1000, ["政策边界", "准入路径", "渠道约束", "合规风险"], "E09|E10|E15", ""),
            BlockSpec("7.1", 7, "7.1 市场预测（方法-假设-情景）", 1700, ["模型方法", "基准假设", "三情景预测", "敏感性", "风险变量"], "E08|E10|E12", "fig_7_1"),
            BlockSpec("7.2", 7, "7.2 战略建议（产品/证据/准入/渠道）", 1700, ["产品定位", "证据策略", "准入策略", "渠道组合", "生命周期管理"], "E09|E10|E15", "fig_7_2"),
            BlockSpec("7.3", 7, "7.3 本章小结", 1000, ["增长判断", "资源配置", "执行优先级", "风险对冲"], "E10|E12|E15", ""),
        ]
    if is_pharyngitis_profile():
        return [
            BlockSpec("1.1", 1, "1.1 疾病定义与病因分层（感染/过敏/刺激暴露）", 1250, ["定义边界", "病因结构", "病程分层", "高危人群", "临床分层"], "E01|E02|E03", "fig_1_1"),
            BlockSpec("1.2", 1, "1.2 发病机制与病理生理（局部炎症-神经敏化）", 1250, ["炎症反应", "神经敏化", "黏膜屏障", "反流刺激", "微生态失衡"], "E01|E03|E13", "fig_1_2|fig_1_3"),
            BlockSpec("1.3", 1, "1.3 本章小结", 900, ["病因分层", "风险链路", "证据边界", "诊疗锚点"], "E01|E02|E13", "fig_1_4"),
            BlockSpec("2.1", 2, "2.1 与免疫、神经及上气道微生态的关联", 1500, ["免疫炎症", "神经反射", "微生态波动", "环境暴露", "系统交互"], "E02|E03|E11", "fig_2_1"),
            BlockSpec("2.2", 2, "2.2 常见并发问题与风险管理", 1500, ["咽喉反流重叠", "睡眠受损", "声音疲劳", "焦虑负担", "复诊负担"], "E03|E11|E13", "fig_2_2|fig_2_3"),
            BlockSpec("2.3", 2, "2.3 本章小结", 1000, ["风险前移", "分层复评", "依从管理", "长期控制"], "E03|E11|E12", ""),
            BlockSpec("3.1", 3, "3.1 临床诊断标准与检查路径（耳鼻喉专科）", 1600, ["病史采集", "喉镜评估", "红旗征识别", "鉴别诊断", "复评路径"], "E02|E03|E11", "fig_3_1"),
            BlockSpec("3.2", 3, "3.2 西医治疗体系（病因控制+症状缓解）", 1600, ["病因控制", "黏膜保护", "抗炎治疗", "合并反流管理", "安全监测"], "E03|E06|E13", "fig_3_2"),
            BlockSpec("3.3", 3, "3.3 中医辨证路径与中西协同管理", 1400, ["辨证分型", "治则治法", "中成药应用", "中西协同", "疗效终点"], "E14|E13|E11", "fig_3_3"),
            BlockSpec("3.4", 3, "3.4 本章小结", 950, ["分层诊疗", "复评管理", "安全边界", "证据闭环"], "E13|E14|E15", ""),
            BlockSpec("4.1", 4, "4.1 治疗药物市场概况", 900, ["渠道规模", "季度趋势", "结构占比", "增长驱动"], "E09|E10|E12", "fig_4_1|fig_4_2"),
            BlockSpec("4.2", 4, "4.2 主要治疗药物分析", 900, ["头部通用名", "渠道差异", "品类结构", "价格带"], "E08|E12|E15", "fig_4_3|fig_4_4"),
            BlockSpec("4.3", 4, "4.3 市场格局与竞争态势", 900, ["集中度", "玩家分层", "竞争壁垒", "切入路径"], "E08|E12|E15", "fig_4_5|fig_4_8"),
            BlockSpec("4.4", 4, "4.4 本章小结", 800, ["口径一致性", "复算路径", "主要结论", "风险提示"], "E09|E10|E12", "fig_4_6|fig_4_7"),
            BlockSpec("5.1", 5, "5.1 患者画像与就诊分层（年龄/职业暴露/症状谱）", 1600, ["年龄结构", "职业暴露", "症状谱", "就诊路径", "用药偏好"], "E04|E08|E13", "fig_5_1"),
            BlockSpec("5.2", 5, "5.2 处方路径与未满足临床需求", 1600, ["处方偏好", "科室差异", "复发管理", "未满足需求", "证据缺口"], "E03|E05|E11", "fig_5_2|fig_5_4"),
            BlockSpec("5.3", 5, "5.3 长期管理与依从性挑战", 1400, ["依从障碍", "环境干预", "复评节点", "院内外协同", "长期风险"], "E04|E05|E13", "fig_5_3"),
            BlockSpec("5.4", 5, "5.4 本章小结", 900, ["人群分层", "处方策略", "长期管理", "证据升级"], "E04|E05|E15", ""),
            BlockSpec("6.1", 6, "6.1 政策环境（全球与中国）", 1700, ["指南更新", "医保支付", "审评监管", "质量标准", "合规传播"], "E09|E10|E11", "fig_6_1"),
            BlockSpec("6.2", 6, "6.2 监管趋势对品类与渠道的影响", 1700, ["审评审批", "说明书边界", "支付规则", "院内准入", "互联网规范"], "E09|E10|E15", "fig_6_2"),
            BlockSpec("6.3", 6, "6.3 本章小结", 1000, ["政策边界", "准入路径", "渠道约束", "合规风险"], "E09|E10|E15", ""),
            BlockSpec("7.1", 7, "7.1 市场预测（方法-假设-情景）", 1700, ["模型方法", "基准假设", "三情景预测", "敏感性", "风险变量"], "E08|E10|E12", "fig_7_1"),
            BlockSpec("7.2", 7, "7.2 战略建议（产品/证据/准入/渠道）", 1700, ["产品定位", "证据策略", "准入策略", "渠道组合", "生命周期管理"], "E09|E10|E15", "fig_7_2"),
            BlockSpec("7.3", 7, "7.3 本章小结", 1000, ["增长判断", "资源配置", "执行优先级", "风险对冲"], "E10|E12|E15", ""),
        ]
    if not is_respiratory_profile():
        return [
            BlockSpec("1.1", 1, "1.1 医学主题定义与应用边界", 1250, ["定义边界", "分型标准", "诊断标准", "病程分层", "适应证场景"], "E01|E02|E03", "fig_1_1"),
            BlockSpec("1.2", 1, "1.2 核心机制与病理生理基础", 1250, ["发病机制", "病理改变", "关键通路", "进展链路", "风险触发因素"], "E01|E03|E13", "fig_1_2|fig_1_3"),
            BlockSpec("1.3", 1, "1.3 本章小结", 900, ["认知框架", "风险分层", "证据导向", "管理边界"], "E01|E02|E13", "fig_1_4"),
            BlockSpec("2.1", 2, "2.1 与相关系统的联系（多系统视角）", 1500, ["系统关联", "代谢影响", "免疫调节", "神经调控", "系统交互"], "E02|E03|E11", "fig_2_1"),
            BlockSpec("2.2", 2, "2.2 常见风险问题与管理要点", 1500, ["并发风险", "共病谱", "复发风险", "治疗耐受性", "复诊负担"], "E03|E11|E13", "fig_2_2|fig_2_3"),
            BlockSpec("2.3", 2, "2.3 本章小结", 1000, ["系统管理", "风险前移", "连续干预", "依从提升"], "E03|E11|E12", ""),
            BlockSpec("3.1", 3, "3.1 临床/应用评估标准与检查路径", 1600, ["分诊评估", "病因筛查", "红旗征识别", "检查路径", "评估量表"], "E02|E03|E11", "fig_3_1"),
            BlockSpec("3.2", 3, "3.2 干预体系与证据结构（药物/器械/疫苗/康复等）", 1600, ["起始治疗", "联合策略", "药物管理", "安全监测", "疗程调整"], "E03|E06|E13", "fig_3_2"),
            BlockSpec("3.3", 3, "3.3 中医辨证与协同管理（如适用）", 1400, ["辨证分型", "治则治法", "中成药应用", "中西协同", "证据等级"], "E14|E13|E11", "fig_3_3"),
            BlockSpec("3.4", 3, "3.4 本章小结", 950, ["规范路径", "证据整合", "风险平衡", "全周期管理"], "E13|E14|E15", ""),
            BlockSpec("4.1", 4, "4.1 治疗药物市场概况", 900, ["渠道规模", "季度趋势", "结构占比", "增长驱动"], "E09|E10|E14", "fig_4_1|fig_4_2"),
            BlockSpec("4.2", 4, "4.2 主要治疗药物分析", 900, ["头部通用名", "渠道差异", "品种生命周期", "结构优化"], "E08|E14|E15", "fig_4_3|fig_4_4"),
            BlockSpec("4.3", 4, "4.3 市场格局与竞争态势", 900, ["集中度", "竞争壁垒", "挑战者路径", "效率竞争"], "E08|E14|E15", "fig_4_5|fig_4_8"),
            BlockSpec("4.4", 4, "4.4 本章小结", 800, ["数据闭环", "经营动作", "跨部门协同", "季度复盘"], "E09|E10|E14", "fig_4_6|fig_4_7"),
            BlockSpec("5.1", 5, "5.1 目标人群结构与终端画像（年龄/地域/负担/偏好）", 1600, ["年龄结构", "就诊场景", "地区差异", "疾病负担", "需求分层"], "E04|E08|E13", "fig_5_1"),
            BlockSpec("5.2", 5, "5.2 终端使用偏好与未满足需求", 1600, ["处方偏好", "科室差异", "证据偏好", "未满足需求", "教育路径"], "E03|E05|E08", "fig_5_2|fig_5_4"),
            BlockSpec("5.3", 5, "5.3 依从性与长期管理（使用现状/影响因素/管理策略）", 1400, ["依从瓶颈", "行为执行", "复购管理", "随访机制", "长期控制"], "E04|E05|E12", "fig_5_3"),
            BlockSpec("5.4", 5, "5.4 本章小结", 900, ["患者中心", "医生协同", "连续管理", "服务化能力"], "E04|E05|E15", ""),
            BlockSpec("6.1", 6, "6.1 医学主题政策环境（全球与中国）", 1700, ["政策主线", "临床规范", "准入与支付", "质量体系", "监管协同"], "E09|E10|E11", "fig_6_1"),
            BlockSpec("6.2", 6, "6.2 监管趋势（审评审批/质量控制/医保支付/行业监管影响）", 1700, ["审评提速", "质量标准", "医保支付", "合规传播", "全链条监管"], "E09|E10|E11", "fig_6_2"),
            BlockSpec("6.3", 6, "6.3 本章小结", 1000, ["政策约束", "策略边界", "合规经营", "长期确定性"], "E09|E10|E15", ""),
            BlockSpec("7.1", 7, "7.1 未来市场预测（负担/市场规模/方案演进；量化+方法+来源）", 1700, ["需求预测", "规模外推", "渠道重构", "证据升级", "方案演进"], "E08|E13|E14", "fig_7_1"),
            BlockSpec("7.2", 7, "7.2 战略建议（面向市场部/战略部：定位、证据、准入、渠道、竞争、生命周期）", 1700, ["定位策略", "证据策略", "准入策略", "渠道策略", "生命周期管理"], "E08|E10|E15", "fig_7_2"),
            BlockSpec("7.3", 7, "7.3 本章小结", 1000, ["前瞻判断", "执行闭环", "组织协同", "风险应对"], "E10|E14|E15", ""),
        ]
    return [
        BlockSpec("1.1", 1, "1.1 疾病定义与分类（中西医角度）", 1250, ["定义边界", "分型标准", "年龄分层", "中西医术语映射", "适应症场景"], "E01|E02|E03", "fig_1_1"),
        BlockSpec("1.2", 1, "1.2 发病机制与病理生理", 1250, ["炎症反应", "分泌物黏稠度", "气道反应性", "病程分期", "风险触发因素"], "E01|E03|E13", "fig_1_2|fig_1_3"),
        BlockSpec("1.3", 1, "1.3 本章小结", 900, ["认知框架", "分层管理", "证据导向", "跨渠道协同"], "E01|E02|E13", "fig_1_4"),
        BlockSpec("2.1", 2, "2.1 与相关系统的联系（呼吸、消化、免疫、神经、内分泌等）", 1500, ["呼吸-免疫联动", "消化吸收影响", "神经调节", "内分泌节律", "系统交互"], "E02|E03|E11", "fig_2_1"),
        BlockSpec("2.2", 2, "2.2 常见并发症与合并症", 1500, ["并发感染", "睡眠受损", "反复咳嗽", "家长焦虑", "复诊负担"], "E03|E11|E13", "fig_2_2|fig_2_3"),
        BlockSpec("2.3", 2, "2.3 本章小结", 1000, ["系统管理", "风险前移", "连续干预", "依从提升"], "E03|E11|E12", ""),
        BlockSpec("3.1", 3, "3.1 临床诊断标准与检查手段", 1600, ["分诊评估", "病因筛查", "红旗征识别", "检查路径", "评估量表"], "E02|E03|E11", "fig_3_1"),
        BlockSpec("3.2", 3, "3.2 西医治疗体系（药物、手术、理疗）", 1600, ["起始治疗", "联合策略", "剂型选择", "安全监测", "疗程调整"], "E03|E06|E13", "fig_3_2"),
        BlockSpec("3.3", 3, "3.3 中医辨证体系与常用方药", 1400, ["辨证分型", "治则治法", "中成药应用", "中西协同", "证据等级"], "E14|E13|E11", "fig_3_3"),
        BlockSpec("3.4", 3, "3.4 本章小结", 950, ["规范路径", "证据整合", "风险平衡", "全周期管理"], "E13|E14|E15", ""),
        BlockSpec("4.1", 4, "4.1 治疗药物市场概况", 900, ["渠道规模", "季度趋势", "结构占比", "增长驱动"], "E09|E10|E14", "fig_4_1|fig_4_2"),
        BlockSpec("4.2", 4, "4.2 主要治疗药物分析", 900, ["头部通用名", "渠道差异", "品种生命周期", "结构优化"], "E08|E14|E15", "fig_4_3|fig_4_4"),
        BlockSpec("4.3", 4, "4.3 市场格局与竞争态势", 900, ["集中度", "竞争壁垒", "挑战者路径", "效率竞争"], "E08|E14|E15", "fig_4_5|fig_4_8"),
        BlockSpec("4.4", 4, "4.4 本章小结", 800, ["数据闭环", "经营动作", "跨部门协同", "季度复盘"], "E09|E10|E14", "fig_4_6|fig_4_7"),
        BlockSpec("5.1", 5, "5.1 患者群体结构与画像（性别/年龄/地域/负担/分层/用药偏好）", 1600, ["年龄结构", "就诊场景", "地区差异", "家庭决策", "负担分层"], "E04|E08|E13", "fig_5_1"),
        BlockSpec("5.2", 5, "5.2 医生用药偏好与诊疗习惯（处方行为/路径差异/未满足需求）", 1600, ["处方偏好", "科室差异", "证据偏好", "未满足需求", "教育路径"], "E03|E05|E08", "fig_5_2|fig_5_4"),
        BlockSpec("5.3", 5, "5.3 患者依从性与长期管理（依从现状/影响因素/管理策略/全周期管理流程图）", 1400, ["依从瓶颈", "家长执行", "复购管理", "随访机制", "长期控制"], "E04|E05|E12", "fig_5_3"),
        BlockSpec("5.4", 5, "5.4 本章小结", 900, ["患者中心", "医生协同", "连续管理", "服务化能力"], "E04|E05|E15", ""),
        BlockSpec("6.1", 6, "6.1 疾病政策环境（6.1.1全球；6.1.2中国）", 1700, ["政策主线", "儿童用药规范", "准入与支付", "质量体系", "监管协同"], "E09|E10|E11", "fig_6_1"),
        BlockSpec("6.2", 6, "6.2 疾病监管趋势（审评审批/质量控制/医保支付/行业监管对路径与用药结构的影响）", 1700, ["审评提速", "质量标准", "医保支付", "合规传播", "全链条监管"], "E09|E10|E11", "fig_6_2"),
        BlockSpec("6.3", 6, "6.3 本章小结", 1000, ["政策约束", "策略边界", "合规经营", "长期确定性"], "E09|E10|E15", ""),
        BlockSpec("7.1", 7, "7.1 未来市场预测（负担/市场规模/方案演进；量化+方法+来源）", 1700, ["需求预测", "规模外推", "渠道重构", "证据升级", "方案演进"], "E08|E13|E14", "fig_7_1"),
        BlockSpec("7.2", 7, "7.2 战略建议（面向市场部/战略部：产品定位、证据策略、准入策略、渠道策略、竞争应对、生命周期管理等，可执行）", 1700, ["定位策略", "证据策略", "准入策略", "渠道策略", "生命周期管理"], "E08|E10|E15", "fig_7_2"),
        BlockSpec("7.3", 7, "7.3 本章小结", 1000, ["前瞻判断", "执行闭环", "组织协同", "风险应对"], "E10|E14|E15", ""),
    ]


def chapter_title(chapter: int) -> str:
    names = {
        1: f"第一章 {DISEASE_NAME}概述与定义边界",
        2: f"第二章 {DISEASE_NAME}机制关联与风险管理",
        3: f"第三章 {DISEASE_NAME}临床应用与干预路径",
        4: f"第四章 {DISEASE_NAME}市场规模与竞争格局",
        5: f"第五章 {DISEASE_NAME}人群画像与终端需求",
        6: f"第六章 {DISEASE_NAME}政策环境与监管趋势",
        7: f"第七章 {DISEASE_NAME}未来展望与战略建议",
    }
    return names[chapter]


def split_paragraphs(text: str) -> List[str]:
    return [p.strip() for p in text.split("\n\n") if p.strip()]


def clean_title_prefix(subtitle: str, text: str) -> str:
    paras = split_paragraphs(text)
    if paras and paras[0].startswith(subtitle):
        paras[0] = paras[0][len(subtitle) :].lstrip("：:，,。 ")
    out = []
    seen = set()
    for p in paras:
        if p not in seen:
            out.append(p)
            seen.add(p)
    return "\n\n".join(out)


def write_evidence_and_refs() -> Tuple[str, str]:
    evidence_text, refs_text = build_evidence_and_refs()
    write_text(OUT_ROOT / "00_evidence.txt", evidence_text + "\n")
    write_text(OUT_ROOT / "refs.txt", refs_text + "\n")
    return evidence_text, refs_text


def _safe_float(value: object, default: float = 0.0) -> float:
    try:
        out = float(value)
    except Exception:
        return default
    if math.isnan(out) or math.isinf(out):
        return default
    return out


def _format_money(value: object) -> str:
    return f"{_safe_float(value):.1f}万元"


def _format_pct(value: object) -> str:
    return f"{_safe_float(value):.1f}%"


def ensure_fig23_codex_spec_ready() -> None:
    path = OUT_ROOT / FIG23_CODEX_SPEC_NAME
    if not path.exists():
        raise FileNotFoundError(
            f"Missing Codex fig23 spec: {path}. Please author it via {OUT_ROOT / FIG23_CODEX_PROMPT_NAME}."
        )
    raw = read_fig23_codex_spec()
    if not raw:
        raise RuntimeError(
            f"Invalid Codex fig23 spec: {path}. Please rewrite it using {OUT_ROOT / FIG23_CODEX_PROMPT_NAME}."
        )


def build_semantic_figure_specs_template() -> Dict[str, Dict[str, object]]:
    if is_cervical_profile():
        fig13_title = f"{DISEASE_NAME}病理生理演进路径"
        fig13_nodes = ["椎间盘退变", "椎体边缘骨赘", "椎管/椎间孔狭窄", "神经结构受压", "疼痛与功能障碍"]
        fig31_title = f"{DISEASE_NAME}临床诊疗流程"
        fig31_nodes = ["首诊分层", "神经体征检查", "影像评估", "保守治疗", "复评分流", "介入/手术"]
        fig53_title = f"{DISEASE_NAME}全周期管理流程"
        fig53_nodes = ["症状识别", "初诊评估", "治疗启动", "复评调整", "复发预防", "长期随访"]
    elif is_gastritis_profile():
        fig13_title = f"{DISEASE_NAME}病理演进链路"
        fig13_nodes = ["病因暴露", "慢性炎症", "腺体萎缩", "肠上皮化生", "风险升级"]
        fig31_title = f"{DISEASE_NAME}临床诊疗流程"
        fig31_nodes = ["症状与红旗征筛查", "Hp检测", "内镜+病理分级", "病因分层", "治疗执行", "复查随访"]
        fig53_title = f"{DISEASE_NAME}长期管理流程"
        fig53_nodes = ["症状出现", "初诊评估", "病因分层", "治疗执行", "复查评估", "长期随访"]
    elif is_respiratory_profile():
        fig13_title = f"{DISEASE_NAME}病理生理演进路径"
        fig13_nodes = ["感染/过敏触发", "炎症与分泌增加", "痰液黏稠", "排痰受阻", "持续咳嗽"]
        fig31_title = f"{DISEASE_NAME}临床诊疗流程"
        fig31_nodes = ["首诊分诊", "病因评估", "风险分层", "治疗启动", "复评调整", "出院/随访"]
        fig53_title = f"{DISEASE_NAME}全周期管理流程"
        fig53_nodes = ["首发症状", "初诊评估", "治疗启动", "功能恢复", "复发预防", "长期随访"]
    else:
        fig13_title = f"{DISEASE_NAME}病理生理演进路径"
        fig13_nodes = ["风险暴露", "病理改变", "功能受损", "症状负担上升", "分层管理"]
        fig31_title = f"{DISEASE_NAME}临床诊疗流程"
        fig31_nodes = ["首诊评估", "病因识别", "风险分层", "治疗启动", "复评校准", "长期管理"]
        fig53_title = f"{DISEASE_NAME}全周期管理流程"
        fig53_nodes = ["症状识别", "初诊评估", "治疗启动", "复评调整", "复发预防", "长期随访"]

    fig62_title = "图表6-2：医保支付与监管联动对用药结构的影响路径"
    fig62_boxes = [
        {"label": "审评审批", "x": 0.10, "y": 0.62},
        {"label": "质量控制", "x": 0.33, "y": 0.62},
        {"label": "医保支付", "x": 0.56, "y": 0.62},
        {"label": "终端执行", "x": 0.79, "y": 0.62},
        {"label": "用药结构优化", "x": 0.45, "y": 0.28},
    ]
    fig62_arrows = [
        {"x1": 0.16, "y1": 0.62, "x2": 0.27, "y2": 0.62},
        {"x1": 0.39, "y1": 0.62, "x2": 0.50, "y2": 0.62},
        {"x1": 0.62, "y1": 0.62, "x2": 0.73, "y2": 0.62},
        {"x1": 0.56, "y1": 0.56, "x2": 0.49, "y2": 0.35},
    ]

    return {
        "fig_1_3": {
            "flow_title": fig13_title,
            "nodes": fig13_nodes,
            "source_line": "REPLACE_WITH_SPECIFIC_SOURCE_LINE",
        },
        "fig_3_1": {
            "flow_title": fig31_title,
            "nodes": fig31_nodes,
            "source_line": "REPLACE_WITH_SPECIFIC_SOURCE_LINE",
        },
        "fig_5_3": {
            "flow_title": fig53_title,
            "nodes": fig53_nodes,
            "source_line": "REPLACE_WITH_SPECIFIC_SOURCE_LINE",
        },
        "fig_6_2": {
            "title": fig62_title,
            "boxes": fig62_boxes,
            "arrows": fig62_arrows,
            "source_line": "REPLACE_WITH_SPECIFIC_SOURCE_LINE",
        },
    }


def build_figure_specs_codex_prompt() -> str:
    return "\n".join(
        [
            "[Codex Semantic Figure Spec Task]",
            f"Goal: if topic semantics need better control, write {OUT_ROOT / 'figure_specs.json'} using {OUT_ROOT / FIGURE_SPECS_CODEX_TEMPLATE_NAME} as the starting point.",
            "",
            "[Scope]",
            "- This file is only for semantic figures such as fig_1_3, fig_3_1, fig_5_3, and fig_6_2.",
            "- Do not use it to invent Chapter-4 market data or override Excel-derived charts.",
            "",
            "[Read First]",
            f"1) {OUT_ROOT / 'ch01.txt'}",
            f"2) {OUT_ROOT / 'ch03.txt'}",
            f"3) {OUT_ROOT / 'ch05.txt'}",
            f"4) {OUT_ROOT / 'ch06.txt'}",
            f"5) {OUT_ROOT / '00_evidence.txt'}",
            f"6) {OUT_ROOT / FIGURE_SPECS_CODEX_TEMPLATE_NAME}",
            "",
            "[Hard Rules]",
            "1) Keep figure meaning aligned with the chapter argument instead of using generic placeholder flowcharts.",
            "2) Prefer fewer, clearer nodes over dense diagrams.",
            "3) Shorten labels before adding more boxes or arrows.",
            "4) For fig_6_2, keep enough whitespace so arrows do not collide with boxes.",
            "5) Every source_line must be specific enough to survive QA.",
            "",
            "[Done Means]",
            f"- Save the final override file to {OUT_ROOT / 'figure_specs.json'}.",
            "- If the built-in defaults are already correct, you may keep the file minimal and only override what truly needs semantic control.",
        ]
    )


def build_semantic_review_prompt() -> str:
    return "\n".join(
        [
            "[Codex Semantic Review Task]",
            "Review the body-summary logic and the key semantic figures together:",
            f"- {OUT_ROOT / 'ch01.txt'}",
            f"- {OUT_ROOT / 'ch03.txt'}",
            f"- {OUT_ROOT / 'ch05.txt'}",
            f"- {OUT_ROOT / 'ch06.txt'}",
            f"- {OUT_ROOT / 'summary.txt'}",
            f"- {FIG_DIR / 'fig_1_3.png'}",
            f"- {FIG_DIR / 'fig_3_1.png'}",
            f"- {FIG_DIR / 'fig_5_3.png'}",
            f"- {FIG_DIR / 'fig_6_2.png'}",
            "",
            "[Check List]",
            "- Does the narrative stay specific to the topic instead of collapsing into generic disease-market language?",
            "- Do the key figures match the chapter logic rather than acting as interchangeable placeholders?",
            "- Are summary conclusions actually supported by the chapter body and figure semantics?",
            "- Are there over-generalized claims, weak causal jumps, or obvious management-jargon filler?",
            "",
            "[Return Format]",
            "1) Verdict: PASS / FAIL",
            "2) Must-fix text issues: none / bullet list",
            "3) Must-fix figure issues: none / bullet list",
            "4) Suggested edits: direct file-level recommendations with the smallest useful changes first",
        ]
    )


def _df_to_records(df: pd.DataFrame, columns: List[str]) -> List[Dict[str, object]]:
    if df.empty:
        return []
    out: List[Dict[str, object]] = []
    for _, row in df[columns].iterrows():
        rec: Dict[str, object] = {}
        for col in columns:
            val = row[col]
            if isinstance(val, str):
                rec[col] = val.strip()
            elif isinstance(val, (np.integer, int)):
                rec[col] = int(val)
            elif isinstance(val, (np.floating, float)):
                rec[col] = round(float(val), 4)
            else:
                rec[col] = val
        out.append(rec)
    return out


def ensure_ch4_extract_ready(xlsx: Path) -> None:
    write_ch4_codex_helper_files(xlsx)
    extract_path = OUT_ROOT / "ch04_codex_extract.json"
    if extract_path.exists():
        try:
            build_ch4_data_from_codex_extract(extract_path)
            return
        except Exception as exc:
            print(f"警告：现有 ch04_codex_extract.json 无法通过校验，已自动重建（{type(exc).__name__}: {exc}）。")
    ch4 = build_ch4_data_from_legacy_parser(xlsx)
    sheet_names = get_workbook_sheet_names(xlsx)
    sheet_name_set = set(sheet_names)

    def mapping_entry(sheet_name: str, missing_note: str, present_note: str = "脚本按标准模板提取") -> Dict[str, str]:
        if sheet_name in sheet_name_set:
            return {"sheet": sheet_name, "status": "自动识别", "header_rows": "自动", "note": present_note}
        return {"sheet": sheet_name, "status": "缺失", "header_rows": "N/A", "note": missing_note}

    payload = {
        "schema_version": "ch4_codex_extract_v1",
        "topic": DISEASE_NAME,
        "disease": DISEASE_NAME,
        "source_workbook": xlsx.name,
        "available_sheets": sheet_names,
        "latest_quarter": ch4.latest_quarter,
        "sheet_mapping": {
            "hospital_category": mapping_entry("医院品类", "未提供该sheet，若其他渠道存在则以0补齐该渠道季度销售额。"),
            "hospital_top": mapping_entry("医院top", "未提供该sheet，医院端Top10与CR5仅保留为空。"),
            "drugstore_category": mapping_entry("药店品类", "未提供该sheet，若其他渠道存在则以0补齐该渠道季度销售额。"),
            "drugstore_top": mapping_entry("药店top", "未提供该sheet，药店端Top10与CR5仅保留为空。"),
            "online_category": mapping_entry("线上品类", "未提供该sheet，若其他渠道存在则以0补齐该渠道季度销售额。"),
            "online_top": mapping_entry("线上top", "未提供该sheet，线上端Top10与CR5仅保留为空。"),
        },
        "tables": {
            "quarterly_channel": _df_to_records(ch4.quarterly, ["quarter", "hospital", "drugstore", "online"]),
            "top10_hospital": _df_to_records(ch4.top_hospital, ["rank", "name", "sales"]),
            "top10_drugstore": _df_to_records(ch4.top_drugstore, ["rank", "name", "sales"]),
            "top10_online": _df_to_records(ch4.top_online, ["rank", "name", "sales"]),
            "cr5_latest": _df_to_records(ch4.cr5_latest, ["channel", "cr5_pct"]),
            "cr5_trend": _df_to_records(ch4.cr5_trend, ["quarter", "channel", "cr5_pct"]),
        },
        "notes": [
            "由脚本按标准sheet自动抽取生成。",
            "若需人工复核，请对照 ch04_workbook_preview.txt 与 ch04_sheet_map.txt。",
            "所有数值均直接来自工作簿，不做平滑、不做外推。",
            "若源工作簿缺少药店端或线上端sheet，季度值按0补齐；缺失渠道的Top10/CR5保留为空。",
        ],
    }
    write_json(extract_path, payload)


def ensure_codex_prep_assets_ready() -> None:
    ensure_ch4_extract_ready(EXCEL_PATH)
    ch4 = build_ch4_data(EXCEL_PATH)
    write_ch4_profile_files(ch4)
    write_text(OUT_ROOT / CH4_NARRATIVE_BRIEF_NAME, build_ch4_narrative_brief(ch4) + "\n")


def make_manifest_files(specs: List[BlockSpec], fig_rows: List[Dict[str, str]]) -> None:
    text_rows = []
    for s in specs:
        text_rows.append(
            {
                "block_id": s.block_id,
                "章": str(s.chapter),
                "标题": s.subtitle,
                "目标字数": str(s.target_chars),
                "可选/必选图表ID": s.fig_ids,
                "证据ID": s.evidence_ids,
                "插入锚点anchor": f"<<S{s.block_id}>>",
            }
        )
    write_csv(
        OUT_ROOT / "manifest_text.csv",
        text_rows,
        ["block_id", "章", "标题", "目标字数", "可选/必选图表ID", "证据ID", "插入锚点anchor"],
    )
    write_csv(
        OUT_ROOT / "manifest_fig.csv",
        fig_rows,
        ["fig_id", "caption", "type", "data_source", "数据表来源", "excel_sheet_or_table", "输出文件名", "插入到哪个block之后", "规则标签", "source_line"],
    )
    ch4_rows = [r for r in fig_rows if r["fig_id"].startswith("fig_4_")]
    write_csv(
        OUT_ROOT / "ch04_manifest_fig.csv",
        ch4_rows,
        ["fig_id", "caption", "type", "data_source", "数据表来源", "excel_sheet_or_table", "输出文件名", "插入到哪个block之后", "规则标签", "source_line"],
    )


def setup_figure_style():
    plt.style.use("seaborn-v0_8-whitegrid")
    # Re-apply Chinese-capable font settings after loading style.
    matplotlib.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "SimSun", "Microsoft JhengHei", "DejaVu Sans"]
    matplotlib.rcParams["font.family"] = "sans-serif"
    matplotlib.rcParams["axes.unicode_minus"] = False


def save_figure(path: Path, fig) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    backup_if_exists(path)
    tight_rect = getattr(fig, "_codex_tight_rect", None)
    if isinstance(tight_rect, (list, tuple)) and len(tight_rect) == 4:
        fig.tight_layout(rect=tight_rect)
    else:
        fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def estimate_flow_box_width(
    text: str,
    min_width: float = 0.10,
    max_width: float = 0.22,
    base: float = 0.064,
    char_step: float = 0.0070,
) -> float:
    raw = normalize_disease_text(text).replace("\n", "").strip()
    visual_len = 0.0
    for char in raw:
        if char.isspace():
            continue
        visual_len += 0.62 if ord(char) < 128 else 1.0
    return min(max_width, max(min_width, base + visual_len * char_step))


def flow_text_visual_len(text: str) -> float:
    raw = normalize_disease_text(text).replace("\n", "").strip()
    visual_len = 0.0
    for char in raw:
        if char.isspace():
            continue
        visual_len += 0.62 if ord(char) < 128 else 1.0
    return visual_len


def wrap_flow_label(text: str, max_visual_per_line: float = 8.6) -> str:
    raw = normalize_disease_text(text).replace("\n", " ").strip()
    if not raw or flow_text_visual_len(raw) <= max_visual_per_line:
        return raw

    tokens = [tok for tok in re.split(r"(\s+)", raw) if tok]
    if len(tokens) > 1 and any(not tok.isspace() for tok in tokens):
        line1: List[str] = []
        current = ""
        for tok in tokens:
            tentative = current + tok
            if current and flow_text_visual_len(tentative) > max_visual_per_line:
                break
            line1.append(tok)
            current = tentative
        left = "".join(line1).strip()
        right = raw[len(left):].strip()
        if left and right:
            return f"{left}\n{right}"

    total = flow_text_visual_len(raw)
    target = total / 2.0
    running = 0.0
    split_idx = max(1, len(raw) // 2)
    for idx, char in enumerate(raw, start=1):
        if not char.isspace():
            running += 0.62 if ord(char) < 128 else 1.0
        if running >= target:
            split_idx = idx
            break
    left = raw[:split_idx].strip()
    right = raw[split_idx:].strip()
    if not left or not right:
        return raw
    return f"{left}\n{right}"


def flow_box_height(text: str, base_height: float = 0.12) -> float:
    return base_height + (0.04 if "\n" in normalize_disease_text(text) else 0.0)


def layout_horizontal_flow_nodes(
    texts: List[str],
    left_margin: float = 0.04,
    right_margin: float = 0.04,
    min_gap: float = 0.08,
    max_gap: float = 0.15,
    min_width: float = 0.105,
    max_width: float = 0.19,
    base: float = 0.062,
    char_step: float = 0.0066,
) -> Tuple[List[str], List[float], float]:
    labels = [wrap_flow_label(text) for text in texts]
    widths = [estimate_flow_box_width(label, min_width=min_width, max_width=max_width, base=base, char_step=char_step) for label in labels]
    n_items = max(len(labels), 1)
    available = 1.0 - left_margin - right_margin
    content_width = sum(widths)
    if n_items > 1:
        gap = max(min_gap, min(max_gap, (available - content_width) / (n_items - 1)))
    else:
        gap = 0.0
    total_width = content_width + gap * max(0, n_items - 1)
    if total_width > available and content_width > 0:
        target_width = max(available - min_gap * max(0, n_items - 1), min_width * n_items)
        scale = min(1.0, target_width / content_width)
        widths = [max(min_width, w * scale) for w in widths]
        content_width = sum(widths)
        if n_items > 1:
            gap = max(0.045, min(max_gap, (available - content_width) / (n_items - 1)))
    return labels, widths, gap


def suggest_flow_figsize(nodes: List[str], direction: str, figsize: Tuple[float, float]) -> Tuple[float, float]:
    width, height = float(figsize[0]), float(figsize[1])
    if direction != "lr":
        return (width, height)
    total_visual = sum(max(1.0, flow_text_visual_len(text)) for text in nodes)
    suggested_width = min(15.2, max(width, 1.85 * max(len(nodes), 1) + 0.11 * total_visual))
    return (suggested_width, height)


def draw_simple_flow(path: Path, title: str, nodes: List[str], direction: str = "lr", color: str = "#2B6CB0", figsize=(10, 3.5)) -> None:
    figsize = suggest_flow_figsize(nodes, direction, figsize)
    fig, ax = plt.subplots(figsize=figsize)
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)

    n_nodes = max(len(nodes), 1)

    if direction == "lr":
        left_margin = 0.04
        labels, widths, gap = layout_horizontal_flow_nodes(
            nodes,
            left_margin=left_margin,
            right_margin=0.04,
            min_gap=0.082 if n_nodes >= 5 else 0.095,
            max_gap=0.15,
            min_width=0.105,
            max_width=0.185,
            base=0.062,
            char_step=0.0064,
        )
        available = 1.0 - left_margin - 0.04
        total_width = sum(widths) + gap * max(0, n_nodes - 1)

        anchors: List[Dict[str, Tuple[float, float]]] = []
        cursor = left_margin + max(0.0, (available - total_width) / 2.0)
        for width, text in zip(widths, labels):
            x = cursor + width / 2.0
            anchors.append(
                draw_box_node(
                    ax,
                    x,
                    0.50,
                    text,
                    width=width,
                    height=flow_box_height(text, 0.118),
                    fc="#E6F2FF",
                    ec=color,
                    lw=1.2,
                    fontsize=8.9 if "\n" in text else 9.2,
                )
            )
            cursor += width + gap

        for i in range(len(anchors) - 1):
            src = anchors[i]["east"]
            dst = anchors[i + 1]["west"]
            draw_poly_arrow(
                ax,
                [src, dst],
                color=color,
                lw=1.6,
                shrink_start_pts=10.0,
                shrink_end_pts=10.0,
            )
    else:
        labels = [wrap_flow_label(text) for text in nodes]
        heights = [estimate_flow_box_width(text, min_width=0.13, max_width=0.22) * 0.72 + (0.03 if "\n" in text else 0.0) for text in labels]
        gap = 0.05
        available = 0.82
        total_height = sum(heights) + gap * max(0, n_nodes - 1)
        if total_height > available and sum(heights) > 0:
            scale = max(0.60, (available - gap * max(0, n_nodes - 1)) / sum(heights))
            heights = [max(0.09, h * scale) for h in heights]

        anchors = []
        cursor = 0.91
        for height, text in zip(heights, labels):
            y = cursor - height / 2.0
            anchors.append(draw_box_node(ax, 0.50, y, text, width=0.30, height=height, fc="#E6F2FF", ec=color, lw=1.2, fontsize=8.9 if "\n" in text else 9.2))
            cursor -= height + gap

        for i in range(len(anchors) - 1):
            src = anchors[i]["south"]
            dst = anchors[i + 1]["north"]
            draw_poly_arrow(
                ax,
                [src, dst],
                color=color,
                lw=1.6,
                shrink_start_pts=10.0,
                shrink_end_pts=10.0,
            )

    ax.set_title(normalize_disease_text(title), fontsize=12, pad=10, fontweight="bold")
    save_figure(path, fig)


def draw_pie_with_leaders(path: Path, title: str, labels: List[str], values: List[float], colors: List[str], figsize=(7.6, 4.6)) -> None:
    fig, ax = plt.subplots(figsize=figsize)
    wedges, _ = ax.pie(values, labels=None, startangle=90, colors=colors, wedgeprops={"linewidth": 1, "edgecolor": "white"})
    ax.axis("equal")

    def distribute_side(items: List[Dict[str, float]], low: float = -1.18, high: float = 1.18, min_gap: float = 0.12) -> List[float]:
        if not items:
            return []
        items = sorted(items, key=lambda item: item["target_y"])
        ys = [max(low, min(high, item["target_y"])) for item in items]
        for idx in range(1, len(ys)):
            ys[idx] = max(ys[idx], ys[idx - 1] + min_gap)
        overflow = ys[-1] - high
        if overflow > 0:
            ys = [y - overflow for y in ys]
        for idx in range(len(ys) - 2, -1, -1):
            ys[idx] = min(ys[idx], ys[idx + 1] - min_gap)
        underflow = low - ys[0]
        if underflow > 0:
            ys = [y + underflow for y in ys]
        return ys

    annotations: List[Dict[str, object]] = []
    for i, w in enumerate(wedges):
        ang = (w.theta2 + w.theta1) / 2.0
        x = math.cos(math.radians(ang))
        y = math.sin(math.radians(ang))
        annotations.append(
            {
                "index": i,
                "x": x,
                "y": y,
                "side": 1 if x >= 0 else -1,
                "target_y": 1.16 * y,
            }
        )

    right_items = [item for item in annotations if int(item["side"]) > 0]
    left_items = [item for item in annotations if int(item["side"]) < 0]
    right_y = distribute_side(right_items)
    left_y = distribute_side(left_items)

    for item, adjusted_y in zip(right_items, right_y):
        item["adjusted_y"] = adjusted_y
    for item, adjusted_y in zip(left_items, left_y):
        item["adjusted_y"] = adjusted_y

    for item in annotations:
        idx = int(item["index"])
        x = float(item["x"])
        y = float(item["y"])
        side = int(item["side"])
        adjusted_y = float(item.get("adjusted_y", 1.16 * y))
        label = f"{normalize_disease_text(labels[idx])} {values[idx]:.1f}%"
        ax.annotate(
            label,
            xy=(x * 0.82, y * 0.82),
            xytext=(1.34 * side, adjusted_y),
            ha="left" if side > 0 else "right",
            va="center",
            fontsize=8.7,
            arrowprops=dict(
                arrowstyle="-",
                color="#444444",
                lw=0.9,
                shrinkA=0,
                shrinkB=0,
                connectionstyle="arc3,rad=0",
            ),
        )

    ax.set_title(normalize_disease_text(title), fontsize=12, fontweight="bold")
    save_figure(path, fig)


def draw_policy_timeline(path: Path, title: str, events: List[Tuple[str, str]], figsize=(8.2, 3.0)) -> None:
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_xlim(0, len(events) + 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.hlines(0.5, 0.8, len(events) + 0.2, color="#2C5282", lw=2.0)
    for idx, (year, txt) in enumerate(events, start=1):
        ax.plot(idx, 0.5, "o", color="#2C5282")
        ax.text(idx, 0.62, year, ha="center", va="bottom", fontsize=9, fontweight="bold")
        ax.text(idx, 0.36, normalize_disease_text(txt), ha="center", va="top", fontsize=8, wrap=True)
    ax.set_title(normalize_disease_text(title), fontsize=12, fontweight="bold", pad=8)
    save_figure(path, fig)


def draw_box_node(
    ax,
    x: float,
    y: float,
    label: str,
    width: float = 0.28,
    height: float = 0.12,
    fc: str = "#EDF2F7",
    ec: str = "#2D3748",
    lw: float = 1.1,
    fontsize: float = 9.5,
) -> Dict[str, Tuple[float, float]]:
    box = FancyBboxPatch(
        (x - width / 2, y - height / 2),
        width,
        height,
        boxstyle="round,pad=0.02",
        fc=fc,
        ec=ec,
        lw=lw,
    )
    ax.add_patch(box)
    ax.text(x, y, normalize_disease_text(label), ha="center", va="center", fontsize=fontsize)
    return {
        "center": (x, y),
        "north": (x, y + height / 2),
        "south": (x, y - height / 2),
        "west": (x - width / 2, y),
        "east": (x + width / 2, y),
    }


def draw_poly_arrow(
    ax,
    points: List[Tuple[float, float]],
    color: str = "#2B6CB0",
    lw: float = 1.2,
    dashed: bool = False,
    shrink_start_pts: float = 0.0,
    shrink_end_pts: float = 0.0,
) -> None:
    if len(points) < 2:
        return
    ls = "--" if dashed else "-"
    for i in range(len(points) - 2):
        x1, y1 = points[i]
        x2, y2 = points[i + 1]
        ax.plot([x1, x2], [y1, y2], color=color, lw=lw, linestyle=ls)
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    ax.annotate(
        "",
        xy=(x2, y2),
        xytext=(x1, y1),
        arrowprops=dict(
            arrowstyle="->",
            lw=lw,
            color=color,
            linestyle=ls,
            shrinkA=max(0.0, float(shrink_start_pts)) if len(points) == 2 else 0.0,
            shrinkB=max(0.0, float(shrink_end_pts)),
        ),
    )


def polyline_point_at(points: List[Tuple[float, float]], frac: float = 0.5) -> Tuple[float, float]:
    if not points:
        return (0.5, 0.5)
    if len(points) == 1:
        return points[0]
    frac = max(0.0, min(1.0, frac))
    seg_lengths: List[float] = []
    total = 0.0
    for idx in range(len(points) - 1):
        x1, y1 = points[idx]
        x2, y2 = points[idx + 1]
        seg = max(math.hypot(x2 - x1, y2 - y1), 1e-9)
        seg_lengths.append(seg)
        total += seg
    target = frac * total
    walked = 0.0
    for idx, seg in enumerate(seg_lengths):
        if walked + seg >= target:
            x1, y1 = points[idx]
            x2, y2 = points[idx + 1]
            t = (target - walked) / seg if seg > 0 else 0.0
            return (x1 + (x2 - x1) * t, y1 + (y2 - y1) * t)
        walked += seg
    return points[-1]


def draw_fig23_layered_path(ax, cfg: Dict[str, object]) -> None:
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    left_title = render_disease_template(str(cfg.get("left_title", "上游决定因素")).strip())
    center_title = render_disease_template(str(cfg.get("center_title", "核心病理/修复过程")).strip())
    right_title = render_disease_template(str(cfg.get("right_title", "临床后果与管理结果")).strip())

    ax.text(0.18, 0.92, left_title, ha="center", va="center", fontsize=10, fontweight="bold")
    ax.text(0.50, 0.92, center_title, ha="center", va="center", fontsize=10, fontweight="bold")
    ax.text(0.82, 0.92, right_title, ha="center", va="center", fontsize=10, fontweight="bold")

    left_nodes_raw = cfg.get("left_nodes", [])
    right_nodes_raw = cfg.get("right_nodes", [])
    core_label = render_disease_template(str(cfg.get("core_label", "核心病理")).strip())

    left_nodes = [render_disease_template(str(x).strip()) for x in left_nodes_raw if str(x).strip()]
    right_nodes = [render_disease_template(str(x).strip()) for x in right_nodes_raw if str(x).strip()]
    if not left_nodes:
        left_nodes = ["上游因素A", "上游因素B", "上游因素C"]
    if not right_nodes:
        right_nodes = ["下游结果A", "下游结果B", "下游结果C"]

    left_y = np.linspace(0.74, 0.26, len(left_nodes))
    right_y = np.linspace(0.74, 0.26, len(right_nodes))
    left_anchors: List[Dict[str, Tuple[float, float]]] = []
    right_anchors: List[Dict[str, Tuple[float, float]]] = []

    for y, text in zip(left_y, left_nodes):
        left_anchors.append(draw_box_node(ax, 0.18, float(y), text, width=0.24, height=0.12, fc="#EDF2F7", ec="#2D3748", lw=1.1, fontsize=9.0))
    core_anchor = draw_box_node(ax, 0.50, 0.50, core_label, width=0.28, height=0.16, fc="#FEEBC8", ec="#C05621", lw=1.2, fontsize=9.0)
    for y, text in zip(right_y, right_nodes):
        right_anchors.append(draw_box_node(ax, 0.82, float(y), text, width=0.24, height=0.12, fc="#EDF2F7", ec="#2D3748", lw=1.1, fontsize=9.0))

    for anchor in left_anchors:
        src = anchor["east"]
        dst = core_anchor["west"]
        mid_x = 0.33
        points = [(src[0] + 0.012, src[1]), (mid_x, src[1]), (mid_x, dst[1]), (dst[0] - 0.018, dst[1])]
        draw_poly_arrow(ax, points, color="#2B6CB0", lw=1.25)

    for anchor in right_anchors:
        src = core_anchor["east"]
        dst = anchor["west"]
        mid_x = 0.67
        points = [(src[0] + 0.018, src[1]), (mid_x, src[1]), (mid_x, dst[1]), (dst[0] - 0.012, dst[1])]
        draw_poly_arrow(ax, points, color="#2B6CB0", lw=1.25)

    ax.text(0.50, 0.08, "注：本图采用分层路径表达，替代高密度关系网络，以避免线条覆盖文本框。", ha="center", va="center", fontsize=8.2, color="#4A5568")


def draw_configured_network_panel(ax, panel_cfg: Dict[str, object], default_edge_color: str = "#2B6CB0") -> None:
    nodes_cfg = panel_cfg.get("nodes", [])
    edges_cfg = panel_cfg.get("edges", [])
    node_map: Dict[str, Dict[str, Tuple[float, float]]] = {}

    if isinstance(nodes_cfg, list):
        for node in nodes_cfg:
            if not isinstance(node, dict):
                continue
            node_id = str(node.get("id", "")).strip()
            if not node_id:
                continue
            x = float(node.get("x", 0.5))
            y = float(node.get("y", 0.5))
            width = float(node.get("width", node.get("w", 0.30)))
            height = float(node.get("height", node.get("h", 0.12)))
            label_tpl = str(node.get("label", node_id)).strip().replace("\\n", "\n")
            label = render_disease_template(label_tpl, default=node_id)
            fc = str(node.get("fc", "#EDF2F7")).strip() or "#EDF2F7"
            ec = str(node.get("ec", "#2D3748")).strip() or "#2D3748"
            lw = float(node.get("lw", 1.1))
            fontsize = float(node.get("fontsize", 9.5))
            node_map[node_id] = draw_box_node(ax, x, y, label, width=width, height=height, fc=fc, ec=ec, lw=lw, fontsize=fontsize)

    if isinstance(edges_cfg, list):
        for edge in edges_cfg:
            if not isinstance(edge, dict):
                continue
            src = str(edge.get("from", "")).strip()
            dst = str(edge.get("to", "")).strip()
            if src not in node_map or dst not in node_map:
                continue
            from_anchor = str(edge.get("from_anchor", "center")).strip() or "center"
            to_anchor = str(edge.get("to_anchor", "center")).strip() or "center"
            src_pt = node_map[src].get(from_anchor, node_map[src]["center"])
            dst_pt = node_map[dst].get(to_anchor, node_map[dst]["center"])

            via_pts: List[Tuple[float, float]] = []
            raw_via = edge.get("via", [])
            if isinstance(raw_via, list):
                for p in raw_via:
                    if isinstance(p, (list, tuple)) and len(p) >= 2:
                        via_pts.append((float(p[0]), float(p[1])))

            points = [src_pt] + via_pts + [dst_pt]
            color = str(edge.get("color", default_edge_color)).strip() or default_edge_color
            lw = float(edge.get("lw", 1.2))
            dashed = bool(edge.get("dashed", False))
            draw_poly_arrow(ax, points, color=color, lw=lw, dashed=dashed)

            sign = str(edge.get("sign", "")).strip()
            if sign:
                raw_sign_xy = edge.get("sign_xy")
                if isinstance(raw_sign_xy, (list, tuple)) and len(raw_sign_xy) >= 2:
                    sign_x, sign_y = float(raw_sign_xy[0]), float(raw_sign_xy[1])
                else:
                    sign_frac = float(edge.get("sign_t", 0.5))
                    sign_x, sign_y = polyline_point_at(points, sign_frac)
                    raw_offset = edge.get("sign_offset", [])
                    if isinstance(raw_offset, (list, tuple)) and len(raw_offset) >= 2:
                        sign_x += float(raw_offset[0])
                        sign_y += float(raw_offset[1])
                bbox = dict(boxstyle="round,pad=0.10", fc="white", ec="none", alpha=0.88) if bool(edge.get("sign_bg", True)) else None
                ax.text(
                    sign_x,
                    sign_y,
                    sign,
                    fontsize=float(edge.get("sign_fontsize", 8.5)),
                    color=color,
                    fontweight="bold",
                    ha="center",
                    va="center",
                    bbox=bbox,
                )

    title = render_disease_template(str(panel_cfg.get("title", "")).strip().replace("\\n", "\n"))
    if title:
        title_y = float(panel_cfg.get("title_y", 1.02))
        ax.text(0.5, title_y, title, transform=ax.transAxes, ha="center", va="bottom", fontsize=10, fontweight="bold")



def generate_figures(ch4: Ch4Data) -> List[Dict[str, str]]:
    setup_figure_style()
    fig_rows: List[Dict[str, str]] = []
    rendered_title_rows: List[Dict[str, str]] = []
    figure_specs = load_figure_specs()
    block_spec_map = {s.block_id: s for s in build_block_specs()}
    evidence_rows, _ = parse_evidence_pool(OUT_ROOT / "00_evidence.txt")
    evidence_map = {str(row.get("evidence_id", "")).strip(): row for row in evidence_rows if str(row.get("evidence_id", "")).strip()}

    def fig_spec(fig_id: str) -> Dict[str, object]:
        v = figure_specs.get(fig_id, {})
        return v if isinstance(v, dict) else {}

    def spec_text(fig_id: str, key: str, default: str) -> str:
        v = fig_spec(fig_id).get(key, default)
        return render_disease_template(str(v), default=default)

    def spec_list(fig_id: str, key: str, default: List[object]) -> List[object]:
        v = fig_spec(fig_id).get(key)
        if isinstance(v, list) and len(v) > 0:
            return v
        return default

    def spec_num_list(fig_id: str, key: str, default: List[float]) -> List[float]:
        v = fig_spec(fig_id).get(key)
        if isinstance(v, list) and len(v) > 0:
            out: List[float] = []
            for x in v:
                try:
                    out.append(float(x))
                except Exception:
                    return default
            return out
        return default

    def spec_matrix(fig_id: str, key: str, default: np.ndarray) -> np.ndarray:
        v = fig_spec(fig_id).get(key)
        if not isinstance(v, list) or not v:
            return default
        try:
            arr = np.array(v, dtype=float)
        except Exception:
            return default
        if arr.ndim != 2 or arr.shape != default.shape:
            return default
        return arr

    def _short_source_title(title: str, max_len: int = 24) -> str:
        title = normalize_disease_text(title).replace("\n", " ").strip()
        title = re.sub(r"\s+", " ", title)
        if len(title) <= max_len:
            return title
        return title[:max_len].rstrip(" ，,;；") + "…"

    def evidence_source_labels(block_id: str) -> List[str]:
        spec = block_spec_map.get(str(block_id).strip())
        if not spec:
            return []
        labels: List[str] = []
        seen: set[str] = set()
        for evidence_id in str(spec.evidence_ids).split("|"):
            row = evidence_map.get(evidence_id.strip())
            if not row:
                continue
            org = normalize_disease_text(str(row.get("org", "")).strip())
            title = _short_source_title(str(row.get("title", "")).strip())
            label = ""
            if org and title:
                label = f"{org}《{title}》"
            elif title:
                label = f"《{title}》"
            elif org:
                label = org
            if not label or label in seen:
                continue
            seen.add(label)
            labels.append(label)
        return labels[:3]

    def evidence_source_line(block_id: str) -> str | None:
        labels = evidence_source_labels(block_id)
        if not labels:
            return None
        return "数据来源：" + "、".join(labels)

    def evidence_data_source(block_id: str) -> str | None:
        labels = evidence_source_labels(block_id)
        if not labels:
            return None
        compact: List[str] = []
        seen: set[str] = set()
        for label in labels:
            item = re.sub(r"《.*?》", "", label).strip(" 《》")
            if not item or item in seen:
                continue
            seen.add(item)
            compact.append(item)
        return " / ".join(compact) if compact else None

    def specific_public_data_source(fig_id: str) -> str | None:
        profile = active_profile_id()
        if profile == "sciatica":
            mapping = {
                "fig_1_1": "NICE/NASS/AANS疾病定义与分型资料",
                "fig_1_2": "WHO/NICE/神经根炎症公开研究",
                "fig_1_3": "NASS/系统综述/疼痛敏化研究",
                "fig_1_4": "WHO/GBD/国家卫健委公开资料",
                "fig_2_1": "NICE/AANS/康复系统综述",
                "fig_2_2": "NICE/红旗征评估/预后系统综述",
                "fig_2_3": "NICE/AANS/系统综述等公开资料",
                "fig_3_1": "NICE/NASS/AANS诊断路径资料",
                "fig_3_2": "NICE/NASS/Cochrane/随机对照研究",
                "fig_3_3": "针刺随机对照研究/康复系统综述",
                "fig_5_1": "国家统计局/国家卫健委/职业负荷公开研究",
                "fig_5_2": "NICE/NASS指南与临床路径公开文献",
                "fig_5_3": "身体活动指南/康复依从性系统综述",
                "fig_5_4": "依从性研究/康复随访公开研究",
                "fig_7_2": "国家统计局/国家医保局/米内网/公开竞争情报",
            }
            return mapping.get(str(fig_id).strip())
        return None

    def specific_public_source_line(fig_id: str) -> str | None:
        profile = active_profile_id()
        if profile == "sciatica":
            mapping = {
                "fig_1_1": "数据来源：NICE NG59、NASS腰椎间盘突出伴神经根病指南、AANS疾病教育资料整理",
                "fig_1_2": "数据来源：WHO《Musculoskeletal conditions》、NICE NG59、神经根炎症与疼痛敏化公开研究整理",
                "fig_1_3": "数据来源：StatPearls《Sciatica》、NASS指南及疼痛敏化系统综述整理",
                "fig_1_4": "数据来源：WHO《Musculoskeletal conditions》、GBD相关负担资料及国家卫健委公开资料整理",
                "fig_2_1": "数据来源：NICE NG59、AANS疾病科普、国家卫健委《身体活动指南》及康复系统综述整理",
                "fig_2_2": "数据来源：NICE NG59、红旗征识别公开文献、慢性腰腿痛预后系统综述整理",
                "fig_2_3": "数据来源：NICE NG59、AANS疾病科普、国家卫健委《身体活动指南》及物理治疗系统综述等公开资料整理",
                "fig_3_1": "数据来源：NICE NG59、NASS腰椎间盘突出伴神经根病指南、AANS临床教育资料整理",
                "fig_3_2": "数据来源：NICE NG59、NASS指南、Cochrane系统综述及随机对照研究整理",
                "fig_3_3": "数据来源：针刺随机对照研究、康复训练系统综述及中西协同公开文献整理",
                "fig_5_1": "数据来源：国家统计局人口年龄结构、国家卫健委慢性疼痛相关资料及职业人群腰背痛公开研究整理",
                "fig_5_2": "数据来源：NICE NG59、NASS指南及骨科/疼痛科/康复科诊疗路径公开文献整理",
                "fig_5_3": "数据来源：国家卫健委《身体活动指南》、康复依从性系统综述及患者教育公开资料整理",
                "fig_5_4": "数据来源：慢性腰痛/坐骨神经痛依从性研究及康复随访公开研究整理",
                "fig_7_2": "数据来源：国家统计局、国家卫健委、国家医保局、米内网项目数据及公开竞争情报整理",
            }
            return mapping.get(str(fig_id).strip())
        return None

    def default_public_data_source(fig_id: str) -> str:
        specific = specific_public_data_source(fig_id)
        if specific:
            return specific
        chapter = 0
        try:
            chapter = int(str(fig_id).split("_")[1])
        except Exception:
            chapter = 0
        profile = active_profile_id()
        if profile == "sciatica":
            mapping = {
                1: "WHO/NICE/NASS等公开资料",
                2: "NICE/AANS/系统综述等公开资料",
                3: "NICE/NASS/AANS/随机对照研究等公开资料",
                5: "国家统计局/国家医保局/公开研究",
                6: "国家医保局/国家卫健委/国家药监局",
                7: "国家统计局/国家卫健委/米内网/公开研究",
            }
            return mapping.get(chapter, "公开资料整理")
        return "公开资料整理"

    def default_public_source_line(fig_id: str, fallback: str) -> str:
        specific = specific_public_source_line(fig_id)
        if specific:
            return specific
        chapter = 0
        try:
            chapter = int(str(fig_id).split("_")[1])
        except Exception:
            chapter = 0
        profile = active_profile_id()
        if profile == "sciatica":
            mapping = {
                1: "数据来源：WHO《Musculoskeletal conditions》、NICE NG59、NASS腰椎间盘突出伴神经根病指南等公开资料整理",
                2: "数据来源：NICE NG59、AANS疾病科普、国家卫健委《身体活动指南》及物理治疗系统综述等公开资料整理",
                3: "数据来源：NICE NG59、NASS指南、AANS资料及针刺/物理治疗公开研究整理",
                5: "数据来源：国家统计局、国家医保局、国家卫健委及公开研究资料整理",
                6: "数据来源：国家医保局、国家卫健委、国家药监局法规与政策文件整理",
                7: "数据来源：国家统计局、国家卫健委、米内网项目数据及公开研究资料整理",
            }
            return mapping.get(chapter, fallback)
        return fallback

    def add_fig_meta(fig_id: str, caption: str, fig_type: str, data_source: str, table_src: str, excel_table: str, block_id: str, rule_tag: str, source_line: str):
        caption_text = spec_text(fig_id, "caption", caption)
        default_source_line = source_line
        default_data_source = data_source
        specific_source = specific_public_source_line(fig_id)
        specific_data = specific_public_data_source(fig_id)
        evidence_source = evidence_source_line(block_id)
        evidence_data = evidence_data_source(block_id)
        if specific_source:
            default_source_line = specific_source
        elif evidence_source and (("公开资料整理" in str(source_line)) or ("指南整理" in str(source_line)) or (not str(source_line).strip())):
            default_source_line = evidence_source
        elif "公开资料整理" in str(source_line):
            default_source_line = default_public_source_line(fig_id, source_line)
        if specific_data:
            default_data_source = specific_data
        elif evidence_data and (str(data_source).strip() in {"公开资料整理", "指南整理"}):
            default_data_source = evidence_data
        elif str(data_source).strip() == "公开资料整理":
            default_data_source = default_public_data_source(fig_id)
        source_text = spec_text(fig_id, "source_line", default_source_line)
        fig_rows.append(
            {
                "fig_id": fig_id,
                "caption": caption_text,
                "type": fig_type,
                "data_source": default_data_source,
                "数据表来源": table_src,
                "excel_sheet_or_table": excel_table,
                "输出文件名": f"{fig_id}.png",
                "插入到哪个block之后": block_id,
                "规则标签": rule_tag,
                "source_line": source_text,
            }
        )

    def set_main_title(ax, fig_id: str, title: str, **kwargs):
        txt = spec_text(fig_id, "title", title)
        ax.set_title(txt, **kwargs)
        rendered_title_rows.append(
            {
                "fig_id": fig_id,
                "rendered_title": txt,
            }
        )

    def set_main_suptitle(fig, fig_id: str, title: str, **kwargs):
        txt = spec_text(fig_id, "title", title)
        fig.suptitle(txt, **kwargs)
        rendered_title_rows.append(
            {
                "fig_id": fig_id,
                "rendered_title": txt,
            }
        )

    # Chapter 1
    if is_cervical_profile():
        cls = ["神经根型", "脊髓型", "椎动脉型", "交感型", "混合型"]
        vals = [46, 18, 14, 9, 13]
        stages = ["早期退变", "症状进展", "功能受限", "慢性稳定"]
        means = [38, 52, 67, 61]
        flow_title = f"{DISEASE_NAME}病理生理演进路径"
        flow_nodes = ["椎间盘退变", "椎体边缘骨赘", "椎管/椎间孔狭窄", "神经结构受压", "疼痛与功能障碍"]
        drivers = ["老龄化与久坐", "影像筛查普及", "康复需求", "门诊量增长", "指南更新"]
        score = [82, 71, 78, 69, 74]
    elif is_gastritis_profile():
        cls = ["Hp相关胃炎", "NSAIDs相关", "胆汁反流相关", "自身免疫相关", "其他病因"]
        vals = [42, 18, 15, 9, 16]
        stages = ["非萎缩期", "萎缩期", "肠化期", "异型增生风险期"]
        means = [44, 56, 68, 77]
        flow_title = f"{DISEASE_NAME}病理演进链路"
        flow_nodes = ["病因暴露", "慢性炎症", "腺体萎缩", "肠上皮化生", "风险升级"]
        drivers = ["中老年人群", "Hp感染管理", "胃镜筛查覆盖", "院外长期用药", "指南更新"]
        score = [79, 74, 71, 65, 72]
    elif is_respiratory_profile():
        cls = ["急性咳嗽", "迁延性咳嗽", "慢性咳嗽", "痰液黏稠型", "痉咳伴喘型"]
        vals = [28, 24, 16, 18, 14]
        stages = ["初发期", "进展期", "缓解期", "复发期"]
        means = [72, 64, 48, 58]
        flow_title = f"{DISEASE_NAME}病理生理演进路径"
        flow_nodes = ["感染/过敏触发", "炎症与分泌增加", "痰液黏稠", "排痰受阻", "持续咳嗽"]
        drivers = ["儿童人口", "门急诊需求", "院外购药", "数字触达", "指南更新"]
        score = [66, 79, 74, 61, 70]
    else:
        cls = ["轻度型", "中度型", "重度型", "慢性管理型", "复发风险型"]
        vals = [26, 31, 18, 14, 11]
        stages = ["初发期", "评估期", "干预期", "稳定期"]
        means = [66, 58, 49, 53]
        flow_title = f"{DISEASE_NAME}病理生理演进路径"
        flow_nodes = ["风险暴露", "病理改变", "功能受损", "症状负担上升", "分层管理"]
        drivers = ["人口结构变化", "就医需求", "院外管理", "数字触达", "指南更新"]
        score = [68, 73, 70, 62, 66]

    cls = [str(x) for x in spec_list("fig_1_1", "categories", cls)]
    vals = spec_num_list("fig_1_1", "values", [float(x) for x in vals])
    if len(cls) != len(vals):
        cls, vals = cls[: min(len(cls), len(vals))], vals[: min(len(cls), len(vals))]
    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    ax.bar(cls, vals, color=["#2B6CB0", "#3182CE", "#63B3ED", "#90CDF4", "#BEE3F8"])
    ax.set_ylabel("占比（%）")
    set_main_title(ax, "fig_1_1", f"图表1-1：{DISEASE_NAME}临床分型结构")
    save_figure(FIG_DIR / "fig_1_1.png", fig)
    add_fig_meta("fig_1_1", f"图表1-1：{DISEASE_NAME}临床分型结构", "柱状图", "公开资料整理", "分型结构整理", "N/A", "1.1", "分型框架", "数据来源：公开资料整理")

    stages = [str(x) for x in spec_list("fig_1_2", "x_labels", stages)]
    means = spec_num_list("fig_1_2", "values", [float(x) for x in means])
    if len(stages) != len(means):
        stages, means = stages[: min(len(stages), len(means))], means[: min(len(stages), len(means))]
    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    ax.plot(stages, means, marker="o", color="#2F855A", lw=2)
    ax.fill_between(stages, means, [min(means) - 12] * len(stages), color="#9AE6B4", alpha=0.35)
    ax.set_ylabel("症状负担指数")
    set_main_title(ax, "fig_1_2", f"图表1-2：{DISEASE_NAME}病程分期与症状负担变化")
    save_figure(FIG_DIR / "fig_1_2.png", fig)
    add_fig_meta("fig_1_2", f"图表1-2：{DISEASE_NAME}病程分期与症状负担变化", "折线图", "公开资料整理", "病程管理要点", "N/A", "1.2", "机制链路", "数据来源：公开资料整理")

    flow_title = spec_text("fig_1_3", "flow_title", flow_title)
    flow_nodes = [str(x) for x in spec_list("fig_1_3", "nodes", flow_nodes)]
    title_1_3 = compose_figure_title("1-3", flow_title)
    draw_simple_flow(FIG_DIR / "fig_1_3.png", title_1_3, flow_nodes, direction="lr", color="#2C5282", figsize=(9.6, 3.2))
    rendered_title_rows.append({"fig_id": "fig_1_3", "rendered_title": normalize_disease_text(title_1_3)})
    add_fig_meta("fig_1_3", title_1_3, "流程图", "公开资料整理", "机制链路", "N/A", "1.2", "机制路径", "数据来源：公开资料整理")

    drivers = [str(x) for x in spec_list("fig_1_4", "categories", drivers)]
    score = spec_num_list("fig_1_4", "values", [float(x) for x in score])
    if len(drivers) != len(score):
        drivers, score = drivers[: min(len(drivers), len(score))], score[: min(len(drivers), len(score))]
    fig, ax = plt.subplots(figsize=(7.8, 4.4))
    ax.barh(drivers, score, color="#805AD5")
    ax.set_xlim(0, 100)
    ax.set_xlabel("驱动强度指数")
    set_main_title(ax, "fig_1_4", f"图表1-4：{DISEASE_NAME}市场需求驱动强度")
    save_figure(FIG_DIR / "fig_1_4.png", fig)
    add_fig_meta("fig_1_4", f"图表1-4：{DISEASE_NAME}市场需求驱动强度", "条形图", "公开资料整理", "需求驱动评分", "N/A", "1.3", "需求结构", "数据来源：公开资料整理")

    # Chapter 2
    if is_cervical_profile():
        systems = ["神经系统", "肌肉骨骼系统", "血管系统", "睡眠系统", "心理行为", "内分泌代谢"]
        influence = [92, 88, 74, 66, 59, 52]
        xlabels = ["神经根痛", "脊髓受压", "眩晕不稳", "睡眠障碍", "焦虑抑郁"]
        ylabels = ["低风险", "中风险", "高风险"]
        matrix = np.array([[0.42, 0.18, 0.15, 0.28, 0.26], [0.41, 0.46, 0.38, 0.49, 0.43], [0.17, 0.36, 0.47, 0.23, 0.31]])
        pos = {
            "神经系统": (0.18, 0.78),
            "肌肉骨骼系统": (0.50, 0.78),
            "血管系统": (0.82, 0.78),
            "睡眠系统": (0.18, 0.28),
            "心理行为": (0.50, 0.28),
            "内分泌代谢": (0.82, 0.28),
        }
        edges = [
            ("神经系统", "肌肉骨骼系统", 0.05),
            ("神经系统", "睡眠系统", -0.08),
            ("血管系统", "神经系统", -0.12),
            ("肌肉骨骼系统", "心理行为", 0.08),
            ("睡眠系统", "心理行为", -0.08),
            ("内分泌代谢", "肌肉骨骼系统", 0.18),
            ("心理行为", "神经系统", 0.14),
        ]
    elif is_gastritis_profile():
        systems = ["消化系统", "免疫系统", "神经系统", "内分泌系统", "胃肠微生态", "血液系统"]
        influence = [94, 82, 68, 63, 71, 56]
        xlabels = ["溃疡风险", "贫血风险", "体重下降", "癌前病变", "复发负担"]
        ylabels = ["低风险", "中风险", "高风险"]
        matrix = np.array([[0.28, 0.22, 0.26, 0.18, 0.24], [0.47, 0.44, 0.43, 0.39, 0.46], [0.25, 0.34, 0.31, 0.43, 0.30]])
        pos = {
            "神经系统": (0.20, 0.78),
            "内分泌系统": (0.50, 0.78),
            "免疫系统": (0.80, 0.78),
            "消化系统": (0.20, 0.28),
            "血液系统": (0.50, 0.28),
            "胃肠微生态": (0.80, 0.28),
        }
        edges = [
            ("免疫系统", "消化系统", 0.12),
            ("内分泌系统", "消化系统", -0.10),
            ("消化系统", "血液系统", 0.08),
            ("消化系统", "胃肠微生态", -0.10),
            ("神经系统", "消化系统", -0.14),
            ("胃肠微生态", "内分泌系统", 0.10),
            ("血液系统", "神经系统", 0.06),
        ]
    elif is_sciatica_profile():
        systems = ["神经系统", "肌肉骨骼系统", "炎症免疫系统", "睡眠系统", "心理行为系统", "运动功能系统"]
        influence = [94, 89, 76, 68, 62, 84]
        xlabels = ["慢性神经痛", "肌力下降", "睡眠障碍", "运动回避", "复诊负担"]
        ylabels = ["低风险", "中风险", "高风险"]
        matrix = np.array([[0.24, 0.18, 0.21, 0.22, 0.19], [0.49, 0.46, 0.43, 0.44, 0.41], [0.27, 0.36, 0.36, 0.34, 0.40]])
        pos = {
            "神经系统": (0.18, 0.78),
            "肌肉骨骼系统": (0.50, 0.78),
            "炎症免疫系统": (0.82, 0.78),
            "睡眠系统": (0.18, 0.28),
            "心理行为系统": (0.50, 0.28),
            "运动功能系统": (0.82, 0.28),
        }
        edges = [
            ("神经系统", "睡眠系统", -0.10, "±"),
            ("神经系统", "心理行为系统", 0.06, "+"),
            ("肌肉骨骼系统", "运动功能系统", 0.00, "+"),
            ("炎症免疫系统", "神经系统", -0.10, "+"),
            ("睡眠系统", "心理行为系统", -0.08, "±"),
            ("心理行为系统", "运动功能系统", 0.08, "±"),
        ]
    elif is_respiratory_profile():
        systems = ["呼吸系统", "免疫系统", "消化系统", "神经系统", "内分泌系统", "肌肉骨骼系统"]
        influence = [92, 81, 63, 58, 46, 34]
        xlabels = ["急性加重", "睡眠受损", "反复就诊", "家长焦虑", "学习受影响"]
        ylabels = ["低风险", "中风险", "高风险"]
        matrix = np.array([[0.35, 0.40, 0.25, 0.31, 0.22], [0.45, 0.42, 0.51, 0.46, 0.39], [0.20, 0.18, 0.24, 0.23, 0.39]])
        pos = {
            "神经系统": (0.20, 0.78),
            "内分泌系统": (0.50, 0.78),
            "免疫系统": (0.80, 0.78),
            "呼吸系统": (0.20, 0.28),
            "消化系统": (0.50, 0.28),
            "肌肉骨骼系统": (0.80, 0.28),
        }
        edges = [
            ("神经系统", "呼吸系统", 0.0, "±"),
            ("内分泌系统", "呼吸系统", -0.18, "±"),
            ("免疫系统", "呼吸系统", 0.18, "+"),
            ("免疫系统", "消化系统", -0.12, "+"),
            ("内分泌系统", "消化系统", 0.10, "±"),
            ("神经系统", "肌肉骨骼系统", 0.16, "+"),
            ("消化系统", "免疫系统", -0.18, "±"),
        ]
    else:
        systems = ["神经系统", "内分泌系统", "免疫系统", "睡眠系统", "心血管系统", "消化系统"]
        influence = [86, 78, 74, 69, 64, 58]
        xlabels = ["急性加重", "睡眠受损", "自主神经症状", "反复就诊", "生活质量下降"]
        ylabels = ["低风险", "中风险", "高风险"]
        matrix = np.array([[0.31, 0.29, 0.26, 0.30, 0.25], [0.48, 0.44, 0.41, 0.45, 0.39], [0.21, 0.27, 0.33, 0.25, 0.36]])
        pos = {
            "神经系统": (0.20, 0.78),
            "内分泌系统": (0.50, 0.78),
            "免疫系统": (0.80, 0.78),
            "睡眠系统": (0.20, 0.28),
            "心血管系统": (0.50, 0.28),
            "消化系统": (0.80, 0.28),
        }
        edges = [
            ("神经系统", "睡眠系统", -0.10, "±"),
            ("内分泌系统", "睡眠系统", -0.08, "±"),
            ("内分泌系统", "心血管系统", 0.08, "±"),
            ("神经系统", "心血管系统", 0.06, "±"),
            ("免疫系统", "消化系统", -0.08, "+"),
        ]

    systems = [str(x) for x in spec_list("fig_2_1", "categories", systems)]
    influence = spec_num_list("fig_2_1", "values", [float(x) for x in influence])
    if len(systems) != len(influence):
        systems, influence = systems[: min(len(systems), len(influence))], influence[: min(len(systems), len(influence))]
    fig, ax = plt.subplots(figsize=(7.8, 4.4))
    ax.bar(systems, influence, color="#2C7A7B")
    ax.set_ylabel("关联强度（0-100）")
    set_main_title(ax, "fig_2_1", f"图表2-1：{DISEASE_NAME}与相关系统关联强度")
    save_figure(FIG_DIR / "fig_2_1.png", fig)
    add_fig_meta("fig_2_1", f"图表2-1：{DISEASE_NAME}与相关系统关联强度", "柱状图", "公开资料整理", "系统关联评分", "N/A", "2.1", "系统关联", "数据来源：公开资料整理")

    xlabels = [str(x) for x in spec_list("fig_2_2", "x_labels", xlabels)]
    ylabels = [str(x) for x in spec_list("fig_2_2", "y_labels", ylabels)]
    matrix_default = matrix.copy()
    matrix_v = fig_spec("fig_2_2").get("matrix")
    if isinstance(matrix_v, list) and matrix_v:
        try:
            m_arr = np.array(matrix_v, dtype=float)
            if m_arr.ndim == 2 and m_arr.shape[0] == len(ylabels) and m_arr.shape[1] == len(xlabels):
                matrix = m_arr
            else:
                matrix = matrix_default
        except Exception:
            matrix = matrix_default
    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    im = ax.imshow(matrix, cmap="YlOrRd", aspect="auto")
    ax.set_xticks(np.arange(len(xlabels)))
    ax.set_xticklabels(xlabels, rotation=20)
    ax.set_yticks(np.arange(len(ylabels)))
    ax.set_yticklabels(ylabels)
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            ax.text(j, i, f"{matrix[i, j]*100:.0f}%", ha="center", va="center", fontsize=8)
    set_main_title(ax, "fig_2_2", "图表2-2：常见并发风险矩阵")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    save_figure(FIG_DIR / "fig_2_2.png", fig)
    add_fig_meta("fig_2_2", "图表2-2：常见并发风险矩阵", "热力图", "公开资料整理", "并发风险评分", "N/A", "2.2", "风险矩阵", "数据来源：公开资料整理")

    fig23_title = fig23_expected_caption()
    if fig23_layout_mode() == "causal_chain":
        fig, ax = plt.subplots(figsize=(11.0, 5.8))
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")

        ax.text(0.5, 0.95, "病因与调节层", ha="center", va="center", fontsize=10, color="#2D3748")
        ax.text(0.5, 0.73, "核心病理层", ha="center", va="center", fontsize=10, color="#2D3748")
        ax.text(0.5, 0.47, "机制传导层", ha="center", va="center", fontsize=10, color="#2D3748")
        ax.text(0.5, 0.18, "临床后果层", ha="center", va="center", fontsize=10, color="#2D3748")

        nodes = {
            "Hp感染": draw_box_node(ax, 0.12, 0.85, "Hp感染", width=0.16, height=0.085, fc="#EBF8FF", ec="#2B6CB0", fontsize=9.2),
            "化学/药物刺激": draw_box_node(ax, 0.30, 0.85, "化学/药物刺激", width=0.18, height=0.085, fc="#EBF8FF", ec="#2B6CB0", fontsize=9.2),
            "胃肠微生态失衡": draw_box_node(ax, 0.48, 0.85, "胃肠微生态失衡", width=0.18, height=0.085, fc="#EBF8FF", ec="#2B6CB0", fontsize=9.2),
            "神经系统": draw_box_node(ax, 0.66, 0.85, "神经系统\n(脑-肠轴)", width=0.16, height=0.095, fc="#EBF8FF", ec="#2B6CB0", fontsize=9.0),
            "内分泌系统": draw_box_node(ax, 0.84, 0.85, "内分泌系统", width=0.16, height=0.085, fc="#EBF8FF", ec="#2B6CB0", fontsize=9.2),
            "免疫系统": draw_box_node(ax, 0.28, 0.66, "免疫系统", width=0.14, height=0.085, fc="#E6FFFA", ec="#2C7A7B", fontsize=9.2),
            "核心": draw_box_node(ax, 0.50, 0.66, "慢性胃炎\n(胃黏膜炎症)", width=0.22, height=0.10, fc="#FEEBC8", ec="#C05621", fontsize=9.8),
            "黏膜屏障": draw_box_node(ax, 0.72, 0.66, "胃黏膜屏障受损", width=0.18, height=0.085, fc="#FEEBC8", ec="#C05621", fontsize=9.1),
            "酸动力": draw_box_node(ax, 0.34, 0.42, "胃酸/胃动力异常", width=0.18, height=0.085, fc="#FFF5F5", ec="#C53030", fontsize=9.1),
            "炎症持续": draw_box_node(ax, 0.50, 0.42, "慢性炎症持续", width=0.16, height=0.085, fc="#FFF5F5", ec="#C53030", fontsize=9.1),
            "修复障碍": draw_box_node(ax, 0.66, 0.42, "黏膜修复障碍", width=0.16, height=0.085, fc="#FFF5F5", ec="#C53030", fontsize=9.1),
            "症状负担": draw_box_node(ax, 0.28, 0.12, "症状负担", width=0.14, height=0.08, fc="#FAF5FF", ec="#6B46C1", fontsize=9.0),
            "营养/贫血风险": draw_box_node(ax, 0.50, 0.12, "营养吸收受限\n/贫血风险", width=0.20, height=0.09, fc="#FAF5FF", ec="#6B46C1", fontsize=8.8),
            "癌前病变风险": draw_box_node(ax, 0.72, 0.12, "癌前病变风险", width=0.16, height=0.08, fc="#FAF5FF", ec="#6B46C1", fontsize=9.0),
        }

        # 病因与调节 -> 核心病理
        draw_poly_arrow(ax, [nodes["Hp感染"]["south"], nodes["核心"]["north"]], color="#2B6CB0")
        draw_poly_arrow(ax, [nodes["化学/药物刺激"]["south"], nodes["核心"]["north"]], color="#2B6CB0")
        draw_poly_arrow(ax, [nodes["胃肠微生态失衡"]["south"], nodes["核心"]["north"]], color="#2B6CB0")
        draw_poly_arrow(ax, [nodes["免疫系统"]["east"], nodes["核心"]["west"]], color="#2C7A7B")
        draw_poly_arrow(ax, [nodes["核心"]["west"], nodes["免疫系统"]["east"]], color="#2C7A7B", dashed=True)
        draw_poly_arrow(ax, [nodes["神经系统"]["south"], (0.60, 0.74), nodes["核心"]["north"]], color="#2B6CB0")
        draw_poly_arrow(ax, [nodes["核心"]["north"], (0.58, 0.79), nodes["神经系统"]["south"]], color="#2B6CB0", dashed=True)
        draw_poly_arrow(ax, [nodes["内分泌系统"]["south"], (0.78, 0.74), nodes["核心"]["north"]], color="#2B6CB0")
        draw_poly_arrow(ax, [nodes["核心"]["north"], (0.74, 0.79), nodes["内分泌系统"]["south"]], color="#2B6CB0", dashed=True)

        # 核心病理 -> 机制传导
        draw_poly_arrow(ax, [nodes["核心"]["east"], nodes["黏膜屏障"]["west"]], color="#C05621")
        draw_poly_arrow(ax, [nodes["核心"]["south"], nodes["炎症持续"]["north"]], color="#C05621")
        draw_poly_arrow(ax, [nodes["核心"]["south"], (0.40, 0.50), nodes["酸动力"]["north"]], color="#C05621")
        draw_poly_arrow(ax, [nodes["黏膜屏障"]["south"], nodes["修复障碍"]["north"]], color="#C05621")

        # 机制传导 -> 临床后果
        draw_poly_arrow(ax, [nodes["酸动力"]["south"], nodes["症状负担"]["north"]], color="#C53030")
        draw_poly_arrow(ax, [nodes["炎症持续"]["south"], nodes["营养/贫血风险"]["north"]], color="#C53030")
        draw_poly_arrow(ax, [nodes["修复障碍"]["south"], nodes["癌前病变风险"]["north"]], color="#C53030")

        # 关系符号标注（+ 促进, ± 双向调节）
        ax.text(0.20, 0.77, "+", fontsize=10, color="#2B6CB0", fontweight="bold")
        ax.text(0.36, 0.77, "+", fontsize=10, color="#2B6CB0", fontweight="bold")
        ax.text(0.52, 0.77, "+", fontsize=10, color="#2B6CB0", fontweight="bold")
        ax.text(0.39, 0.66, "±", fontsize=10, color="#2C7A7B", fontweight="bold")
        ax.text(0.60, 0.77, "±", fontsize=10, color="#2B6CB0", fontweight="bold")
        ax.text(0.76, 0.77, "±", fontsize=10, color="#2B6CB0", fontweight="bold")
        ax.text(0.60, 0.66, "+", fontsize=10, color="#C05621", fontweight="bold")
        ax.text(0.50, 0.53, "+", fontsize=10, color="#C05621", fontweight="bold")
        ax.text(0.43, 0.32, "+", fontsize=10, color="#C53030", fontweight="bold")
        ax.text(0.56, 0.32, "+", fontsize=10, color="#C53030", fontweight="bold")
        ax.text(0.72, 0.32, "+", fontsize=10, color="#C53030", fontweight="bold")

        set_main_title(ax, "fig_2_3", fig23_title, fontsize=12, fontweight="bold")
    elif fig23_layout_mode() == "layered_path":
        fig, ax = plt.subplots(figsize=(11.0, 5.0))
        draw_fig23_layered_path(ax, fig23_layered_path_config())
        fig._codex_tight_rect = (0, 0, 1, 0.96)
        set_main_title(ax, "fig_2_3", fig23_title, fontsize=12, fontweight="bold")
    elif fig23_layout_mode() == "dual_panel":
        fig, axes = plt.subplots(1, 2, figsize=(10.8, 4.8))
        ax_l, ax_r = axes
        for ax in axes:
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
            ax.axis("off")

        panel_cfg = fig23_dual_panel_config()
        left_cfg = panel_cfg.get("left", {}) if isinstance(panel_cfg, dict) else {}
        right_cfg = panel_cfg.get("right", {}) if isinstance(panel_cfg, dict) else {}
        if isinstance(left_cfg, dict):
            draw_configured_network_panel(ax_l, left_cfg, default_edge_color="#2B6CB0")
        if isinstance(right_cfg, dict):
            draw_configured_network_panel(ax_r, right_cfg, default_edge_color="#2B6CB0")
        set_main_suptitle(fig, "fig_2_3", fig23_title, fontsize=12, fontweight="bold", y=0.98)
    else:
        fig, ax = plt.subplots(figsize=(8.2, 4.8))
        ax.axis("off")
        fig23_cfg = fig_spec("fig_2_3")
        disallow_terms = fig23_disallow_nodes()
        extra_disallow = fig23_cfg.get("disallow_nodes")
        if isinstance(extra_disallow, list):
            disallow_terms.extend([str(x).strip() for x in extra_disallow if str(x).strip()])
        pos_cfg = fig23_cfg.get("pos")
        if isinstance(pos_cfg, dict):
            new_pos: Dict[str, Tuple[float, float]] = {}
            for k, v in pos_cfg.items():
                if not isinstance(v, (list, tuple)) or len(v) < 2:
                    continue
                try:
                    new_pos[str(k)] = (float(v[0]), float(v[1]))
                except Exception:
                    continue
            if new_pos:
                pos = new_pos
        edges_cfg = fig23_cfg.get("edges")
        if isinstance(edges_cfg, list):
            new_edges = []
            for e in edges_cfg:
                if not isinstance(e, (list, tuple)) or len(e) < 3:
                    continue
                try:
                    src = str(e[0]).strip()
                    dst = str(e[1]).strip()
                    rad = float(e[2])
                    sign = str(e[3]).strip() if len(e) >= 4 else ""
                    if src and dst:
                        new_edges.append((src, dst, rad, sign))
                except Exception:
                    continue
            if new_edges:
                edges = new_edges
        systems = [x for x in systems if x not in disallow_terms]
        for term in disallow_terms:
            if term in pos:
                pos.pop(term, None)
        filtered_edges = []
        for e in edges:
            s, t, rad = e[0], e[1], e[2]
            sign = e[3] if len(e) >= 4 else ""
            if s in disallow_terms or t in disallow_terms:
                continue
            if s not in pos or t not in pos:
                continue
            filtered_edges.append((s, t, rad, sign))

        core_drawn = False
        if fig23_require_core_node():
            core_pos = (0.50, 0.53)
            core_drawn = True
            core_label = fig23_core_label()
            cfg_core_label = fig23_cfg.get("core_label")
            if cfg_core_label is not None:
                core_label = render_disease_template(str(cfg_core_label), default=core_label)
            ax.text(
                core_pos[0],
                core_pos[1],
                core_label,
                ha="center",
                va="center",
                fontsize=9.2,
                bbox=dict(boxstyle="round,pad=0.36", fc="#FEEBC8", ec="#C05621", lw=1.2),
            )
            top_links = [n for n in fig23_top_to_core_nodes() if n in pos]
            cfg_top = fig23_cfg.get("top_to_core")
            if isinstance(cfg_top, list) and cfg_top:
                top_links = [str(x).strip() for x in cfg_top if str(x).strip() in pos]
            for n in top_links:
                x1, y1 = pos[n]
                ax.annotate(
                    "",
                    xy=(core_pos[0], core_pos[1] + 0.07),
                    xytext=(x1, y1 - 0.06),
                    arrowprops=dict(arrowstyle="->", lw=1.1, color="#2B6CB0", connectionstyle="arc3,rad=0.0"),
                )
                ax.text((x1 + core_pos[0]) / 2, (y1 + core_pos[1]) / 2 + 0.03, "±", fontsize=8.5, color="#2B6CB0", fontweight="bold")
            bottom_links = [n for n in fig23_core_to_bottom_nodes() if n in pos]
            cfg_bottom = fig23_cfg.get("core_to_bottom")
            if isinstance(cfg_bottom, list) and cfg_bottom:
                bottom_links = [str(x).strip() for x in cfg_bottom if str(x).strip() in pos]
            for n in bottom_links:
                x2, y2 = pos[n]
                ax.annotate(
                    "",
                    xy=(x2, y2 + 0.06),
                    xytext=(core_pos[0], core_pos[1] - 0.07),
                    arrowprops=dict(arrowstyle="->", lw=1.1, color="#2B6CB0", connectionstyle="arc3,rad=0.0"),
                )
                ax.text((x2 + core_pos[0]) / 2, (y2 + core_pos[1]) / 2 - 0.03, "+", fontsize=8.5, color="#2B6CB0", fontweight="bold")

        for name, (x, y) in pos.items():
            ax.text(x, y, name, ha="center", va="center", fontsize=10, bbox=dict(boxstyle="round,pad=0.35", fc="#EDF2F7", ec="#2D3748", lw=1.1))
        for s, t, rad, sign in filtered_edges:
            x1, y1 = pos[s]
            x2, y2 = pos[t]
            ax.annotate("", xy=(x2, y2 + 0.06), xytext=(x1, y1 - 0.06), arrowprops=dict(arrowstyle="->", lw=1.2, color="#2B6CB0", connectionstyle=f"arc3,rad={rad}"))
            if sign:
                ax.text((x1 + x2) / 2, (y1 + y2) / 2, sign, fontsize=8.5, color="#2B6CB0", fontweight="bold")
        if core_drawn:
            ax.text(0.86, 0.92, "注：+ 促进；± 双向调节", fontsize=8.2, color="#4A5568", ha="right")
        set_main_title(ax, "fig_2_3", fig23_title, fontsize=12, fontweight="bold")

    save_figure(FIG_DIR / "fig_2_3.png", fig)
    fig_2_3_caption = fig23_title
    add_fig_meta("fig_2_3", fig_2_3_caption, "关系图", "公开资料整理", "系统交互机制", "N/A", "2.2", "关系网络", "数据来源：公开资料整理")

    # Chapter 3
    if is_cervical_profile():
        diag_flow_title = f"{DISEASE_NAME}临床诊疗流程"
        diag_flow_nodes = ["首诊分层", "神经体征检查", "影像评估", "保守治疗", "复评分流", "介入/手术"]
        schemes = ["药物+康复", "理疗+牵引", "介入治疗", "手术减压融合"]
        high = [31, 22, 18, 19]
        mid = [37, 40, 45, 44]
        low = [32, 38, 37, 37]
        pie_title = f"{DISEASE_NAME}常用治疗组合偏好"
        pie_labels = ["药物治疗", "物理治疗", "康复训练", "介入治疗", "手术治疗"]
        pie_vals = [32.0, 26.0, 21.0, 9.0, 12.0]
    elif is_gastritis_profile():
        diag_flow_title = f"{DISEASE_NAME}临床诊疗流程"
        diag_flow_nodes = ["症状与红旗征筛查", "Hp检测", "内镜+病理分级", "病因分层", "治疗执行", "复查随访"]
        schemes = ["抑酸+黏膜保护", "Hp根除方案", "促动力/消化酶", "中西联合方案"]
        high = [35, 41, 22, 18]
        mid = [42, 39, 46, 47]
        low = [23, 20, 32, 35]
        pie_title = f"{DISEASE_NAME}治疗路径构成"
        pie_labels = ["抑酸治疗", "黏膜保护", "Hp根除", "促动力/消化酶", "中医协同"]
        pie_vals = [29.0, 24.0, 22.0, 15.0, 10.0]
    elif is_respiratory_profile():
        diag_flow_title = f"{DISEASE_NAME}临床诊疗流程"
        diag_flow_nodes = ["首诊分诊", "病因评估", "风险分层", "治疗启动", "复评调整", "出院/随访"]
        schemes = ["祛痰+止咳", "抗炎+祛痰", "支气管舒张+止咳", "中西联合方案"]
        high = [28, 24, 19, 15]
        mid = [34, 36, 41, 44]
        low = [38, 40, 40, 41]
        pie_title = f"{DISEASE_NAME}常用剂型偏好"
        pie_labels = ["口服液", "颗粒剂", "糖浆", "片剂", "雾化支持"]
        pie_vals = [34.0, 26.0, 21.0, 9.0, 10.0]
    else:
        diag_flow_title = f"{DISEASE_NAME}临床诊疗流程"
        diag_flow_nodes = ["首诊评估", "病因识别", "风险分层", "治疗启动", "复评校准", "长期管理"]
        schemes = ["单药治疗", "联合治疗", "分层管理", "多学科协同"]
        high = [29, 25, 22, 18]
        mid = [39, 41, 44, 46]
        low = [32, 34, 34, 36]
        pie_title = f"{DISEASE_NAME}治疗方案偏好"
        pie_labels = ["药物治疗", "非药物干预", "联合方案", "长期管理", "其他支持"]
        pie_vals = [33.0, 21.0, 24.0, 14.0, 8.0]

    diag_flow_title = spec_text("fig_3_1", "flow_title", diag_flow_title)
    diag_flow_nodes = [str(x) for x in spec_list("fig_3_1", "nodes", diag_flow_nodes)]
    title_3_1 = compose_figure_title("3-1", diag_flow_title)
    draw_simple_flow(FIG_DIR / "fig_3_1.png", title_3_1, diag_flow_nodes, direction="lr", color="#276749", figsize=(10.2, 3.2))
    rendered_title_rows.append({"fig_id": "fig_3_1", "rendered_title": normalize_disease_text(title_3_1)})
    add_fig_meta("fig_3_1", title_3_1, "流程图", "指南整理", "诊疗路径", "N/A", "3.1", "诊断路径", "数据来源：临床指南与公开资料整理")

    schemes = [str(x) for x in spec_list("fig_3_2", "categories", schemes)]
    high = spec_num_list("fig_3_2", "high", [float(x) for x in high])
    mid = spec_num_list("fig_3_2", "mid", [float(x) for x in mid])
    low = spec_num_list("fig_3_2", "low", [float(x) for x in low])
    n3 = min(len(schemes), len(high), len(mid), len(low))
    schemes, high, mid, low = schemes[:n3], high[:n3], mid[:n3], low[:n3]
    fig, ax = plt.subplots(figsize=(8.2, 5.0))
    ax.bar(schemes, high, label="高证据", color="#2B6CB0")
    ax.bar(schemes, mid, bottom=high, label="中证据", color="#63B3ED")
    ax.bar(schemes, low, bottom=np.array(high) + np.array(mid), label="低证据", color="#BEE3F8")
    ax.set_ylabel("占比（%）")
    ax.set_ylim(0, 105)
    set_main_title(ax, "fig_3_2", "图表3-2：主要治疗方案证据等级结构", pad=12)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.12), ncol=3, frameon=False, borderaxespad=0.0)
    fig._codex_tight_rect = (0, 0.08, 1, 1)
    save_figure(FIG_DIR / "fig_3_2.png", fig)
    add_fig_meta("fig_3_2", "图表3-2：主要治疗方案证据等级结构", "堆叠柱状图", "公开资料整理", "方案证据分层", "N/A", "3.2", "治疗评估", "数据来源：公开资料整理")

    pie_title = spec_text("fig_3_3", "pie_title", pie_title)
    pie_labels = [str(x) for x in spec_list("fig_3_3", "labels", pie_labels)]
    pie_vals = spec_num_list("fig_3_3", "values", [float(x) for x in pie_vals])
    n33 = min(len(pie_labels), len(pie_vals))
    pie_labels, pie_vals = pie_labels[:n33], pie_vals[:n33]
    title_3_3 = compose_figure_title("3-3", pie_title)
    draw_pie_with_leaders(FIG_DIR / "fig_3_3.png", title_3_3, pie_labels, pie_vals, ["#3182CE", "#63B3ED", "#90CDF4", "#A0AEC0", "#2F855A"], figsize=(7.5, 4.4))
    rendered_title_rows.append({"fig_id": "fig_3_3", "rendered_title": normalize_disease_text(title_3_3)})
    add_fig_meta("fig_3_3", title_3_3, "饼图", "公开资料整理", "剂型偏好", "N/A", "3.3", "剂型结构", "数据来源：公开资料整理")

    # Chapter 4 (Excel-driven)
    q = ch4.quarterly.copy()
    q["label"] = q["quarter"]
    fig, ax = plt.subplots(figsize=(8.6, 4.7))
    ax.plot(q["label"], q["hospital"], label="医院端", color="#2B6CB0", lw=2)
    ax.plot(q["label"], q["drugstore"], label="药店端", color="#DD6B20", lw=2)
    ax.plot(q["label"], q["online"], label="线上端", color="#2F855A", lw=2)
    step = max(1, len(q) // 8)
    ax.set_xticks(range(0, len(q), step))
    ax.set_xticklabels(q["label"].iloc[::step], rotation=35, ha="right")
    ax.set_ylabel("销售额（万元）")
    set_main_title(ax, "fig_4_1", f"图表4-1：三端季度销售额趋势（{q['quarter'].iloc[0]}-{q['quarter'].iloc[-1]}）")
    ax.legend()
    save_figure(FIG_DIR / "fig_4_1.png", fig)
    add_fig_meta("fig_4_1", f"图表4-1：三端季度销售额趋势（{q['quarter'].iloc[0]}-{q['quarter'].iloc[-1]}）", "折线图", "米内网", "quarterly_channel", "quarterly_channel", "4.1", "第4章数据专线", "数据来源：米内网")

    latest = ch4.latest_share
    title_4_2 = spec_text("fig_4_2", "title", f"图表4-2：{ch4.latest_quarter}三端销售结构占比")
    draw_pie_with_leaders(FIG_DIR / "fig_4_2.png", title_4_2, latest["channel"].tolist(), latest["share_pct"].tolist(), ["#2B6CB0", "#DD6B20", "#2F855A"], figsize=(7.5, 4.4))
    rendered_title_rows.append({"fig_id": "fig_4_2", "rendered_title": normalize_disease_text(title_4_2)})
    add_fig_meta("fig_4_2", f"图表4-2：{ch4.latest_quarter}三端销售结构占比", "饼图", "米内网", "latest_share", "latest_share", "4.1", "第4章数据专线", "数据来源：米内网")

    annual = ch4.annual
    fig, ax = plt.subplots(figsize=(8.0, 4.7))
    x = np.arange(len(annual))
    w = 0.25
    ax.bar(x - w, annual["hospital"], width=w, label="医院端", color="#2B6CB0")
    ax.bar(x, annual["drugstore"], width=w, label="药店端", color="#DD6B20")
    ax.bar(x + w, annual["online"], width=w, label="线上端", color="#2F855A")
    ax.set_xticks(x)
    ax.set_xticklabels(annual["year"].astype(int).astype(str))
    ax.set_ylabel("销售额（万元）")
    set_main_title(ax, "fig_4_3", "图表4-3：年度三端销售额对比")
    ax.legend()
    save_figure(FIG_DIR / "fig_4_3.png", fig)
    add_fig_meta("fig_4_3", "图表4-3：年度三端销售额对比", "分组柱状图", "米内网", "annual_channel", "annual_channel", "4.2", "第4章数据专线", "数据来源：米内网")

    fig, ax = plt.subplots(figsize=(7.8, 4.2))
    yoy = ch4.yoy_latest
    ax.bar(yoy["channel"], yoy["yoy_pct"], color=["#2B6CB0", "#DD6B20", "#2F855A"])
    ax.axhline(0, color="#4A5568", lw=1)
    ax.set_ylabel("同比增速（%）")
    set_main_title(ax, "fig_4_4", f"图表4-4：{ch4.latest_quarter}三端同比增速")
    for i, v in enumerate(yoy["yoy_pct"]):
        if pd.notna(v):
            ax.text(i, v + (0.8 if v >= 0 else -1.2), f"{v:.1f}%", ha="center", va="bottom" if v >= 0 else "top", fontsize=9)
    save_figure(FIG_DIR / "fig_4_4.png", fig)
    add_fig_meta("fig_4_4", f"图表4-4：{ch4.latest_quarter}三端同比增速", "柱状图", "米内网", "latest_yoy", "latest_yoy", "4.2", "第4章数据专线", "数据来源：米内网")

    def top10_bar(df: pd.DataFrame, title: str, path: Path, fig_id: str):
        fig, ax = plt.subplots(figsize=(8.2, 4.8))
        d = df.copy()
        if d.empty:
            set_main_title(ax, fig_id, title)
            ax.axis("off")
            ax.text(0.5, 0.5, "源表未提供该渠道TOP数据", ha="center", va="center", fontsize=12, color="#4A5568")
            save_figure(path, fig)
            return
        d = d.sort_values("sales", ascending=True)
        ax.barh(d["name"], d["sales"], color="#3182CE")
        ax.set_xlabel("销售额（万元）")
        set_main_title(ax, fig_id, title)
        save_figure(path, fig)

    top10_bar(ch4.top_hospital, f"图表4-5：医院端TOP10通用名（{ch4.latest_quarter}）", FIG_DIR / "fig_4_5.png", "fig_4_5")
    add_fig_meta("fig_4_5", f"图表4-5：医院端TOP10通用名（{ch4.latest_quarter}）", "横向柱状图", "米内网", "top10_hospital", "top10_hospital", "4.3", "第4章数据专线", "数据来源：米内网")
    top10_bar(ch4.top_drugstore, f"图表4-6：药店端TOP10通用名（{ch4.latest_quarter}）", FIG_DIR / "fig_4_6.png", "fig_4_6")
    add_fig_meta("fig_4_6", f"图表4-6：药店端TOP10通用名（{ch4.latest_quarter}）", "横向柱状图", "米内网", "top10_drugstore", "top10_drugstore", "4.4", "第4章数据专线", "数据来源：米内网")
    top10_bar(ch4.top_online, f"图表4-7：线上端TOP10通用名（{ch4.latest_quarter}）", FIG_DIR / "fig_4_7.png", "fig_4_7")
    add_fig_meta("fig_4_7", f"图表4-7：线上端TOP10通用名（{ch4.latest_quarter}）", "横向柱状图", "米内网", "top10_online", "top10_online", "4.4", "第4章数据专线", "数据来源：米内网")

    fig, ax = plt.subplots(figsize=(8.2, 4.5))
    cr5 = ch4.cr5_latest
    ax.bar(cr5["channel"], cr5["cr5_pct"], color=["#2B6CB0", "#DD6B20", "#2F855A"])
    ax.set_ylabel("CR5（%）")
    valid_cr5 = cr5["cr5_pct"].dropna()
    upper = float(valid_cr5.max()) * 1.25 if not valid_cr5.empty else 10.0
    ax.set_ylim(0, max(10, upper))
    set_main_title(ax, "fig_4_8", f"图表4-8：{ch4.latest_quarter}三端市场集中度（CR5）")
    for i, v in enumerate(cr5["cr5_pct"]):
        if pd.notna(v):
            ax.text(i, v + 0.8, f"{v:.1f}%", ha="center", fontsize=9)
        else:
            ax.text(i, 0.8, "N/A", ha="center", fontsize=9, color="#718096")
    save_figure(FIG_DIR / "fig_4_8.png", fig)
    add_fig_meta("fig_4_8", f"图表4-8：{ch4.latest_quarter}三端市场集中度（CR5）", "柱状图", "米内网", "cr5_latest", "cr5_latest", "4.3", "第4章数据专线", "数据来源：米内网")

    # Chapter 5
    if is_cervical_profile():
        age_groups = ["18-39岁", "40-49岁", "50-59岁", "60岁及以上"]
        male = [22, 28, 24, 16]
        female = [18, 26, 25, 17]
        factors = ["疼痛缓解证据", "神经功能改善", "起效速度", "复发控制", "康复便利性", "支付可及性"]
        vals = [89, 86, 78, 73, 68, 61]
        journey_title = f"{DISEASE_NAME}全周期管理流程"
        journey_nodes = ["症状识别", "首诊分层", "保守治疗", "功能康复", "复评决策", "长期管理"]
        labels = ["工作姿势负荷", "居家训练依从", "疼痛波动", "复诊可及性", "康复资源不足", "心理压力"]
        impact = [76, 71, 67, 59, 56, 53]
    elif is_gastritis_profile():
        age_groups = ["18-39岁", "40-49岁", "50-59岁", "60岁及以上"]
        male = [18, 24, 23, 17]
        female = [16, 22, 24, 20]
        factors = ["症状控制证据", "Hp根除证据", "病理风险管理", "价格可及性", "复查便利性", "安全性"]
        vals = [88, 91, 83, 66, 72, 79]
        journey_title = f"{DISEASE_NAME}长期管理流程"
        journey_nodes = ["症状出现", "初诊评估", "病因分层", "治疗执行", "复查评估", "长期随访"]
        labels = ["疗程长度", "复查依从", "不良反应担忧", "信息不一致", "生活方式执行", "复诊可及性"]
        impact = [71, 69, 58, 56, 61, 54]
    elif is_respiratory_profile():
        age_groups = ["0-2岁", "3-5岁", "6-9岁", "10-14岁"]
        male = [18, 29, 24, 13]
        female = [15, 26, 21, 14]
        factors = ["安全性证据", "起效速度", "口感依从性", "价格可及性", "家长教育支持", "复诊衔接"]
        vals = [91, 84, 78, 66, 62, 58]
        journey_title = f"{DISEASE_NAME}全周期管理流程"
        journey_nodes = ["首发症状", "初诊评估", "治疗启动", "功能恢复", "复发预防", "长期随访"]
        labels = ["家长时间投入", "用药频次复杂", "口感接受度", "信息不一致", "疗程提醒不足", "随访中断"]
        impact = [72, 69, 63, 58, 54, 49]
    else:
        age_groups = ["18-39岁", "40-49岁", "50-59岁", "60岁及以上"]
        male = [21, 27, 23, 16]
        female = [19, 25, 24, 18]
        factors = ["安全性证据", "疗效持续性", "起效速度", "支付可及性", "患者教育支持", "复诊衔接"]
        vals = [88, 82, 74, 68, 63, 60]
        journey_title = f"{DISEASE_NAME}全周期管理流程"
        journey_nodes = ["症状识别", "初诊评估", "治疗启动", "复评调整", "复发预防", "长期随访"]
        labels = ["治疗复杂度", "执行负担", "信息不一致", "疗程提醒不足", "随访可及性", "生活方式约束"]
        impact = [70, 67, 61, 57, 55, 50]

    age_groups = [str(x) for x in spec_list("fig_5_1", "x_labels", age_groups)]
    male = spec_num_list("fig_5_1", "male", [float(x) for x in male])
    female = spec_num_list("fig_5_1", "female", [float(x) for x in female])
    n51 = min(len(age_groups), len(male), len(female))
    age_groups, male, female = age_groups[:n51], male[:n51], female[:n51]

    factors = [str(x) for x in spec_list("fig_5_2", "categories", factors)]
    vals = spec_num_list("fig_5_2", "values", [float(x) for x in vals])
    n52 = min(len(factors), len(vals))
    factors, vals = factors[:n52], vals[:n52]

    labels = [str(x) for x in spec_list("fig_5_4", "x_labels", labels)]
    impact = spec_num_list("fig_5_4", "values", [float(x) for x in impact])
    n54 = min(len(labels), len(impact))
    labels, impact = labels[:n54], impact[:n54]

    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    x = np.arange(len(age_groups))
    ax.bar(x, male, width=0.35, label="男", color="#3182CE")
    ax.bar(x + 0.35, female, width=0.35, label="女", color="#90CDF4")
    ax.set_xticks(x + 0.175)
    ax.set_xticklabels(age_groups)
    ax.set_ylabel("占比（%）")
    set_main_title(ax, "fig_5_1", f"图表5-1：{DISEASE_NAME}就诊人群年龄与性别结构")
    ax.legend()
    save_figure(FIG_DIR / "fig_5_1.png", fig)
    add_fig_meta("fig_5_1", f"图表5-1：{DISEASE_NAME}就诊人群年龄与性别结构", "分组柱状图", "公开资料整理", "患者画像", "N/A", "5.1", "患者画像", "数据来源：公开资料整理")

    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    ax.barh(factors[::-1], vals[::-1], color="#D69E2E")
    ax.set_xlim(0, 100)
    ax.set_xlabel("影响强度指数")
    set_main_title(ax, "fig_5_2", "图表5-2：医生处方偏好与决策要素")
    save_figure(FIG_DIR / "fig_5_2.png", fig)
    add_fig_meta("fig_5_2", "图表5-2：医生处方偏好与决策要素", "横向柱状图", "公开资料整理", "医生偏好", "N/A", "5.2", "医生偏好", "数据来源：公开资料整理")

    journey_title = spec_text("fig_5_3", "flow_title", journey_title)
    journey_nodes = [str(x) for x in spec_list("fig_5_3", "nodes", journey_nodes)]
    title_5_3 = compose_figure_title("5-3", journey_title)
    draw_simple_flow(FIG_DIR / "fig_5_3.png", title_5_3, journey_nodes, direction="lr", color="#2F855A", figsize=(10.6, 3.1))
    rendered_title_rows.append({"fig_id": "fig_5_3", "rendered_title": normalize_disease_text(title_5_3)})
    add_fig_meta("fig_5_3", title_5_3, "流程图", "公开资料整理", "患者旅程", "N/A", "5.3", "全周期管理", "数据来源：公开资料整理")

    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    xpos = np.arange(len(labels))
    ax.bar(xpos, impact, color="#805AD5")
    ax.set_ylabel("对依从性的影响（指数）")
    ax.set_xticks(xpos)
    ax.set_xticklabels(labels, rotation=20, ha="right")
    set_main_title(ax, "fig_5_4", "图表5-4：依从性影响因素分解")
    save_figure(FIG_DIR / "fig_5_4.png", fig)
    add_fig_meta("fig_5_4", "图表5-4：依从性影响因素分解", "柱状图", "公开资料整理", "依从性因素", "N/A", "5.2", "依从因素", "数据来源：公开资料整理")

    # Chapter 6
    if is_cervical_profile():
        policy_title = f"{DISEASE_NAME}相关政策时间线"
        policy_events = [("2018", "骨科分级诊疗"), ("2020", "康复服务规范"), ("2022", "集采与支付协同"), ("2024", "慢病管理强化"), ("2025", "互联网复诊规范")]
    elif is_gastritis_profile():
        policy_title = f"{DISEASE_NAME}相关政策时间线"
        policy_events = [("2017", "慢性胃炎共识更新"), ("2020", "消化内镜质量规范"), ("2022", "Hp处理共识更新"), ("2024", "医保目录调整"), ("2025", "互联网复诊规范化")]
    elif is_respiratory_profile():
        policy_title = f"{DISEASE_NAME}相关政策时间线"
        policy_events = [("2019", "儿童用药规范化"), ("2021", "药品审评优化"), ("2023", "质量监管强化"), ("2024", "支付政策调整"), ("2025", "分级诊疗协同")]
    else:
        policy_title = f"{DISEASE_NAME}相关政策时间线"
        policy_events = [("2018", "诊疗规范更新"), ("2020", "支付政策优化"), ("2022", "质量与监管协同"), ("2024", "慢病管理强化"), ("2025", "分级诊疗与数字化协同")]
    policy_title = spec_text("fig_6_1", "timeline_title", policy_title)
    policy_events_cfg = fig_spec("fig_6_1").get("events")
    if isinstance(policy_events_cfg, list) and policy_events_cfg:
        events: List[Tuple[str, str]] = []
        for it in policy_events_cfg:
            if isinstance(it, (list, tuple)) and len(it) >= 2:
                events.append((str(it[0]), render_disease_template(str(it[1]))))
            elif isinstance(it, dict) and ("year" in it) and ("event" in it):
                events.append((str(it.get("year", "")), render_disease_template(str(it.get("event", "")))))
        if events:
            policy_events = events
    title_6_1 = compose_figure_title("6-1", policy_title)
    draw_policy_timeline(FIG_DIR / "fig_6_1.png", title_6_1, policy_events, figsize=(8.2, 3.0))
    rendered_title_rows.append({"fig_id": "fig_6_1", "rendered_title": normalize_disease_text(title_6_1)})
    add_fig_meta("fig_6_1", title_6_1, "时间轴", "政府公开文件", "政策环境", "N/A", "6.1", "政策环境", "数据来源：国家卫健委、国家药监局、国家医保局")

    fig, ax = plt.subplots(figsize=(10.4, 3.8))
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    boxes = [("审评审批", (0.10, 0.62)), ("质量控制", (0.33, 0.62)), ("医保支付", (0.56, 0.62)), ("终端执行", (0.79, 0.62)), ("用药结构优化", (0.45, 0.28))]
    boxes_cfg = fig_spec("fig_6_2").get("boxes")
    if isinstance(boxes_cfg, list) and boxes_cfg:
        new_boxes: List[Tuple[str, Tuple[float, float]]] = []
        for b in boxes_cfg:
            if isinstance(b, (list, tuple)) and len(b) >= 3:
                try:
                    new_boxes.append((render_disease_template(str(b[0])), (float(b[1]), float(b[2]))))
                except Exception:
                    continue
            elif isinstance(b, dict):
                try:
                    new_boxes.append(
                        (
                            render_disease_template(str(b.get("label", ""))),
                            (float(b.get("x", 0.5)), float(b.get("y", 0.5))),
                        )
                    )
                except Exception:
                    continue
        if new_boxes:
            boxes = new_boxes
    box_anchors: List[Dict[str, Tuple[float, float]]] = []
    custom_layout = isinstance(boxes_cfg, list) and bool(boxes_cfg)
    if (not custom_layout) and len(boxes) >= 5:
        top_labels, top_widths, top_gap = layout_horizontal_flow_nodes(
            [boxes[0][0], boxes[1][0], boxes[2][0], boxes[3][0]],
            left_margin=0.035,
            right_margin=0.035,
            min_gap=0.085,
            max_gap=0.16,
            min_width=0.11,
            max_width=0.18,
            base=0.066,
            char_step=0.0066,
        )
        available = 1.0 - 0.035 - 0.035
        total_width = sum(top_widths) + top_gap * 3
        cursor = 0.035 + max(0.0, (available - total_width) / 2.0)
        for label, width in zip(top_labels, top_widths):
            x = cursor + width / 2.0
            box_anchors.append(
                draw_box_node(
                    ax,
                    x,
                    0.64,
                    label,
                    width=width,
                    height=flow_box_height(label, 0.115),
                    fc="#F7FAFC",
                    ec="#2D3748",
                    lw=1.2,
                    fontsize=9.8 if "\n" in label else 10.0,
                )
            )
            cursor += width + top_gap
        bottom_label = wrap_flow_label(boxes[4][0])
        bottom_width = max(0.18, estimate_flow_box_width(bottom_label, min_width=0.14, max_width=0.20, base=0.068, char_step=0.0068))
        box_anchors.append(
            draw_box_node(
                ax,
                0.50,
                0.26,
                bottom_label,
                width=bottom_width,
                height=flow_box_height(bottom_label, 0.118),
                fc="#F7FAFC",
                ec="#2D3748",
                lw=1.2,
                fontsize=9.8 if "\n" in bottom_label else 10.0,
            )
        )
        arrows = [
            [box_anchors[0]["east"], box_anchors[1]["west"]],
            [box_anchors[1]["east"], box_anchors[2]["west"]],
            [box_anchors[2]["east"], box_anchors[3]["west"]],
            [
                box_anchors[2]["south"],
                (box_anchors[2]["south"][0], box_anchors[2]["south"][1] - 0.08),
                (box_anchors[4]["north"][0] + 0.08, box_anchors[4]["north"][1] + 0.10),
                box_anchors[4]["north"],
            ],
        ]
    else:
        for idx, (text, (x, y)) in enumerate(boxes):
            label = wrap_flow_label(text)
            width = estimate_flow_box_width(label, min_width=0.11, max_width=0.20, base=0.068, char_step=0.0072)
            if idx == len(boxes) - 1:
                width = max(width, 0.16)
            box_anchors.append(
                draw_box_node(
                    ax,
                    x,
                    y,
                    label,
                    width=width,
                    height=flow_box_height(label, 0.115),
                    fc="#F7FAFC",
                    ec="#2D3748",
                    lw=1.2,
                    fontsize=9.8 if "\n" in label else 10.0,
                )
            )
        arrows = [
            [box_anchors[0]["east"], box_anchors[1]["west"]],
            [box_anchors[1]["east"], box_anchors[2]["west"]],
            [box_anchors[2]["east"], box_anchors[3]["west"]],
            [box_anchors[2]["south"], box_anchors[4]["north"]],
        ]
    arrows_cfg = fig_spec("fig_6_2").get("arrows")
    if isinstance(arrows_cfg, list) and arrows_cfg:
        new_arrows: List[List[Tuple[float, float]]] = []
        for a in arrows_cfg:
            if isinstance(a, (list, tuple)) and len(a) >= 4:
                try:
                    if len(a) > 4 and len(a) % 2 == 0:
                        pts: List[Tuple[float, float]] = []
                        for i in range(0, len(a), 2):
                            pts.append((float(a[i]), float(a[i + 1])))
                        new_arrows.append(pts)
                    else:
                        new_arrows.append([(float(a[0]), float(a[1])), (float(a[2]), float(a[3]))])
                except Exception:
                    continue
            elif isinstance(a, dict):
                try:
                    via_points: List[Tuple[float, float]] = []
                    raw_via = a.get("via", [])
                    if isinstance(raw_via, list):
                        for item in raw_via:
                            if isinstance(item, (list, tuple)) and len(item) >= 2:
                                via_points.append((float(item[0]), float(item[1])))
                    new_arrows.append(
                        [(float(a.get("x1", 0.0)), float(a.get("y1", 0.0)))] + via_points + [(float(a.get("x2", 0.0)), float(a.get("y2", 0.0)))]
                    )
                except Exception:
                    continue
        if new_arrows:
            arrows = new_arrows
    for points in arrows:
        draw_poly_arrow(ax, points, color="#2B6CB0", lw=1.5, shrink_start_pts=10.0, shrink_end_pts=10.0)
    set_main_title(ax, "fig_6_2", "图表6-2：医保支付与监管联动对用药结构的影响路径", fontsize=12, fontweight="bold")
    save_figure(FIG_DIR / "fig_6_2.png", fig)
    add_fig_meta("fig_6_2", "图表6-2：医保支付与监管联动对用药结构的影响路径", "路径图", "政策公开文件整理", "监管趋势", "N/A", "6.2", "监管趋势", "数据来源：国家医保局、国家药监局")

    # Chapter 7
    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    years = np.array([2025, 2026, 2027, 2028, 2029, 2030])
    base = np.array([100, 108, 116, 124, 132, 141])
    optimistic = np.array([100, 111, 121, 132, 144, 157])
    conservative = np.array([100, 105, 110, 116, 121, 127])
    y_cfg = spec_num_list("fig_7_1", "years", years.astype(float).tolist())
    b_cfg = spec_num_list("fig_7_1", "base", base.astype(float).tolist())
    o_cfg = spec_num_list("fig_7_1", "optimistic", optimistic.astype(float).tolist())
    c_cfg = spec_num_list("fig_7_1", "conservative", conservative.astype(float).tolist())
    n71 = min(len(y_cfg), len(b_cfg), len(o_cfg), len(c_cfg))
    years = np.array(y_cfg[:n71])
    base = np.array(b_cfg[:n71])
    optimistic = np.array(o_cfg[:n71])
    conservative = np.array(c_cfg[:n71])
    ax.plot(years, base, label="基准情景", color="#2B6CB0", lw=2)
    ax.plot(years, optimistic, label="乐观情景", color="#2F855A", lw=2)
    ax.plot(years, conservative, label="审慎情景", color="#DD6B20", lw=2)
    ax.set_ylabel("市场规模指数（2025=100）")
    set_main_title(ax, "fig_7_1", f"图表7-1：{DISEASE_NAME}市场规模预测（2026-2030）")
    ax.legend()
    save_figure(FIG_DIR / "fig_7_1.png", fig)
    add_fig_meta("fig_7_1", f"图表7-1：{DISEASE_NAME}市场规模预测（2026-2030）", "折线图", "米内网+趋势测算", "预测模型", "annual_channel", "7.1", "市场预测", "数据来源：米内网与趋势测算")

    fig, ax = plt.subplots(figsize=(7.8, 4.5))
    if is_cervical_profile():
        measures = ["证据升级", "康复网络", "数字随访", "依从管理", "术式优化", "准入协同"]
    elif is_gastritis_profile():
        measures = ["证据升级", "Hp管理", "病理随访", "准入协同", "渠道组合", "患者教育"]
    elif is_respiratory_profile():
        measures = ["证据升级", "渠道协同", "家长教育", "依从管理", "数字化运营", "准入优化"]
    else:
        measures = ["证据升级", "分层管理", "渠道协同", "依从管理", "数字化随访", "准入优化"]
    measures = [str(x) for x in spec_list("fig_7_2", "labels", measures)]
    x = spec_num_list("fig_7_2", "x", [82, 74, 62, 77, 69, 58])
    y = spec_num_list("fig_7_2", "y", [86, 79, 72, 75, 84, 65])
    sizes = spec_num_list("fig_7_2", "sizes", [420, 360, 300, 380, 340, 280])
    n72 = min(len(measures), len(x), len(y), len(sizes))
    measures, x, y, sizes = measures[:n72], x[:n72], y[:n72], sizes[:n72]
    ax.scatter(x, y, s=sizes, c=["#2B6CB0", "#3182CE", "#63B3ED", "#2F855A", "#38A169", "#D69E2E"], alpha=0.75)
    for i, m in enumerate(measures):
        ax.text(x[i] + 0.8, y[i] + 0.6, m, fontsize=8)
    ax.set_xlabel("战略价值")
    ax.set_ylabel("落地可行性")
    set_main_title(ax, "fig_7_2", "图表7-2：战略举措优先级矩阵")
    xlim_cfg = spec_num_list("fig_7_2", "xlim", [50, 90])
    ylim_cfg = spec_num_list("fig_7_2", "ylim", [60, 90])
    if len(xlim_cfg) >= 2:
        ax.set_xlim(xlim_cfg[0], xlim_cfg[1])
    else:
        ax.set_xlim(50, 90)
    if len(ylim_cfg) >= 2:
        ax.set_ylim(ylim_cfg[0], ylim_cfg[1])
    else:
        ax.set_ylim(60, 90)
    save_figure(FIG_DIR / "fig_7_2.png", fig)
    add_fig_meta("fig_7_2", "图表7-2：战略举措优先级矩阵", "气泡图", "项目分析", "战略举措评分", "N/A", "7.2", "战略建议", "数据来源：项目分析整理")

    title_rows: List[Dict[str, str]] = []
    for row in rendered_title_rows:
        title = normalize_disease_text(str(row.get("rendered_title", "")).strip())
        title_rows.append(
            {
                "fig_id": str(row.get("fig_id", "")).strip(),
                "rendered_title": title,
                "has_serial_prefix": "1" if FIG_TITLE_SERIAL_RE.match(title) else "0",
            }
        )
    write_csv(
        OUT_ROOT / "figure_title_registry.csv",
        title_rows,
        ["fig_id", "rendered_title", "has_serial_prefix"],
    )

    return fig_rows


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    paragraph._p = paragraph._element = None


def set_para_text(paragraph, text: str, bold: bool = False, center: bool = False):
    paragraph.text = normalize_disease_text(text)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        for r in paragraph.runs:
            r.bold = bold


def apply_reference_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)


def insert_block_with_figures(doc: Document, subtitle: str, content: str, fig_ids: List[str], fig_meta: Dict[str, Dict[str, str]]) -> None:
    h2 = doc.add_paragraph(style="二级目录")
    set_para_text(h2, subtitle)
    paragraphs = split_paragraphs(content)
    if not paragraphs:
        return

    insert_positions: List[int] = []
    if len(fig_ids) == 1:
        insert_positions = [max(1, len(paragraphs) // 2)]
    elif len(fig_ids) >= 2:
        p1 = max(1, len(paragraphs) // 3)
        p2 = max(p1 + 1, len(paragraphs) * 2 // 3)
        insert_positions = [p1, p2]
        if len(fig_ids) > 2:
            insert_positions.extend([len(paragraphs)] * (len(fig_ids) - 2))

    fig_ptr = 0
    for i, para_text in enumerate(paragraphs, start=1):
        p = doc.add_paragraph(style="数据报告正文")
        set_para_text(p, para_text)
        while fig_ptr < len(fig_ids) and i == insert_positions[min(fig_ptr, len(insert_positions) - 1)]:
            fid = fig_ids[fig_ptr]
            meta = fig_meta[fid]
            cap = doc.add_paragraph(style="数据报告正文")
            set_para_text(cap, meta["caption"], bold=True, center=True)
            doc.add_picture(str(FIG_DIR / meta["输出文件名"]), width=Inches(5.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            src = doc.add_paragraph(style="数据报告正文")
            set_para_text(src, meta["source_line"], center=True)
            fig_ptr += 1

    while fig_ptr < len(fig_ids):
        fid = fig_ids[fig_ptr]
        meta = fig_meta[fid]
        cap = doc.add_paragraph(style="数据报告正文")
        set_para_text(cap, meta["caption"], bold=True, center=True)
        doc.add_picture(str(FIG_DIR / meta["输出文件名"]), width=Inches(5.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        src = doc.add_paragraph(style="数据报告正文")
        set_para_text(src, meta["source_line"], center=True)
        fig_ptr += 1


def assemble_docx(specs: List[BlockSpec], block_text: Dict[str, str], summary: str, refs_text: str, fig_rows: List[Dict[str, str]]) -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Missing template: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))
    first_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        style_name = p.style.name if p.style is not None else ""
        # Use body heading anchor only; avoid matching TOC lines that also start with "第一章".
        if style_name == "一级目录" and t.startswith("第一章"):
            first_idx = i
            break
    if first_idx is None:
        first_idx = len(doc.paragraphs)
    for i in range(len(doc.paragraphs) - 1, first_idx - 1, -1):
        remove_paragraph(doc.paragraphs[i])

    fig_map: Dict[str, Dict[str, str]] = {r["fig_id"]: r for r in fig_rows}
    specs_by_ch: Dict[int, List[BlockSpec]] = {}
    for s in specs:
        specs_by_ch.setdefault(s.chapter, []).append(s)

    for ch in range(1, 8):
        h1 = doc.add_paragraph(style="一级目录")
        set_para_text(h1, chapter_title(ch))
        for s in specs_by_ch[ch]:
            figs = [x for x in s.fig_ids.split("|") if x]
            insert_block_with_figures(doc, s.subtitle, block_text[s.block_id], figs, fig_map)

    h_sum = doc.add_paragraph(style="一级目录")
    set_para_text(h_sum, "总结")
    for ptxt in split_paragraphs(summary):
        p = doc.add_paragraph(style="数据报告正文")
        set_para_text(p, ptxt)

    h_ref = doc.add_paragraph(style="一级目录")
    set_para_text(h_ref, "参考文献")
    for line in refs_text.splitlines():
        if line.strip():
            p = doc.add_paragraph(style="数据报告正文")
            set_para_text(p, normalize_reference_line(line))
            apply_reference_paragraph_format(p)

    for section in doc.sections:
        # Only touch default header to avoid creating extra empty first/even headers.
        hdr = section.header
        for p in hdr.paragraphs:
            txt = p.text
            txt = txt.replace("XXX", DISEASE_NAME)
            txt = txt.replace("《XXX疾病市场分析报告》", REPORT_TITLE)
            txt = txt.replace("《XXX市场分析报告》", REPORT_TITLE)
            txt = txt.replace("XXX疾病市场分析报告", f"{DISEASE_NAME}市场分析报告")
            txt = txt.replace("XXX市场分析报告", f"{DISEASE_NAME}市场分析报告")
            if txt != p.text:
                set_para_text(p, txt)

    # Ensure body sections without explicit header inherit the previous section header.
    for i in range(1, len(doc.sections)):
        prev_hdr_txt = "".join(x.text.strip() for x in doc.sections[i - 1].header.paragraphs)
        cur_hdr_txt = "".join(x.text.strip() for x in doc.sections[i].header.paragraphs)
        if (not cur_hdr_txt) and prev_hdr_txt:
            doc.sections[i].header.is_linked_to_previous = True

    FINAL_DOCX.parent.mkdir(parents=True, exist_ok=True)
    backup_if_exists(FINAL_DOCX)
    doc.save(str(FINAL_DOCX))


def post_process_docx_xml(docx_path: Path) -> None:
    def footer_rel_map(rels_xml: str) -> Dict[str, str]:
        rels: Dict[str, str] = {}
        pattern = re.compile(
            r'<Relationship[^>]*\bId="([^"]+)"[^>]*\bType="([^"]+)"[^>]*\bTarget="([^"]+)"[^>]*/>'
        )
        for rid, rel_type, target in pattern.findall(rels_xml):
            if not rel_type.endswith("/footer"):
                continue
            p = target.replace("\\", "/")
            if p.startswith("/"):
                p = p.lstrip("/")
            elif not p.startswith("word/"):
                p = f"word/{p}"
            rels[rid] = p
        return rels

    def footer_has_page_field(footer_xml: str) -> bool:
        return bool(
            re.search(r'<w:instrText[^>]*>\s*[^<]*\bPAGE\b', footer_xml)
            or re.search(r'<w:fldSimple[^>]*w:instr="[^"]*\bPAGE\b', footer_xml)
        )

    def build_plain_page_footer_xml() -> str:
        return (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:ftr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
            "<w:p>"
            "<w:pPr><w:jc w:val=\"center\"/></w:pPr>"
            "<w:r><w:rPr><w:rFonts w:hint=\"default\" w:ascii=\"Times New Roman\" w:hAnsi=\"Times New Roman\" w:cs=\"Times New Roman\"/>"
            "<w:b/><w:bCs/><w:sz w:val=\"21\"/><w:szCs w:val=\"32\"/></w:rPr><w:fldChar w:fldCharType=\"begin\"/></w:r>"
            "<w:r><w:rPr><w:rFonts w:hint=\"default\" w:ascii=\"Times New Roman\" w:hAnsi=\"Times New Roman\" w:cs=\"Times New Roman\"/>"
            "<w:b/><w:bCs/><w:sz w:val=\"21\"/><w:szCs w:val=\"32\"/></w:rPr><w:instrText xml:space=\"preserve\"> PAGE \\* MERGEFORMAT </w:instrText></w:r>"
            "<w:r><w:rPr><w:rFonts w:hint=\"default\" w:ascii=\"Times New Roman\" w:hAnsi=\"Times New Roman\" w:cs=\"Times New Roman\"/>"
            "<w:b/><w:bCs/><w:sz w:val=\"21\"/><w:szCs w:val=\"32\"/></w:rPr><w:fldChar w:fldCharType=\"separate\"/></w:r>"
            "<w:r><w:rPr><w:rFonts w:hint=\"default\" w:ascii=\"Times New Roman\" w:hAnsi=\"Times New Roman\" w:cs=\"Times New Roman\"/>"
            "<w:b/><w:bCs/><w:sz w:val=\"21\"/><w:szCs w:val=\"32\"/></w:rPr><w:t>1</w:t></w:r>"
            "<w:r><w:rPr><w:rFonts w:hint=\"default\" w:ascii=\"Times New Roman\" w:hAnsi=\"Times New Roman\" w:cs=\"Times New Roman\"/>"
            "<w:b/><w:bCs/><w:sz w:val=\"21\"/><w:szCs w:val=\"32\"/></w:rPr><w:fldChar w:fldCharType=\"end\"/></w:r>"
            "</w:p>"
            "</w:ftr>"
        )

    def pick_main_footer_rid(doc_xml: str, rels_xml: str, xml_parts: Dict[str, str]) -> str | None:
        tags = re.findall(r"<w:footerReference[^>]*/>", doc_xml)
        if not tags:
            return None

        refs: List[Tuple[str, str]] = []
        for tag in tags:
            rid_m = re.search(r'r:id="([^"]+)"', tag)
            if not rid_m:
                continue
            type_m = re.search(r'w:type="([^"]+)"', tag)
            refs.append((rid_m.group(1), type_m.group(1) if type_m else "default"))
        if not refs:
            return None

        ordered: List[str] = []
        ordered_default: List[str] = []
        for rid, ftype in refs:
            if rid not in ordered:
                ordered.append(rid)
            if ftype == "default" and rid not in ordered_default:
                ordered_default.append(rid)
        # Prefer later sections' default footer first (body footer is usually in later sections).
        candidates_in_order = ordered_default + [rid for rid in ordered if rid not in ordered_default]
        candidates = list(reversed(candidates_in_order))

        rel_map = footer_rel_map(rels_xml)

        def footer_xml(rid: str) -> str:
            target = rel_map.get(rid, "")
            return xml_parts.get(target, "")

        for rid in candidates:
            fxml = footer_xml(rid)
            if fxml and footer_has_page_field(fxml) and "<w:txbxContent" not in fxml:
                return rid
        for rid in candidates:
            fxml = footer_xml(rid)
            if fxml and footer_has_page_field(fxml):
                return rid
        return candidates[0] if candidates else None

    tmp = docx_path.with_suffix(".tmp.docx")
    replacements = {
        "<<<在此填写疾病名>>>": DISEASE_NAME,
        "<<<疾病名>>>": DISEASE_NAME,
        "<<<在此填写医学主题>>>": DISEASE_NAME,
        "<<<医学主题>>>": DISEASE_NAME,
        "《XXX疾病市场分析报告》": REPORT_TITLE,
        "《XXX市场分析报告》": REPORT_TITLE,
        "XXX疾病市场分析报告": f"{DISEASE_NAME}市场分析报告",
        "XXX市场分析报告": f"{DISEASE_NAME}市场分析报告",
        "XXX": DISEASE_NAME,
        "AAA": "",
    }
    for token in LEGACY_DISEASE_TOKENS:
        replacements[token] = DISEASE_NAME
    with zipfile.ZipFile(docx_path, "r") as zin:
        infos = zin.infolist()
        binary_parts = {item.filename: zin.read(item.filename) for item in infos}

    xml_parts: Dict[str, str] = {}
    for name, data in binary_parts.items():
        if not name.endswith(".xml"):
            continue
        txt = data.decode("utf-8")
        for old, new in replacements.items():
            txt = txt.replace(old, new)
        xml_parts[name] = txt

    rels_xml = ""
    if "word/_rels/document.xml.rels" in binary_parts:
        rels_xml = binary_parts["word/_rels/document.xml.rels"].decode("utf-8")

    doc_xml = xml_parts.get("word/document.xml")
    if doc_xml is not None:
        # Normalize section pagination strategy:
        # - use a single PAGE-capable default footer for all sections;
        # - clear section-level page-number reset (w:start) to keep continuous numbering.
        main_footer_rid = pick_main_footer_rid(doc_xml, rels_xml, xml_parts)
        if main_footer_rid:
            rel_map = footer_rel_map(rels_xml)
            main_footer_target = rel_map.get(main_footer_rid, "")

            def normalize_footer_reference(tag: str) -> str:
                if 'r:id="' in tag:
                    return re.sub(r'r:id="[^"]+"', f'r:id="{main_footer_rid}"', tag, count=1)
                return tag.replace("/>", f' r:id="{main_footer_rid}"/>')

            doc_xml = re.sub(r"<w:footerReference[^>]*/>", lambda m: normalize_footer_reference(m.group(0)), doc_xml)
            if main_footer_target:
                xml_parts[main_footer_target] = build_plain_page_footer_xml()
        doc_xml = re.sub(r'(<w:pgNumType\b[^>]*?)\s+w:start="[^"]+"', r"\1", doc_xml)
        xml_parts["word/document.xml"] = doc_xml

    settings_xml = xml_parts.get("word/settings.xml")
    if settings_xml is not None:
        if "<w:updateFields" not in settings_xml:
            settings_xml = settings_xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
        else:
            settings_xml = re.sub(r"<w:updateFields[^>]*/>", '<w:updateFields w:val="true"/>', settings_xml)
        xml_parts["word/settings.xml"] = settings_xml

    with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in infos:
            if item.filename in xml_parts:
                data = xml_parts[item.filename].encode("utf-8")
            else:
                data = binary_parts[item.filename]
            zout.writestr(item, data)
    backup_if_exists(docx_path)
    shutil.move(tmp, docx_path)


def extract_docx_text_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(docx_path, "r") as zf:
        all_xml = []
        for name in zf.namelist():
            if name.endswith(".xml") and name.startswith("word/"):
                all_xml.append(zf.read(name).decode("utf-8", errors="ignore"))
        return "\n".join(all_xml)


def normalize_sentence(s: str) -> str:
    s = re.sub(r"\[\d+\]", "", s)
    s = re.sub(r"\s+", "", s)
    s = s.replace("，", ",").replace("。", ".").replace("；", ";")
    return s.strip(" .;,:")


def sentence_repeat_stats(text: str) -> Tuple[int, List[Tuple[str, int]]]:
    chunks = re.split(r"[。！？!?\n]+", text)
    counter: Dict[str, int] = {}
    samples: Dict[str, str] = {}
    for raw in chunks:
        s = raw.strip()
        if len(s) < 18:
            continue
        if s.startswith("图表") or s.startswith("数据来源："):
            continue
        norm = normalize_sentence(s)
        if len(norm) < 16:
            continue
        counter[norm] = counter.get(norm, 0) + 1
        samples.setdefault(norm, s)
    top = sorted([(samples[k], v) for k, v in counter.items()], key=lambda x: x[1], reverse=True)
    max_dup = top[0][1] if top else 0
    return max_dup, top[:5]


def paragraph_has_anchor(text: str) -> bool:
    patterns = [
        r"20\d{2}",
        r"\d+(?:\.\d+)?%",
        r"\d+(?:\.\d+)?(?:万元|万|元|例|岁|天|小时|周|月|年|季度)",
        r"\[\d+\]",
        r"(评分|发生率|恢复时间|稳定率|完成率|复评|复查|红旗征|禁忌|适应证|说明书|内镜|病理|幽门螺杆菌|CR5|CAGR|YTD)",
    ]
    return any(re.search(p, text) for p in patterns)


MEDICAL_PATTERNS = {
    "red_flag": r"(红旗征|警示症状|进行性消瘦|贫血|黑便|呕血|呼吸困难|发绀|低氧|意识改变|脱水)",
    "contra": r"(年龄限制|禁忌|适应证|说明书|妊娠|肝肾功能)",
    "review": r"(复评|复查|48小时|72小时|3-5天|1周|3个月|6个月|1年|随访|内镜|病理)",
    "evidence": r"(指南|共识|证据等级|推荐)",
    "safety": r"(不良反应|安全监测|警示|药物相互作用|风险)",
}


STALE_PHRASES = [
    "并非单点波动",
    "可追踪、可复算、可解释",
    "先做试点，再做跨区域放大",
    "只有把结论沉淀到标准流程",
    "前端分诊质量会改变后续处方稳定性",
    "关键耦合关系在于",
]

MANAGEMENT_DRIFT_PHRASES = [
    "跨部门关键动作拆解为可追踪任务并设置责任人",
    "谁来做、何时做、如何复盘",
    "同一话术和同一KPI",
    "结论-动作-复盘-迭代",
]


def collect_metric_logic_issues(text: str) -> List[str]:
    issues: List[str] = []
    if re.search(r"无效换药率[^。；\n]{0,40}(>=|>|不低于|高于)", text):
        issues.append("无效换药率阈值方向疑似错误")
    if re.search(r"复购间隔中位数[^。；\n]{0,40}%", text):
        issues.append("复购间隔中位数与百分比单位冲突")
    if re.search(r"(中位数)[^。；\n]{0,20}(>=|<=)[^。；\n]{0,10}%", text):
        issues.append("中位数指标疑似错误使用百分比阈值")
    return issues


def is_nonspecific_reference_url(url: str) -> bool:
    try:
        p = urlparse(url.strip())
    except Exception:
        return False
    domain = (p.netloc or "").lower()
    path = (p.path or "").strip("/")
    if not domain:
        return False
    strict_domains = {"www.who.int", "who.int", "www.nhc.gov.cn", "nhc.gov.cn", "guide.medlive.cn"}
    return (domain in strict_domains) and (path == "")


def parse_evidence_pool(path: Path) -> Tuple[List[Dict[str, object]], List[str]]:
    """
    Parse 00_evidence.txt lines:
      证据ID|标题|机构/作者|年份|要点|可追溯来源
    """
    if not path.exists():
        return [], [f"missing:{path.name}"]

    lines = [x.rstrip("\n") for x in path.read_text(encoding="utf-8").splitlines() if x.strip()]
    if not lines:
        return [], [f"empty:{path.name}"]

    rows: List[Dict[str, object]] = []
    errs: List[str] = []
    start_idx = 1 if lines[0].startswith("证据ID|") else 0

    for i, line in enumerate(lines[start_idx:], start=start_idx + 1):
        parts = [p.strip() for p in line.split("|", 5)]
        if len(parts) < 6:
            errs.append(f"line{i}:field_count<{6}")
            continue
        eid, title, org, year, keypoint, source = parts
        m = re.match(r"^E(\d{2})$", eid)
        if not m:
            errs.append(f"line{i}:bad_evidence_id={eid}")
            continue
        num = int(m.group(1))
        year_ok = bool(re.search(r"(19|20)\d{2}", year))
        source_ok = ("http" in source) or (EXCEL_PATH.name in source) or (normalize_disease_text(EXCEL_PATH.name) in source)
        rows.append(
            {
                "line": i,
                "evidence_id": eid,
                "num": num,
                "title": title,
                "org": org,
                "year": year,
                "year_ok": year_ok,
                "keypoint": keypoint,
                "source": source,
                "source_ok": source_ok,
            }
        )
    return rows, errs


def parse_reference_list(path: Path) -> Tuple[List[Dict[str, object]], List[str]]:
    """
    Parse refs.txt lines:
      [1] 机构. 标题[类型]. 年份. URL/文件
    """
    if not path.exists():
        return [], [f"missing:{path.name}"]

    lines = [x.strip() for x in path.read_text(encoding="utf-8").splitlines() if x.strip()]
    if not lines:
        return [], [f"empty:{path.name}"]

    rows: List[Dict[str, object]] = []
    errs: List[str] = []
    for i, line in enumerate(lines, start=1):
        m = re.match(r"^\[(\d+)\]\s+(.+)$", line)
        if not m:
            errs.append(f"line{i}:bad_ref_prefix")
            continue
        num = int(m.group(1))
        body = m.group(2).strip()
        year_ok = bool(re.search(r"(19|20)\d{2}", body))
        url_m = re.search(r"https?://[^\s]+", body)
        has_url_or_file = bool(url_m) or (EXCEL_PATH.name in body) or (normalize_disease_text(EXCEL_PATH.name) in body)
        rows.append(
            {
                "line": i,
                "num": num,
                "raw": line,
                "body": body,
                "year_ok": year_ok,
                "has_url_or_file": has_url_or_file,
                "url": url_m.group(0) if url_m else "",
            }
        )
    return rows, errs


def collect_reference_chain_metrics(specs: List[BlockSpec], block_text: Dict[str, str], summary_text: str) -> Dict[str, object]:
    evidence_rows, evidence_parse_errors = parse_evidence_pool(OUT_ROOT / "00_evidence.txt")
    ref_rows, ref_parse_errors = parse_reference_list(OUT_ROOT / "refs.txt")

    evidence_nums = [int(r["num"]) for r in evidence_rows]
    ref_nums = [int(r["num"]) for r in ref_rows]
    evidence_set = set(evidence_nums)
    ref_set = set(ref_nums)

    evidence_id_dup = sorted([x for x in evidence_set if evidence_nums.count(x) > 1])
    ref_id_dup = sorted([x for x in ref_set if ref_nums.count(x) > 1])

    evidence_seq_ok = evidence_nums == list(range(1, len(evidence_nums) + 1))
    ref_seq_ok = ref_nums == list(range(1, len(ref_nums) + 1))
    evidence_ref_gap = sorted(evidence_set.symmetric_difference(ref_set))

    all_text = "\n".join([block_text[s.block_id] for s in specs]) + "\n" + summary_text
    cited_nums_all = [int(x) for x in re.findall(r"\[(\d+)\]", all_text)]
    cited_set = set(cited_nums_all)
    dangling_cites = sorted([x for x in cited_set if x not in ref_set])
    uncited_refs = sorted([x for x in ref_set if x not in cited_set])
    citation_coverage = (len(ref_set - set(uncited_refs)) / len(ref_set)) if ref_set else 0.0

    evidence_bad_year = sorted([int(r["num"]) for r in evidence_rows if not bool(r.get("year_ok"))])
    evidence_bad_source = sorted([int(r["num"]) for r in evidence_rows if not bool(r.get("source_ok"))])
    ref_bad_year = sorted([int(r["num"]) for r in ref_rows if not bool(r.get("year_ok"))])
    ref_bad_source = sorted([int(r["num"]) for r in ref_rows if not bool(r.get("has_url_or_file"))])

    return {
        "evidence_count": len(evidence_rows),
        "ref_count": len(ref_rows),
        "evidence_parse_errors": evidence_parse_errors,
        "ref_parse_errors": ref_parse_errors,
        "evidence_id_dup": evidence_id_dup,
        "ref_id_dup": ref_id_dup,
        "evidence_seq_ok": evidence_seq_ok,
        "ref_seq_ok": ref_seq_ok,
        "evidence_ref_gap": evidence_ref_gap,
        "cited_count": len(cited_set),
        "dangling_cites": dangling_cites,
        "uncited_refs": uncited_refs,
        "citation_coverage": citation_coverage,
        "evidence_bad_year": evidence_bad_year,
        "evidence_bad_source": evidence_bad_source,
        "ref_bad_year": ref_bad_year,
        "ref_bad_source": ref_bad_source,
    }


def collect_text_quality_metrics(specs: List[BlockSpec], block_text: Dict[str, str]) -> Dict[str, object]:
    all_text = "\n".join([block_text[s.block_id] for s in specs])
    stale_counts = {ph: all_text.count(ph) for ph in STALE_PHRASES}
    drift_counts = {ph: all_text.count(ph) for ph in MANAGEMENT_DRIFT_PHRASES}
    max_sentence_dup, top_sentence_dups = sentence_repeat_stats(all_text)
    metric_logic_issues = collect_metric_logic_issues(all_text)

    dup_prefix_hits = 0
    for s in specs:
        paras = split_paragraphs(block_text[s.block_id])
        if paras and paras[0].startswith(s.subtitle):
            dup_prefix_hits += 1

    chapter_stats: List[Tuple[int, int, int, float, int, int]] = []
    low_anchor_chapters: List[int] = []
    chapter_dup_fails: List[int] = []
    chapter_no_cites: List[int] = []
    chapter_chars: Dict[int, int] = {}
    chapter_len_fails: List[str] = []
    for ch in range(1, 8):
        ch_blocks = [s for s in specs if s.chapter == ch]
        ch_text = "\n".join([block_text[s.block_id] for s in ch_blocks])
        ch_paras = []
        for s in ch_blocks:
            ch_paras.extend(split_paragraphs(block_text[s.block_id]))
        ch_chars = len(re.sub(r"\s+", "", ch_text))
        chapter_chars[ch] = ch_chars
        anchored = sum(1 for p in ch_paras if paragraph_has_anchor(p))
        anchor_cov = anchored / len(ch_paras) if ch_paras else 0.0
        ch_dup, _ = sentence_repeat_stats(ch_text)
        cite_cnt = len(re.findall(r"\[\d+\]", ch_text))
        chapter_stats.append((ch, ch_chars, len(ch_paras), anchor_cov, ch_dup, cite_cnt))
        if anchor_cov < 0.7:
            low_anchor_chapters.append(ch)
        if ch_dup >= 4:
            chapter_dup_fails.append(ch)
        if cite_cnt == 0:
            chapter_no_cites.append(ch)
        min_chars = CHAPTER_MIN_CHARS[ch]
        shortfall = chapter_char_shortfall(ch, ch_chars)
        if shortfall > CHAPTER_CHAR_TOLERANCE:
            chapter_len_fails.append(f"{ch}({ch_chars}<{min_chars}; shortfall={shortfall}>tol={CHAPTER_CHAR_TOLERANCE})")

    anchor_cov_by_block: List[Tuple[str, float]] = []
    for s in specs:
        paras = split_paragraphs(block_text[s.block_id])
        if not paras:
            anchor_cov_by_block.append((s.block_id, 0.0))
            continue
        anchored = sum(1 for p in paras if paragraph_has_anchor(p))
        anchor_cov_by_block.append((s.block_id, anchored / len(paras)))
    low_anchor_blocks = [f"{bid}={cov*100:.0f}%" for bid, cov in anchor_cov_by_block if cov < 0.7]

    medical_density_failed = []
    for s in specs:
        if s.chapter > 3:
            continue
        txt = block_text[s.block_id]
        hit_types = sum(1 for p in MEDICAL_PATTERNS.values() if re.search(p, txt))
        if hit_types < 2:
            medical_density_failed.append(s.block_id)

    cagr_logic_ok = True
    ch41 = block_text.get("4.1", "")
    ch44 = block_text.get("4.4", "")
    cagr_match = re.search(r"年复合增速约为(-?\d+(?:\.\d+)?)%", ch41)
    if cagr_match:
        cagr_val = float(cagr_match.group(1))
        if cagr_val < 0 and re.search(r"(总体扩张|保持增长|持续增长)", ch44):
            cagr_logic_ok = False
        if cagr_val > 0 and re.search(r"(总体收缩|持续收缩)", ch44):
            cagr_logic_ok = False

    cr5_logic_ok = True
    ch43 = block_text.get("4.3", "")
    cr5_match = re.search(
        r"医院端CR5为(-?\d+(?:\.\d+)?)%，药店端CR5为(-?\d+(?:\.\d+)?)%，线上端CR5为(-?\d+(?:\.\d+)?)%。按当前口径，(医院端|药店端|线上端)集中度最高",
        ch43,
    )
    if cr5_match:
        cr5_vals = {
            "医院端": float(cr5_match.group(1)),
            "药店端": float(cr5_match.group(2)),
            "线上端": float(cr5_match.group(3)),
        }
        stated_top = cr5_match.group(4)
        expected_top = sorted(cr5_vals.items(), key=lambda x: x[1], reverse=True)[0][0]
        cr5_logic_ok = stated_top == expected_top

    return {
        "all_text": all_text,
        "stale_counts": stale_counts,
        "drift_counts": drift_counts,
        "max_sentence_dup": max_sentence_dup,
        "top_sentence_dups": top_sentence_dups,
        "metric_logic_issues": metric_logic_issues,
        "dup_prefix_hits": dup_prefix_hits,
        "chapter_stats": chapter_stats,
        "low_anchor_chapters": low_anchor_chapters,
        "chapter_dup_fails": chapter_dup_fails,
        "chapter_no_cites": chapter_no_cites,
        "chapter_chars": chapter_chars,
        "chapter_len_fails": chapter_len_fails,
        "anchor_cov_by_block": anchor_cov_by_block,
        "low_anchor_blocks": low_anchor_blocks,
        "medical_density_failed": medical_density_failed,
        "cagr_logic_ok": cagr_logic_ok,
        "cr5_logic_ok": cr5_logic_ok,
    }


def build_codex_rewrite_prompt(specs: List[BlockSpec], metrics: Dict[str, object], summary_text: str) -> str:
    chapter_chars: Dict[int, int] = metrics["chapter_chars"]  # type: ignore[assignment]
    chapter_len_fails: List[str] = metrics["chapter_len_fails"]  # type: ignore[assignment]
    low_anchor_blocks: List[str] = metrics["low_anchor_blocks"]  # type: ignore[assignment]
    medical_density_failed: List[str] = metrics["medical_density_failed"]  # type: ignore[assignment]
    total_chars = sum(chapter_chars.values()) + len(re.sub(r"\s+", "", summary_text))
    total_gap = max(0, 30000 - total_chars)
    summary_chars = len(re.sub(r"\s+", "", summary_text))

    chapter_to_specs: Dict[int, List[BlockSpec]] = {}
    for spec in specs:
        chapter_to_specs.setdefault(spec.chapter, []).append(spec)

    chapter_rows: List[Tuple[int, int, int, int]] = []
    for chapter in range(1, 8):
        current_chars = int(chapter_chars.get(chapter, 0))
        min_chars = int(CHAPTER_MIN_CHARS.get(chapter, 0))
        gap = max(0, min_chars - current_chars)
        chapter_rows.append((chapter, current_chars, min_chars, gap))

    priority_chapters = [f"Chapter {chapter} ({gap} chars short)" for chapter, _, _, gap in sorted(chapter_rows, key=lambda item: item[3], reverse=True) if gap > 0]

    lines = [
        "[Codex Rewrite Task]",
        f"Topic: {DISEASE_NAME}",
        f"Current total chars (chapters + summary, no whitespace): {total_chars}",
        "Target total chars: 30000-34000",
        f"Minimum chars still needed: {total_gap}",
        f"Current summary chars: {summary_chars} (recommended: 1200-1500)",
        "",
        "[Read First]",
        f"- {OUT_ROOT / CODEX_GAP_PANEL_NAME}",
        f"- {OUT_ROOT / CHAPTER_PRECHECK_NAME}",
        f"- {OUT_ROOT / CH4_NARRATIVE_BRIEF_NAME}",
        "",
        "[Primary Goal]",
        "- Directly rewrite and overwrite ch01.txt ~ ch07.txt and summary.txt.",
        "- Hit the gate in one pass instead of writing an obviously under-length draft first.",
        "- Any newly added content must be grounded in the evidence pool, existing figure scope, or chapter-4 structured data.",
        "",
        "[Hard Rules]",
        f"1) Every chapter should meet the script floor; a shortfall within {CHAPTER_CHAR_TOLERANCE} chars is acceptable, and the full draft must stay within 30000-34000 chars.",
        "2) Chapter 4 may only use Excel-derived facts, ch04_codex_extract.json, and generated figure scope. Do not invent new market numbers.",
        "3) New medical, policy, or pathway claims must be traceable to 00_evidence.txt, refs.txt, manifest_fig.csv, or the existing draft context.",
        "4) Keep citation numbering consistent. Do not fabricate citation IDs.",
        "5) Prioritize cause stratification, mechanism, diagnosis, evidence, boundaries, and red flags in Chapters 1-3.",
        "6) Prioritize segmentation, pathway differences, adherence mechanisms, policy constraints, forecast assumptions, and strategy actions in Chapters 5-7.",
        "7) Non-chapter-4 figure source lines must be specific. Do not use only 'public source??' style placeholders.",
        "8) Avoid empty management jargon. Each paragraph should carry a real medical or market anchor.",
        "",
        "[Suggested Order]",
        ("- Fix under-length chapters first: " + ", ".join(priority_chapters)) if priority_chapters else "- Chapter lengths already pass; focus on weak blocks and summary quality first.",
        "- Then fix low-anchor blocks: " + (", ".join(low_anchor_blocks) if low_anchor_blocks else "none"),
        "- Then fix low-medical-density blocks: " + (", ".join(medical_density_failed) if medical_density_failed else "none"),
        "- Rewrite the summary last so it reflects the final body instead of duplicating the outline.",
        "",
        "[Chapter Tasks]",
    ]
    for chapter, current_chars, min_chars, gap in chapter_rows:
        status = "below gate" if gap > CHAPTER_CHAR_TOLERANCE else ("within tolerance" if gap > 0 else "pass")
        lines.append(f"- Chapter {chapter}: current={current_chars}, minimum={min_chars}, shortfall={gap}, tolerance={CHAPTER_CHAR_TOLERANCE} ({status})")
        for spec in chapter_to_specs.get(chapter, []):
            topic_text = ", ".join(spec.topics)
            fig_text = spec.fig_ids if spec.fig_ids else "none"
            lines.append(
                f"  - {spec.block_id} {spec.subtitle} | block target={spec.target_chars} | focus={topic_text} | evidence={spec.evidence_ids} | figures={fig_text}"
            )
    lines.extend(
        [
            "",
            "[Current Weak Spots]",
            "- Under-length chapters: " + (", ".join(chapter_len_fails) if chapter_len_fails else "none"),
            "- Low-anchor blocks: " + (", ".join(low_anchor_blocks) if low_anchor_blocks else "none"),
            "- Low medical density blocks (chapters 1-3): " + (", ".join(medical_density_failed) if medical_density_failed else "none"),
            "",
            "[Delivery]",
            f"- Overwrite: {OUT_ROOT / 'ch01.txt'} ~ {OUT_ROOT / 'ch07.txt'}",
            f"- Overwrite: {OUT_ROOT / 'summary.txt'}",
            "- Do not write any new artifact to the repo root.",
            f"- Re-run after rewriting: python scripts/run_pipeline.py --topic \"{DISEASE_NAME}\"",
        ]
    )
    return "\n".join(lines)



def build_codex_content_blueprint(specs: List[BlockSpec]) -> str:
    chapter_to_specs: Dict[int, List[BlockSpec]] = {}
    for spec in specs:
        chapter_to_specs.setdefault(spec.chapter, []).append(spec)

    chapter_target_map: Dict[int, int] = {chapter: sum(item.target_chars for item in items) for chapter, items in chapter_to_specs.items()}
    total_target = sum(chapter_target_map.values())
    summary_target_floor = 1200
    draft_floor = max(31200, total_target + summary_target_floor)
    draft_ceiling = 33000

    lines = [
        "[Codex Writing Blueprint]",
        f"Topic: {DISEASE_NAME}",
        f"Recommended first-draft total chars: {draft_floor}-{draft_ceiling}",
        "Goal: land a gate-ready first draft instead of backfilling later.",
        "Rule of thumb: keep each chapter 150-300 chars above the hard floor; summary should usually be 1200-1500 chars.",
        "",
        "[Read First]",
        f"1) {OUT_ROOT / '00_evidence.txt'}",
        f"2) {OUT_ROOT / 'manifest_text.csv'}",
        f"3) {OUT_ROOT / 'manifest_fig.csv'}",
        f"4) {OUT_ROOT / 'ch04_codex_extract.json'}",
        f"5) {OUT_ROOT / CH4_NARRATIVE_BRIEF_NAME}",
        f"6) {OUT_ROOT / CODEX_GAP_PANEL_NAME}",
        f"7) {OUT_ROOT / CHAPTER_PRECHECK_NAME}",
        f"8) {OUT_ROOT / 'figure_specs.json'} (if present)",
        "",
        "[Output Files]",
        f"- {OUT_ROOT / 'ch01.txt'} ~ {OUT_ROOT / 'ch07.txt'}",
        f"- {OUT_ROOT / 'summary.txt'}",
        "- Keep all generated or overwritten files inside the current autofile topic folder.",
        "",
        "[Hard Writing Rules]",
        "1) Finish the body first, then write the summary.",
        "2) Chapter 4 must stay inside Excel-derived scope only. Do not add new market numbers, shares, or growth rates.",
        "3) Each paragraph should contain concrete anchors: mechanism, pathway, evidence, policy basis, figure scope, or market fact.",
        "4) Keep citations consistent and distributed across paragraphs instead of clustering them at the end.",
        "5) Non-chapter-4 figure source lines must be specific to guideline/review/institution level.",
        "6) If evidence is weak, narrow the claim boundary instead of padding with generic language.",
        "7) If fig_2_3 is not ready, write fig23_codex_spec.json first before stage3-stage5.",
        "",
        "[Chapter Targets]",
    ]
    for chapter in range(1, 8):
        floor = int(CHAPTER_MIN_CHARS.get(chapter, 0))
        target = int(max(chapter_target_map.get(chapter, floor), floor + 200))
        lines.append(f"- Chapter {chapter}: hard floor={floor}; recommended first-draft target={target}-{target + 200}")
        for spec in chapter_to_specs.get(chapter, []):
            topic_text = ", ".join(spec.topics)
            fig_text = spec.fig_ids if spec.fig_ids else "none"
            lines.append(
                f"  - {spec.block_id} {spec.subtitle} | block target={spec.target_chars} | focus={topic_text} | evidence={spec.evidence_ids} | figures={fig_text}"
            )
    lines.extend(
        [
            "",
            "[Suggested Workflow]",
            "1) Read the evidence pool, text manifest, figure manifest, and chapter-4 structured data before writing.",
            "2) Write Chapters 1-3 first, then Chapter 4, then Chapters 5-7.",
            "3) Write the summary last so it reflects the final logic instead of repeating the table of contents.",
            "4) Run stage3-stage5 after the body is complete.",
            f"5) Command: python scripts/run_pipeline.py --topic \"{DISEASE_NAME}\"",
        ]
    )
    return "\n".join(lines)



def build_fig23_codex_spec_template() -> Dict[str, object]:
    return {
        "schema_version": "fig23_codex_v1",
        "topic": DISEASE_NAME,
        "disease": DISEASE_NAME,
        "authored_by": "codex",
        "layout_mode": "layered_path",
        "title": compose_figure_title("2-3", f"{DISEASE_NAME}关键关系分层路径图"),
        "caption": compose_figure_title("2-3", f"{DISEASE_NAME}关键关系分层路径图"),
        "source_line": "数据来源：请由当前 Codex 会话根据正文与证据池填写为具体来源行",
        "layered_path": {
            "left_title": "上游决定因素",
            "center_title": "核心病理/修复过程",
            "right_title": "下游后果与管理结果",
            "left_nodes": [
                "待由Codex替换",
                "待由Codex替换",
                "待由Codex替换"
            ],
            "core_label": "待由Codex替换",
            "right_nodes": [
                "待由Codex替换",
                "待由Codex替换",
                "待由Codex替换"
            ]
        }
    }


def build_fig23_codex_prompt() -> str:
    return "\n".join(
        [
            "[Codex fig_2_3 Authoring Task]",
            f"Goal: write only {OUT_ROOT / FIG23_CODEX_SPEC_NAME}. Do not create extra markdown or free-form explanation files.",
            "",
            "[Read First]",
            f"1) {OUT_ROOT / 'ch02.txt'}",
            f"2) {OUT_ROOT / '00_evidence.txt'}",
            f"3) {OUT_ROOT / 'figure_specs.json'} (if present)",
            f"4) {OUT_ROOT / FIG23_CODEX_SPEC_TEMPLATE_NAME}",
            "",
            "[Output Rules]",
            "1) authored_by must be codex.",
            "2) Prefer layout_mode=layered_path; use dual_panel only when layered_path cannot keep the figure readable.",
            "3) title/caption may be full or short, because the script normalizes figure numbering.",
            "4) source_line must be specific to guideline/review/institution level.",
            "5) If using layered_path, fill left_nodes, core_label, and right_nodes.",
            "6) If using dual_panel, key edges must use via/sign_xy to avoid geometry collisions.",
            "",
            "[Semantic Rules]",
            "- Prefer a simple upstream -> core mechanism -> downstream outcome/management path.",
            "- Keep 2-3 nodes per side whenever possible.",
            "- Reduce node count before adding more geometry complexity.",
            "- Keep layer semantics consistent; do not mix systems, mechanisms, and management outcomes at the same level without reason.",
            "",
            "[Visual Rules]",
            "- No same-layer crowding.",
            "- No arrows touching text boxes too closely or crossing over label text.",
            "- Avoid unnecessary line crossings.",
            "- If a label is too long, shorten it first; only then consider explicit line breaks.",
            "",
            "[Done Means]",
            f"- The JSON at {OUT_ROOT / FIG23_CODEX_SPEC_NAME} can be consumed by stage3 directly.",
            "- If uncertain, delete weak nodes and preserve readability.",
        ]
    )



def write_codex_preflight_assets() -> None:
    specs = build_block_specs()
    write_text(OUT_ROOT / CODEX_CONTENT_BLUEPRINT_NAME, build_codex_content_blueprint(specs) + "\n")
    write_json(OUT_ROOT / FIG23_CODEX_SPEC_TEMPLATE_NAME, build_fig23_codex_spec_template())
    write_text(OUT_ROOT / FIG23_CODEX_PROMPT_NAME, build_fig23_codex_prompt() + "\n")
    write_json(OUT_ROOT / FIGURE_SPECS_CODEX_TEMPLATE_NAME, build_semantic_figure_specs_template())
    write_text(OUT_ROOT / FIGURE_SPECS_CODEX_PROMPT_NAME, build_figure_specs_codex_prompt() + "\n")
    write_text(OUT_ROOT / SEMANTIC_REVIEW_PROMPT_NAME, build_semantic_review_prompt() + "\n")
    write_codex_progress_assets(specs)


def build_fig23_review_prompt() -> str:
    return "\n".join(
        [
            "[fig_2_3 Visual Review Task]",
            "Review both the rendered figure and its config:",
            f"1) {FIG_DIR / 'fig_2_3.png'}",
            f"2) {OUT_ROOT / FIG23_CODEX_SPEC_NAME}",
            "",
            "[Check List]",
            "- Is the causal direction clinically and logically sound?",
            "- Are there overlaps, arrow-to-box collisions, text coverage, or labels that feel crowded?",
            "- Is there enough whitespace between nodes, especially within the same layer?",
            "- Are bidirectional/dashed relations still readable after geometry routing?",
            "- Are layer semantics internally consistent?",
            "- Do title, caption, and source_line match the actual figure meaning?",
            "",
            "[Return Format]",
            "Use this exact structure:",
            "1) Verdict: PASS / FAIL",
            "2) Must-fix issues: list only the items that block acceptance; write 'none' if empty",
            "3) Nice-to-improve issues: optional improvements; write 'none' if empty",
            "4) Suggested field rewrites: give directly usable config edits, prioritizing layered_path.left_nodes / core_label / right_nodes; only mention dual_panel via/sign_xy when needed",
            "",
            "[Decision Rule]",
            "- If the figure is already clear, say PASS explicitly and avoid change-for-change's-sake edits.",
            "- If it fails, prefer the smallest field changes that restore readability.",
            "- If the graph is too dense, remove nodes before adding more routing lines.",
        ]
    )



def run_txt_stage_checks(specs: List[BlockSpec], block_text: Dict[str, str], summary_text: str) -> Tuple[str, bool]:
    metrics = collect_text_quality_metrics(specs, block_text)
    stale_counts: Dict[str, int] = metrics["stale_counts"]  # type: ignore[assignment]
    drift_counts: Dict[str, int] = metrics["drift_counts"]  # type: ignore[assignment]
    max_sentence_dup: int = metrics["max_sentence_dup"]  # type: ignore[assignment]
    top_sentence_dups: List[Tuple[str, int]] = metrics["top_sentence_dups"]  # type: ignore[assignment]
    metric_logic_issues: List[str] = metrics["metric_logic_issues"]  # type: ignore[assignment]
    dup_prefix_hits: int = metrics["dup_prefix_hits"]  # type: ignore[assignment]
    chapter_stats: List[Tuple[int, int, int, float, int, int]] = metrics["chapter_stats"]  # type: ignore[assignment]
    low_anchor_chapters: List[int] = metrics["low_anchor_chapters"]  # type: ignore[assignment]
    chapter_dup_fails: List[int] = metrics["chapter_dup_fails"]  # type: ignore[assignment]
    chapter_no_cites: List[int] = metrics["chapter_no_cites"]  # type: ignore[assignment]
    chapter_len_fails: List[str] = metrics["chapter_len_fails"]  # type: ignore[assignment]
    medical_density_failed: List[str] = metrics["medical_density_failed"]  # type: ignore[assignment]
    cagr_logic_ok: bool = metrics["cagr_logic_ok"]  # type: ignore[assignment]
    cr5_logic_ok: bool = metrics["cr5_logic_ok"]  # type: ignore[assignment]
    total_chars = sum(len(re.sub(r"\s+", "", block_text[s.block_id])) for s in specs) + len(re.sub(r"\s+", "", summary_text))

    fail_reasons: List[str] = []
    if max(stale_counts.values()) > 8:
        fail_reasons.append("高频套话超阈值")
    if max(drift_counts.values()) > 0:
        fail_reasons.append("报告定位跑偏（管理话术残留）")
    if max_sentence_dup >= 4:
        fail_reasons.append("全书句级重复超阈值")
    if metric_logic_issues:
        fail_reasons.append("指标逻辑或单位冲突")
    if dup_prefix_hits > 0:
        fail_reasons.append("标题重复前缀残留")
    if low_anchor_chapters:
        fail_reasons.append("章节事实锚点覆盖不足")
    if chapter_dup_fails:
        fail_reasons.append("章节内句级重复超阈值")
    if chapter_no_cites:
        fail_reasons.append("章节引用缺失")
    if not (30000 <= total_chars <= 34000):
        fail_reasons.append("总字数未达到30000-34000")
    if chapter_len_fails:
        fail_reasons.append("分章最低字数未达标")
    if medical_density_failed:
        fail_reasons.append("第1-3章医学要素不足")
    if not cagr_logic_ok:
        fail_reasons.append("第四章CAGR逻辑冲突")
    if not cr5_logic_ok:
        fail_reasons.append("第四章CR5叙述冲突")

    passed = len(fail_reasons) == 0

    lines = [
        "【TXT阶段质量检查】",
        f"流程模式：{WORKFLOW_MODE}",
        f"总字数（章节+总结，去空白）：{total_chars}",
        f"标题重复前缀命中数：{dup_prefix_hits}",
        "高频套话统计：" + ", ".join([f"{k}={v}" for k, v in stale_counts.items()]),
        "管理话术命中：" + ", ".join([f"{k}={v}" for k, v in drift_counts.items()]),
        f"句级最大重复次数：{max_sentence_dup}",
        "句级重复TOP5：" + ("; ".join([f"{c}x:{s[:36]}..." for s, c in top_sentence_dups]) if top_sentence_dups else "无"),
        "指标逻辑冲突：" + ("；".join(metric_logic_issues) if metric_logic_issues else "无"),
        "第四章逻辑一致性：" + ("通过" if cagr_logic_ok else "不通过"),
        "第四章CR5叙述一致性：" + ("通过" if cr5_logic_ok else "不通过"),
        "第1-3章医学密度不足block：" + (", ".join(medical_density_failed) if medical_density_failed else "无"),
        "",
        "【分章统计】",
    ]
    for ch, chars, para_cnt, anchor_cov, ch_dup, cite_cnt in chapter_stats:
        lines.append(
            f"第{ch}章：字数={chars}，段落数={para_cnt}，锚点覆盖={anchor_cov*100:.1f}%"
            f"，句级最大重复={ch_dup}，引用数={cite_cnt}"
        )
    lines.extend(
        [
            "",
            "【分章问题】",
            "锚点覆盖不足章节：" + (", ".join([str(x) for x in low_anchor_chapters]) if low_anchor_chapters else "无"),
            "章节内句级重复超阈值章节：" + (", ".join([str(x) for x in chapter_dup_fails]) if chapter_dup_fails else "无"),
            "引用缺失章节：" + (", ".join([str(x) for x in chapter_no_cites]) if chapter_no_cites else "无"),
            f"总字数30000-34000：{'通过' if (30000 <= total_chars <= 34000) else '不通过'}",
            "分章字数下限未达标：" + (", ".join(chapter_len_fails) if chapter_len_fails else "无"),
            "",
            "【TXT闸门判定】",
            f"结果：{'通过' if passed else '不通过'}",
            "失败原因：" + ("；".join(fail_reasons) if fail_reasons else "无"),
        ]
    )
    report = "\n".join(lines)
    write_text(OUT_ROOT / "txt_stage_qa.txt", report + "\n")
    write_text(OUT_ROOT / "codex_rewrite_prompt.txt", build_codex_rewrite_prompt(specs, metrics, summary_text) + "\n")
    write_text(OUT_ROOT / "fig23_review_prompt.txt", build_fig23_review_prompt() + "\n")
    write_codex_progress_assets(specs, block_text, summary_text)
    return report, passed



def run_checks(specs: List[BlockSpec], block_text: Dict[str, str], fig_rows: List[Dict[str, str]], summary_text: str) -> Tuple[str, bool]:
    total_chars = sum(len(re.sub(r"\s+", "", block_text[s.block_id])) for s in specs) + len(re.sub(r"\s+", "", summary_text))
    fig_files = sorted(FIG_DIR.glob("fig_*.png"))
    fig_count = len(fig_files)
    ch4_fig_count = len([f for f in fig_files if f.stem.startswith("fig_4_")])

    doc_xml = extract_docx_text_xml(FINAL_DOCX)
    current_disease_token = DISEASE_NAME.strip()
    legacy_placeholder_tokens = [t for t in LEGACY_DISEASE_TOKENS if t.strip() and t.strip() != current_disease_token]
    placeholder_keys = ["XXX", "<<<", "AAA"] + legacy_placeholder_tokens
    placeholder_hits = {k: doc_xml.count(k) for k in placeholder_keys}

    with zipfile.ZipFile(FINAL_DOCX, "r") as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore") if "word/document.xml" in names else ""
        settings_xml = zf.read("word/settings.xml").decode("utf-8", errors="ignore") if "word/settings.xml" in names else ""
        rels_xml = (
            zf.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
            if "word/_rels/document.xml.rels" in names
            else ""
        )

        pgnum_start_present = bool(re.search(r"<w:pgNumType[^>]*\bw:start=\"[^\"]+\"", document_xml))
        footer_ids = re.findall(r"<w:footerReference[^>]*r:id=\"([^\"]+)\"", document_xml)
        unique_footer_ids = sorted(set(footer_ids))
        footer_uniform_ok = len(unique_footer_ids) == 1

        rel_map: Dict[str, str] = {}
        rel_pattern = re.compile(
            r'<Relationship[^>]*\bId="([^"]+)"[^>]*\bType="([^"]+)"[^>]*\bTarget="([^"]+)"[^>]*/>'
        )
        for rid, rel_type, target in rel_pattern.findall(rels_xml):
            if not rel_type.endswith("/footer"):
                continue
            p = target.replace("\\", "/")
            if p.startswith("/"):
                p = p.lstrip("/")
            elif not p.startswith("word/"):
                p = f"word/{p}"
            rel_map[rid] = p

        footer_has_page = False
        footer_targets_with_page: List[str] = []
        footer_textbox_targets: List[str] = []
        for rid in unique_footer_ids:
            footer_target = rel_map.get(rid, "")
            if not footer_target or footer_target not in names:
                continue
            footer_xml = zf.read(footer_target).decode("utf-8", errors="ignore")
            has_page = bool(
                re.search(r'<w:instrText[^>]*>\s*[^<]*\bPAGE\b', footer_xml)
                or re.search(r'<w:fldSimple[^>]*w:instr="[^"]*\bPAGE\b', footer_xml)
            )
            if has_page:
                footer_targets_with_page.append(footer_target)
                if ("<w:txbxContent" in footer_xml) or ('txBox="1"' in footer_xml):
                    footer_textbox_targets.append(footer_target)
        footer_has_page = len(footer_targets_with_page) > 0
        footer_target = ",".join(sorted(set(footer_targets_with_page)))
        footer_page_in_textbox = len(footer_textbox_targets) > 0
        footer_textbox_target = ",".join(sorted(set(footer_textbox_targets)))

        update_fields_ok = bool(re.search(r'<w:updateFields[^>]*w:val="true"', settings_xml))

    doc = Document(str(FINAL_DOCX))
    cap_count = sum(1 for p in doc.paragraphs if p.text.strip().startswith("图表"))
    src_count = sum(1 for p in doc.paragraphs if p.text.strip().startswith("数据来源："))
    src_ch4_count = sum(1 for p in doc.paragraphs if p.text.strip() == "数据来源：米内网")

    title_registry_error = ""
    title_registry: Dict[str, str] = {}
    title_missing: List[str] = []
    title_no_serial: List[str] = []
    title_caption_mismatch: List[str] = []
    try:
        title_registry = load_figure_title_registry()
    except Exception as e:
        title_registry_error = str(e)
    if not title_registry_error:
        for row in fig_rows:
            fid = str(row.get("fig_id", "")).strip()
            expected_caption = normalize_disease_text(str(row.get("caption", "")).strip())
            rendered_title = title_registry.get(fid, "")
            if not rendered_title:
                title_missing.append(fid)
                continue
            if not FIG_TITLE_SERIAL_RE.match(rendered_title):
                title_no_serial.append(fid)
            if rendered_title != expected_caption:
                title_caption_mismatch.append(fid)

    fig23_expected = fig23_expected_caption()
    fig23_forbidden = fig23_forbidden_captions()
    disallow_terms = fig23_disallow_nodes()
    disallow_term_hits: List[str] = []
    fig23_manifest_caption = ""
    fig23_registry_caption = ""
    fig23_manifest_ok = True
    fig23_registry_ok = True
    fig23_forbidden_hits: List[str] = []

    fig23_manifest_row = next((r for r in fig_rows if str(r.get("fig_id", "")).strip() == "fig_2_3"), None)
    if fig23_manifest_row:
        fig23_manifest_caption = normalize_disease_text(str(fig23_manifest_row.get("caption", "")).strip())
        fig23_manifest_ok = fig23_manifest_caption == fig23_expected
    else:
        fig23_manifest_ok = False

    fig23_registry_caption = normalize_disease_text(title_registry.get("fig_2_3", "").strip())
    fig23_registry_ok = fig23_registry_caption == fig23_expected
    for term in disallow_terms:
        if not term:
            continue
        if ((term in fig23_expected) or (term in fig23_manifest_caption) or (term in fig23_registry_caption)) and term not in disallow_term_hits:
            disallow_term_hits.append(term)

    for p in doc.paragraphs:
        ptxt = normalize_disease_text(p.text.strip())
        if ptxt and ptxt in fig23_forbidden and ptxt not in fig23_forbidden_hits:
            fig23_forbidden_hits.append(ptxt)

    fig23_struct = validate_fig23_structural_rules()
    fig23_spec_source = fig23_spec_origin()
    fig23_codex_ok = fig23_codex_authored_ok()
    fig23_layout_mode_active = str(fig23_struct.get("layout_mode", ""))
    fig23_layout_required_ok = fig23_layout_mode_active in {"dual_panel", "layered_path"}
    fig23_causal_issues: List[str] = list(fig23_struct.get("causal_direction_issues", []) or [])
    fig23_overlap_issues: List[str] = list(fig23_struct.get("same_track_overlap_issues", []) or [])
    fig23_bidir_issues: List[str] = list(fig23_struct.get("bidirectional_readability_issues", []) or [])
    fig23_layer_issues: List[str] = list(fig23_struct.get("layer_consistency_issues", []) or [])
    fig23_node_spacing_issues: List[str] = list(fig23_struct.get("node_spacing_issues", []) or [])

    metrics = collect_text_quality_metrics(specs, block_text)
    stale_counts: Dict[str, int] = metrics["stale_counts"]  # type: ignore[assignment]
    drift_counts: Dict[str, int] = metrics["drift_counts"]  # type: ignore[assignment]
    max_sentence_dup: int = metrics["max_sentence_dup"]  # type: ignore[assignment]
    top_sentence_dups: List[Tuple[str, int]] = metrics["top_sentence_dups"]  # type: ignore[assignment]
    metric_logic_issues: List[str] = metrics["metric_logic_issues"]  # type: ignore[assignment]
    dup_prefix_hits: int = metrics["dup_prefix_hits"]  # type: ignore[assignment]
    low_anchor_blocks: List[str] = metrics["low_anchor_blocks"]  # type: ignore[assignment]
    chapter_chars: Dict[int, int] = metrics["chapter_chars"]  # type: ignore[assignment]
    chapter_len_fails: List[str] = metrics["chapter_len_fails"]  # type: ignore[assignment]
    cagr_logic_ok: bool = metrics["cagr_logic_ok"]  # type: ignore[assignment]
    cr5_logic_ok: bool = metrics["cr5_logic_ok"]  # type: ignore[assignment]
    medical_density_failed: List[str] = metrics["medical_density_failed"]  # type: ignore[assignment]

    # Reference traceability checks.
    refs_lines = [normalize_reference_line(x) for x in (OUT_ROOT / "refs.txt").read_text(encoding="utf-8").splitlines() if x.strip()]
    normalized_excel_name = normalize_disease_text(EXCEL_PATH.name)
    bad_ref_rows: List[str] = []
    for idx, line in enumerate(refs_lines, start=1):
        has_year = bool(re.search(r"(19|20)\d{2}", line))
        has_url_or_file = ("http" in line) or (EXCEL_PATH.name in line) or (normalized_excel_name in line)
        has_org_and_title = line.startswith("[") and len(line.split(". ")) >= 2
        url_m = re.search(r"https?://[^\s]+", line)
        nonspecific = is_nonspecific_reference_url(url_m.group(0)) if url_m else False
        if (not (has_year and has_url_or_file and has_org_and_title)) or nonspecific:
            bad_ref_rows.append(str(idx))

    ref_chain = collect_reference_chain_metrics(specs, block_text, summary_text)
    evidence_count: int = int(ref_chain["evidence_count"])  # type: ignore[assignment]
    ref_count: int = int(ref_chain["ref_count"])  # type: ignore[assignment]
    evidence_parse_errors: List[str] = ref_chain["evidence_parse_errors"]  # type: ignore[assignment]
    ref_parse_errors: List[str] = ref_chain["ref_parse_errors"]  # type: ignore[assignment]
    evidence_id_dup: List[int] = ref_chain["evidence_id_dup"]  # type: ignore[assignment]
    ref_id_dup: List[int] = ref_chain["ref_id_dup"]  # type: ignore[assignment]
    evidence_seq_ok: bool = bool(ref_chain["evidence_seq_ok"])
    ref_seq_ok: bool = bool(ref_chain["ref_seq_ok"])
    evidence_ref_gap: List[int] = ref_chain["evidence_ref_gap"]  # type: ignore[assignment]
    cited_count: int = int(ref_chain["cited_count"])  # type: ignore[assignment]
    dangling_cites: List[int] = ref_chain["dangling_cites"]  # type: ignore[assignment]
    uncited_refs: List[int] = ref_chain["uncited_refs"]  # type: ignore[assignment]
    citation_coverage: float = float(ref_chain["citation_coverage"])  # type: ignore[assignment]
    evidence_bad_year: List[int] = ref_chain["evidence_bad_year"]  # type: ignore[assignment]
    evidence_bad_source: List[int] = ref_chain["evidence_bad_source"]  # type: ignore[assignment]
    ref_bad_year: List[int] = ref_chain["ref_bad_year"]  # type: ignore[assignment]
    ref_bad_source: List[int] = ref_chain["ref_bad_source"]  # type: ignore[assignment]

    qa_fail_reasons: List[str] = []
    if not (20 <= fig_count <= 30):
        qa_fail_reasons.append("图表总量不在20-30")
    if not (6 <= ch4_fig_count <= 8):
        qa_fail_reasons.append("第四章图表数量不在6-8")
    if cap_count != src_count or src_count != fig_count:
        qa_fail_reasons.append("图表标题/来源行数与图表数不一致")
    if src_ch4_count != ch4_fig_count:
        qa_fail_reasons.append("第四章数据来源“米内网”行数不一致")
    if not all(v == 0 for v in placeholder_hits.values()):
        qa_fail_reasons.append("存在占位符残留")
    if not footer_has_page:
        qa_fail_reasons.append("页脚PAGE域缺失")
    if footer_page_in_textbox:
        qa_fail_reasons.append("页脚PAGE域仍位于文本框中")
    if pgnum_start_present:
        qa_fail_reasons.append("检测到页码重置配置(w:start)")
    if not footer_uniform_ok:
        qa_fail_reasons.append("页脚引用不一致（多节footerReference）")
    if not update_fields_ok:
        qa_fail_reasons.append("settings.updateFields未启用")
    if dup_prefix_hits != 0:
        qa_fail_reasons.append("标题重复前缀残留")
    if max(stale_counts.values()) > 8:
        qa_fail_reasons.append("高频套话超阈值")
    if max(drift_counts.values()) > 0:
        qa_fail_reasons.append("报告定位跑偏（管理话术残留）")
    if max_sentence_dup >= 4:
        qa_fail_reasons.append("句级重复超阈值")
    if metric_logic_issues:
        qa_fail_reasons.append("指标逻辑或单位冲突")
    if not (30000 <= total_chars <= 34000):
        qa_fail_reasons.append("总字数未达到30000-34000")
    if chapter_len_fails:
        qa_fail_reasons.append("分章最低字数未达标")
    if low_anchor_blocks:
        qa_fail_reasons.append("事实锚点覆盖率不足")
    if not cagr_logic_ok:
        qa_fail_reasons.append("第四章CAGR逻辑冲突")
    if not cr5_logic_ok:
        qa_fail_reasons.append("第四章CR5叙述冲突")
    if medical_density_failed:
        qa_fail_reasons.append("第1-3章医学密度不足")
    if bad_ref_rows:
        qa_fail_reasons.append("参考文献可核验性不足")
    if evidence_parse_errors or ref_parse_errors:
        qa_fail_reasons.append("证据池或参考文献解析失败")
    if evidence_count == 0 or ref_count == 0:
        qa_fail_reasons.append("证据池或参考文献为空")
    if evidence_count != ref_count:
        qa_fail_reasons.append("证据池条数与参考文献条数不一致")
    if evidence_id_dup or ref_id_dup:
        qa_fail_reasons.append("证据ID或参考编号重复")
    if (not evidence_seq_ok) or (not ref_seq_ok):
        qa_fail_reasons.append("证据ID或参考编号未按连续序号编排")
    if evidence_ref_gap:
        qa_fail_reasons.append("证据ID与参考编号集合不一致")
    if dangling_cites:
        qa_fail_reasons.append("正文存在悬空引用编号")
    if evidence_bad_year or evidence_bad_source or ref_bad_year or ref_bad_source:
        qa_fail_reasons.append("证据或参考存在年份/来源缺失")
    if ref_count > 0 and citation_coverage < 0.60:
        qa_fail_reasons.append("正文引用覆盖率不足（<60%）")
    if title_registry_error:
        qa_fail_reasons.append("缺少图题注册表")
    if title_missing:
        qa_fail_reasons.append("图题注册缺失")
    if title_no_serial:
        qa_fail_reasons.append("图片本体图题缺少图表序号")
    if title_caption_mismatch:
        qa_fail_reasons.append("图片本体图题与manifest图注不一致")
    if fig23_spec_source != "codex_spec":
        qa_fail_reasons.append("图表2-3未使用Codex专用配置文件")
    if not fig23_codex_ok:
        qa_fail_reasons.append("图表2-3配置未标记为Codex authored")
    if not fig23_layout_required_ok:
        qa_fail_reasons.append("图表2-3未使用允许的分层替代表达")
    if not fig23_manifest_ok:
        qa_fail_reasons.append("图表2-3 manifest图注未匹配当前疾病画像语义口径")
    if not fig23_registry_ok:
        qa_fail_reasons.append("图表2-3图片主标题未匹配当前疾病画像语义口径")
    if fig23_forbidden_hits:
        qa_fail_reasons.append("图表2-3命中禁用旧标题")
    if disallow_term_hits:
        qa_fail_reasons.append("文档出现关系图禁用节点术语")
    if fig23_causal_issues:
        qa_fail_reasons.append("图表2-3因果方向检查不通过")
    if fig23_overlap_issues:
        qa_fail_reasons.append("图表2-3存在同轨重叠线")
    if fig23_bidir_issues:
        qa_fail_reasons.append("图表2-3双向关系可读性不足")
    if fig23_layer_issues:
        qa_fail_reasons.append("图表2-3层级一致性检查不通过")
    if fig23_node_spacing_issues:
        qa_fail_reasons.append("图表2-3节点间距或留白不足")
    qa_passed = len(qa_fail_reasons) == 0

    lines = [
        "【QA检查结果】",
        f"流程模式：{WORKFLOW_MODE}",
        f"总字数（章节+总结，去空白）：{total_chars}",
        f"图表总数：{fig_count}",
        f"第四章图表数：{ch4_fig_count}",
        f"manifest_fig行数：{len(fig_rows)}",
        f"文档图表标题行数：{cap_count}",
        f"文档数据来源行数：{src_count}",
        f"第四章“数据来源：米内网”行数：{src_ch4_count}",
        f"标题重复前缀命中数：{dup_prefix_hits}",
        "占位符残留统计：" + ", ".join([f"{k}={v}" for k, v in placeholder_hits.items()]),
        f"页码重置配置(w:start)：{'检测到（存在重置风险）' if pgnum_start_present else '未检测到'}",
        f"footerReference统一ID（信息项）：{'是' if footer_uniform_ok else '否'}（IDs={','.join(unique_footer_ids) if unique_footer_ids else '无'}）",
        f"引用页脚含PAGE域：{'通过' if footer_has_page else '不通过'}（{footer_target if footer_target else '未解析'}）",
        f"页脚PAGE域位于文本框：{'是' if footer_page_in_textbox else '否'}（{footer_textbox_target if footer_textbox_target else '无'}）",
        f"settings.updateFields=true：{'通过' if update_fields_ok else '不通过'}",
        "高频套话统计：" + ", ".join([f"{k}={v}" for k, v in stale_counts.items()]),
        "管理话术命中：" + ", ".join([f"{k}={v}" for k, v in drift_counts.items()]),
        f"句级最大重复次数：{max_sentence_dup}",
        "句级重复TOP5：" + ("; ".join([f"{c}x:{s[:36]}..." for s, c in top_sentence_dups]) if top_sentence_dups else "无"),
        "指标逻辑冲突：" + ("；".join(metric_logic_issues) if metric_logic_issues else "无"),
        "事实锚点覆盖不足block：" + (", ".join(low_anchor_blocks) if low_anchor_blocks else "无"),
        "分章字数下限未达标：" + (", ".join(chapter_len_fails) if chapter_len_fails else "无"),
        "第四章逻辑一致性：" + ("通过" if cagr_logic_ok else "不通过"),
        "第四章CR5叙述一致性：" + ("通过" if cr5_logic_ok else "不通过"),
        "第1-3章医学密度不足block：" + (", ".join(medical_density_failed) if medical_density_failed else "无"),
        "参考文献可核验异常行：" + (", ".join(bad_ref_rows) if bad_ref_rows else "无"),
        "证据池条数：" + str(evidence_count),
        "参考文献条数：" + str(ref_count),
        "正文引用编号数（去重）：" + str(cited_count),
        "引用覆盖率（正文命中参考）：" + f"{citation_coverage*100:.1f}%",
        "证据池解析异常：" + (", ".join(evidence_parse_errors) if evidence_parse_errors else "无"),
        "参考文献解析异常：" + (", ".join(ref_parse_errors) if ref_parse_errors else "无"),
        "证据ID重复：" + (", ".join([str(x) for x in evidence_id_dup]) if evidence_id_dup else "无"),
        "参考编号重复：" + (", ".join([str(x) for x in ref_id_dup]) if ref_id_dup else "无"),
        "证据ID连续性：" + ("通过" if evidence_seq_ok else "不通过"),
        "参考编号连续性：" + ("通过" if ref_seq_ok else "不通过"),
        "证据ID与参考编号差异：" + (", ".join([str(x) for x in evidence_ref_gap]) if evidence_ref_gap else "无"),
        "正文悬空引用编号：" + (", ".join([str(x) for x in dangling_cites]) if dangling_cites else "无"),
        "未被正文使用的参考编号：" + (", ".join([str(x) for x in uncited_refs]) if uncited_refs else "无"),
        "证据池年份异常编号：" + (", ".join([str(x) for x in evidence_bad_year]) if evidence_bad_year else "无"),
        "证据池来源异常编号：" + (", ".join([str(x) for x in evidence_bad_source]) if evidence_bad_source else "无"),
        "参考年份异常编号：" + (", ".join([str(x) for x in ref_bad_year]) if ref_bad_year else "无"),
        "参考来源异常编号：" + (", ".join([str(x) for x in ref_bad_source]) if ref_bad_source else "无"),
        "图题注册表：" + (title_registry_error if title_registry_error else "已加载"),
        "图题注册缺失fig_id：" + (", ".join(title_missing) if title_missing else "无"),
        "图片本体图题缺少序号fig_id：" + (", ".join(title_no_serial) if title_no_serial else "无"),
        "图题与manifest图注不一致fig_id：" + (", ".join(title_caption_mismatch) if title_caption_mismatch else "无"),
        "当前疾病画像：" + active_profile_id(),
        "图表2-3来源配置：" + fig23_spec_source,
        "图表2-3 authored_by=codex：" + ("通过" if fig23_codex_ok else "不通过"),
        "图表2-3预期图注：" + fig23_expected,
        "图表2-3 manifest图注：" + (fig23_manifest_caption if fig23_manifest_caption else "缺失"),
        "图表2-3 图片主标题：" + (fig23_registry_caption if fig23_registry_caption else "缺失"),
        "图表2-3 图注口径：" + ("通过" if fig23_manifest_ok else "不通过"),
        "图表2-3 主标题口径：" + ("通过" if fig23_registry_ok else "不通过"),
        "图表2-3 禁用标题命中：" + (", ".join(fig23_forbidden_hits) if fig23_forbidden_hits else "无"),
        "图表关系图禁用节点命中：" + (", ".join(disallow_term_hits) if disallow_term_hits else "无"),
        "图表2-3结构检查布局模式：" + (fig23_layout_mode_active if fig23_layout_mode_active else "N/A"),
        "图表2-3因果方向问题：" + ("；".join(fig23_causal_issues) if fig23_causal_issues else "无"),
        "图表2-3同轨重叠问题：" + ("；".join(fig23_overlap_issues) if fig23_overlap_issues else "无"),
        "图表2-3双向可读性问题：" + ("；".join(fig23_bidir_issues) if fig23_bidir_issues else "无"),
        "图表2-3层级一致性问题：" + ("；".join(fig23_layer_issues) if fig23_layer_issues else "无"),
        "图表2-3节点间距问题：" + ("；".join(fig23_node_spacing_issues) if fig23_node_spacing_issues else "无"),
        "",
        "【约束判定】",
        f"字数是否在30000-34000：{'通过' if (30000 <= total_chars <= 34000) else '不通过'}",
        f"分章最低字数（差{CHAPTER_CHAR_TOLERANCE}字内允许通过）：{'通过' if (not chapter_len_fails) else '不通过'}",
        f"第1章>=3000或差距<=100：{'通过' if chapter_char_gate_ok(1, chapter_chars[1]) else '不通过'}",
        f"第2章>=3500或差距<=100：{'通过' if chapter_char_gate_ok(2, chapter_chars[2]) else '不通过'}",
        f"第3章>=4800或差距<=100：{'通过' if chapter_char_gate_ok(3, chapter_chars[3]) else '不通过'}",
        f"第4章>=3000或差距<=100：{'通过' if chapter_char_gate_ok(4, chapter_chars[4]) else '不通过'}",
        f"第5章>=4800或差距<=100：{'通过' if chapter_char_gate_ok(5, chapter_chars[5]) else '不通过'}",
        f"第6章>=4800或差距<=100：{'通过' if chapter_char_gate_ok(6, chapter_chars[6]) else '不通过'}",
        f"第7章>=4800或差距<=100：{'通过' if chapter_char_gate_ok(7, chapter_chars[7]) else '不通过'}",
        f"图表总量20-30：{'通过' if 20 <= fig_count <= 30 else '不通过'}",
        f"第四章图表6-8：{'通过' if 6 <= ch4_fig_count <= 8 else '不通过'}",
        f"标题行与来源行一致：{'通过' if cap_count == src_count == fig_count else '不通过'}",
        f"第四章来源行固定米内网：{'通过' if src_ch4_count == ch4_fig_count else '不通过'}",
        f"占位符清洗：{'通过' if all(v == 0 for v in placeholder_hits.values()) else '不通过'}",
        f"页码连续性（无w:start重置）：{'通过' if not pgnum_start_present else '不通过'}",
        f"节页脚统一引用：{'通过' if footer_uniform_ok else '不通过'}",
        f"页脚PAGE域存在：{'通过' if footer_has_page else '不通过'}",
        f"页脚PAGE域不在文本框：{'通过' if not footer_page_in_textbox else '不通过'}",
        f"updateFields=true：{'通过' if update_fields_ok else '不通过'}",
        f"重复前缀清洗：{'通过' if dup_prefix_hits == 0 else '不通过'}",
        f"高频套话压缩：{'通过' if max(stale_counts.values()) <= 8 else '不通过'}",
        f"管理话术清洗：{'通过' if max(drift_counts.values()) == 0 else '不通过'}",
        f"句级重复阈值(<4)：{'通过' if max_sentence_dup < 4 else '不通过'}",
        f"指标逻辑一致性：{'通过' if not metric_logic_issues else '不通过'}",
        f"事实锚点覆盖率阈值(>=70%)：{'通过' if not low_anchor_blocks else '不通过'}",
        f"第四章逻辑一致性：{'通过' if cagr_logic_ok else '不通过'}",
        f"第四章CR5叙述一致性：{'通过' if cr5_logic_ok else '不通过'}",
        f"第1-3章医学密度校验：{'通过' if not medical_density_failed else '不通过'}",
        f"引用可核验性：{'通过' if not bad_ref_rows else '不通过'}",
        f"证据池/参考解析：{'通过' if not (evidence_parse_errors or ref_parse_errors) else '不通过'}",
        f"证据池与参考条数一致：{'通过' if evidence_count == ref_count and evidence_count > 0 else '不通过'}",
        f"证据与参考编号连续：{'通过' if (evidence_seq_ok and ref_seq_ok) else '不通过'}",
        f"证据与参考编号集合一致：{'通过' if not evidence_ref_gap else '不通过'}",
        f"正文悬空引用清洗：{'通过' if not dangling_cites else '不通过'}",
        f"引用覆盖率>=60%：{'通过' if (ref_count > 0 and citation_coverage >= 0.60) else '不通过'}",
        f"证据年份/来源完整：{'通过' if not (evidence_bad_year or evidence_bad_source) else '不通过'}",
        f"参考年份/来源完整：{'通过' if not (ref_bad_year or ref_bad_source) else '不通过'}",
        f"图题注册表可用：{'通过' if not title_registry_error else '不通过'}",
        f"图题注册完整性：{'通过' if not title_missing else '不通过'}",
        f"图片本体图题序号：{'通过' if not title_no_serial else '不通过'}",
        f"图题与manifest一致性：{'通过' if not title_caption_mismatch else '不通过'}",
        f"图表2-3图注语义口径：{'通过' if fig23_manifest_ok else '不通过'}",
        f"图表2-3主标题语义口径：{'通过' if fig23_registry_ok else '不通过'}",
        f"图表2-3使用Codex专用配置：{'通过' if fig23_spec_source == 'codex_spec' else '不通过'}",
        f"图表2-3配置标记为Codex authored：{'通过' if fig23_codex_ok else '不通过'}",
        f"图表2-3允许的分层替代表达：{'通过' if fig23_layout_required_ok else '不通过'}",
        f"图表2-3禁用标题清洗：{'通过' if not fig23_forbidden_hits else '不通过'}",
        f"图表关系图禁用节点清洗：{'通过' if not disallow_term_hits else '不通过'}",
        f"图表2-3因果方向合理性：{'通过' if not fig23_causal_issues else '不通过'}",
        f"图表2-3同轨重叠检查：{'通过' if not fig23_overlap_issues else '不通过'}",
        f"图表2-3双向关系可读性：{'通过' if not fig23_bidir_issues else '不通过'}",
        f"图表2-3层级一致性：{'通过' if not fig23_layer_issues else '不通过'}",
        f"图表2-3节点间距/留白：{'通过' if not fig23_node_spacing_issues else '不通过'}",
        "",
        "【最终判定】",
        f"结果：{'通过' if qa_passed else '不通过'}",
        "失败原因：" + ("；".join(qa_fail_reasons) if qa_fail_reasons else "无"),
        "",
        "备注：目录与页码字段已设置updateFields=true，打开Word后全选F9可刷新显示。",
    ]
    report = "\n".join(lines)
    write_text(OUT_ROOT / "fig23_review_prompt.txt", build_fig23_review_prompt() + "\n")
    write_text(OUT_ROOT / "qa_check.txt", report + "\n")
    return report, qa_passed


def ensure_inputs(require_excel: bool = False, require_template: bool = False) -> None:
    if require_excel and not EXCEL_PATH.exists():
        raise FileNotFoundError(f"Missing input Excel: {EXCEL_PATH}")
    if require_template and not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Missing template docx: {TEMPLATE_PATH}")


def load_block_text_from_files(specs: List[BlockSpec]) -> Dict[str, str]:
    pattern = re.compile(
        r"\[\[BLOCK_ID=(?P<id>[^\]]+)\]\]\n(?P<title>[^\n]*)\n(?P<body>.*?)\n\[\[END_BLOCK_ID=(?P=id)\]\]",
        re.S,
    )
    block_text: Dict[str, str] = {}
    for ch in range(1, 8):
        path = OUT_ROOT / f"ch0{ch}.txt"
        if not path.exists():
            raise FileNotFoundError(f"Missing chapter text file: {path}")
        text = path.read_text(encoding="utf-8")
        for m in pattern.finditer(text):
            block_text[m.group("id")] = m.group("body").strip()

    missing = [s.block_id for s in specs if s.block_id not in block_text]
    if missing:
        raise ValueError(f"Missing BLOCK_ID content in chapter txt files: {', '.join(missing)}")
    return block_text


def load_summary_and_refs() -> Tuple[str, str]:
    summary_path = OUT_ROOT / "summary.txt"
    refs_path = OUT_ROOT / "refs.txt"
    if not summary_path.exists():
        raise FileNotFoundError(f"Missing summary file: {summary_path}")
    if not refs_path.exists():
        raise FileNotFoundError(f"Missing refs file: {refs_path}")
    return summary_path.read_text(encoding="utf-8").strip(), refs_path.read_text(encoding="utf-8")


def load_manifest_fig_rows() -> List[Dict[str, str]]:
    manifest_path = OUT_ROOT / "manifest_fig.csv"
    if not manifest_path.exists():
        raise FileNotFoundError(f"Missing manifest file: {manifest_path}")
    with manifest_path.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))
    if not rows:
        raise ValueError(f"Manifest has no rows: {manifest_path}")
    required = ["fig_id", "caption", "输出文件名", "source_line"]
    miss = [k for k in required if k not in rows[0]]
    if miss:
        raise ValueError(f"Manifest is missing columns: {', '.join(miss)}")
    return rows


def load_figure_title_registry() -> Dict[str, str]:
    path = OUT_ROOT / "figure_title_registry.csv"
    if not path.exists():
        raise FileNotFoundError(f"Missing figure title registry: {path}")
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
    required = ["fig_id", "rendered_title", "has_serial_prefix"]
    for c in required:
        if c not in (reader.fieldnames or []):
            raise ValueError(f"Invalid figure title registry, missing column: {c}")
    return {str(r.get("fig_id", "")).strip(): normalize_disease_text(str(r.get("rendered_title", "")).strip()) for r in rows if str(r.get("fig_id", "")).strip()}


def ensure_figure_files(fig_rows: List[Dict[str, str]]) -> None:
    missing: List[str] = []
    for row in fig_rows:
        name = row.get("输出文件名", "")
        if not name:
            missing.append(f"{row.get('fig_id', 'UNKNOWN')}:<empty>")
            continue
        fig_path = FIG_DIR / name
        if not fig_path.exists():
            missing.append(f"{row.get('fig_id', 'UNKNOWN')}:{name}")
    if missing:
        raise FileNotFoundError("Missing figure files for stage4/stage5: " + ", ".join(missing))


def cleanup_intermediate_outputs() -> None:
    keep_files = {FINAL_DOCX.name, "qa_check.txt", "role_run.log"}
    for p in OUT_ROOT.iterdir():
        if p.is_file():
            if p.name in keep_files:
                continue
            try:
                p.unlink()
            except Exception:
                pass
            continue
        if p.is_dir():
            try:
                shutil.rmtree(p)
            except Exception:
                pass


def run_stage1_evidence() -> None:
    ensure_runtime_dirs()
    write_evidence_and_refs()
    write_codex_preflight_assets()
    print(f"阶段1完成：{OUT_ROOT / '00_evidence.txt'}")


def run_stage3_ch4_and_figures() -> None:
    ensure_runtime_dirs()
    ensure_inputs(require_excel=True)
    ch4 = build_ch4_data(EXCEL_PATH)
    write_ch4_profile_files(ch4)
    write_text(OUT_ROOT / CH4_NARRATIVE_BRIEF_NAME, build_ch4_narrative_brief(ch4) + "\n")
    write_codex_preflight_assets()
    specs = build_block_specs()
    try:
        block_text = load_block_text_from_files(specs)
        summary_path = OUT_ROOT / "summary.txt"
        if not summary_path.exists():
            raise FileNotFoundError(f"Missing summary file: {summary_path}")
        summary_text = summary_path.read_text(encoding="utf-8").strip()
        txt_report, txt_passed = run_txt_stage_checks(specs, block_text, summary_text)
        print(txt_report)
        if not txt_passed:
            raise RuntimeError("TXT gate failed. Please use txt_stage_qa.txt and codex_rewrite_prompt.txt to repair length and chapter gaps before rerunning.")
        ensure_fig23_codex_spec_ready()
        if not fig23_codex_authored_ok():
            raise RuntimeError(
                f"Missing Codex fig23 spec: {OUT_ROOT / FIG23_CODEX_SPEC_NAME}."
                f"Please read {OUT_ROOT / FIG23_CODEX_PROMPT_NAME} and rewrite it using {OUT_ROOT / FIG23_CODEX_SPEC_TEMPLATE_NAME}."
            )
        if not (OUT_ROOT / "figure_specs.json").exists():
            print(
                f"Hint: if you want Codex to control fig_1_3/fig_3_1/fig_5_3/fig_6_2, review {OUT_ROOT / FIGURE_SPECS_CODEX_PROMPT_NAME}. "
                f"Then write {OUT_ROOT / 'figure_specs.json'} using {OUT_ROOT / FIGURE_SPECS_CODEX_TEMPLATE_NAME}."
            )
        print("Stage3: reusing existing chapter text and summary.")
    except (FileNotFoundError, ValueError) as exc:
        raise RuntimeError(
            "Missing reusable body text. "
            f"Write ch01~ch07.txt/summary.txt/refs.txt by following {OUT_ROOT / CODEX_CONTENT_BLUEPRINT_NAME}; "
            f"if rewriting is needed, use {OUT_ROOT / 'codex_rewrite_prompt.txt'}; "
            f"write fig23 via {OUT_ROOT / FIG23_CODEX_PROMPT_NAME} -> {OUT_ROOT / FIG23_CODEX_SPEC_NAME}; "
            f"and for other semantic figures use {OUT_ROOT / FIGURE_SPECS_CODEX_PROMPT_NAME} -> {OUT_ROOT / 'figure_specs.json'}."
        ) from exc
    fig_rows = generate_figures(ch4)
    make_manifest_files(specs, fig_rows)
    print(f"阶段3完成（复用文本）：{OUT_ROOT / 'ch04_agg_tables.xlsx'}，图表数={len(fig_rows)}")


def run_stage4_assemble_docx() -> None:
    ensure_runtime_dirs()
    ensure_inputs(require_template=True)
    cleanup_stale_final_docx()
    specs = build_block_specs()
    block_text = load_block_text_from_files(specs)
    summary_text, refs_text = load_summary_and_refs()
    fig_rows = load_manifest_fig_rows()
    ensure_figure_files(fig_rows)
    assemble_docx(specs, block_text, summary_text, refs_text, fig_rows)
    post_process_docx_xml(FINAL_DOCX)
    print(f"阶段4完成：{FINAL_DOCX}")


def run_stage5_qa() -> None:
    ensure_runtime_dirs()
    specs = build_block_specs()
    if not FINAL_DOCX.exists():
        raise FileNotFoundError(f"Missing final docx: {FINAL_DOCX}. 请先执行阶段4。")
    block_text = load_block_text_from_files(specs)
    summary_text, _ = load_summary_and_refs()
    fig_rows = load_manifest_fig_rows()
    ensure_figure_files(fig_rows)
    qa, qa_passed = run_checks(specs, block_text, fig_rows, summary_text)
    print(qa)
    if not qa_passed:
        raise RuntimeError("QA未通过硬门槛，流程已中断。请先按qa_check.txt修订后重新执行。")
    if LITE_OUTPUT:
        cleanup_intermediate_outputs()
        print("已启用轻量输出：中间产物已清理，仅保留final docx与qa结果。")
    print(f"\n阶段5完成：{OUT_ROOT / 'qa_check.txt'}")


def ensure_prewritten_text_ready() -> None:
    specs = build_block_specs()
    _ = load_block_text_from_files(specs)
    summary_path = OUT_ROOT / "summary.txt"
    if not summary_path.exists():
        raise FileNotFoundError(f"Missing summary file: {summary_path}")


def run_assist_pipeline() -> None:
    ensure_runtime_dirs()
    try:
        ensure_prewritten_text_ready()
    except Exception as exc:
        raise RuntimeError(
            "Assist mode requires body text written by the current Codex session (ch01~ch07.txt and summary.txt)."
            f"Please use {OUT_ROOT / CODEX_CONTENT_BLUEPRINT_NAME} and {OUT_ROOT / 'codex_rewrite_prompt.txt'} before rerunning."
        ) from exc

    plan = [
        {
            "role_id": "content",
            "role_name": "Content Agent",
            "description": "复用会话正文，执行第四章数据专线+图表与清单",
            "stage_hint": "stage3(reuse-only)",
            "runner": run_stage3_ch4_and_figures,
        },
        {
            "role_id": "docx",
            "role_name": "Docx Agent",
            "description": "装配并后处理最终Word文档",
            "stage_hint": "stage4",
            "runner": run_stage4_assemble_docx,
        },
        {
            "role_id": "qa",
            "role_name": "QA Agent",
            "description": "执行最终质量检查并输出qa_check",
            "stage_hint": "stage5",
            "runner": run_stage5_qa,
        },
    ]

    log_lines = [
        "【Role执行日志】",
        f"开始时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "执行模式：assist",
        f"流程模式：{WORKFLOW_MODE}",
        "计划角色：" + ", ".join([f"{x['role_id']}({x['stage_hint']})" for x in plan]),
        "",
    ]

    total = len(plan)
    for idx, spec in enumerate(plan, start=1):
        started = datetime.now()
        print(f"[角色 {idx}/{total}] {spec['role_name']} ({spec['role_id']}) 开始：{spec['description']}")
        try:
            spec["runner"]()
        except Exception as exc:
            elapsed = (datetime.now() - started).total_seconds()
            log_lines.append(
                f"[FAIL] {spec['role_id']} | stage={spec['stage_hint']} | duration={elapsed:.1f}s | error={type(exc).__name__}: {exc}"
            )
            write_text(OUT_ROOT / "role_run.log", "\n".join(log_lines) + "\n")
            raise
        elapsed = (datetime.now() - started).total_seconds()
        log_lines.append(f"[PASS] {spec['role_id']} | stage={spec['stage_hint']} | duration={elapsed:.1f}s")
        print(f"[角色 {idx}/{total}] {spec['role_name']} ({spec['role_id']}) 完成，用时{elapsed:.1f}s")

    log_lines.append("")
    log_lines.append(f"结束时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    write_text(OUT_ROOT / "role_run.log", "\n".join(log_lines) + "\n")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Medical-topic report pipeline runner")
    parser.add_argument("--topic", default=None, help="医学主题，例如：肠黏膜修复、儿童流感疫苗市场")
    parser.add_argument("--disease", default=None, help="兼容旧参数：等同于 --topic")
    parser.add_argument("--all-topics", action="store_true", help="遍历 data 目录下全部 *.xlsx，并按文件名自动生成所有医学主题报告")
    parser.add_argument("--data-dir", default="data", help="医学主题 Excel 目录，默认 data")
    parser.add_argument("--from-readme", action="store_true", help="从README读取“医学主题：”配置（未传--topic时生效，兼容“疾病名：”）")
    parser.add_argument("--readme", default="README.md", help="README文件路径（默认README.md）")
    parser.add_argument("--xlsx", default=None, help="第四章Excel路径，默认：data/<医学主题>.xlsx")
    parser.add_argument("--template", default="template.docx", help="Word模板路径")
    parser.add_argument("--out-base", default="autofile", help="输出根目录（默认autofile）")
    parser.add_argument("--lite-output", action="store_true", help="轻量输出：流程结束后清理中间产物，仅保留final docx与qa结果")
    return parser.parse_args()


def run(
    topic: str,
    xlsx: str | None = None,
    template: str | None = "template.docx",
    out_base: str | None = "autofile",
    lite_output: bool = False,
) -> None:
    configure_runtime(
        disease_name=topic,
        excel_path=Path(xlsx) if xlsx else None,
        template_path=Path(template) if template else None,
        out_base=Path(out_base) if out_base else None,
    )
    configure_output_mode(lite_output)
    print(f"流程模式：{WORKFLOW_MODE}（固定）")
    print("Hint: the script prepares evidence, chapter-4 structured data, and Codex prompt assets; body text, fig23 config, and semantic figure overrides must be authored by the current Codex session.")
    print("QA闸门：严格（固定，失败即中断）")
    run_stage1_evidence()
    ensure_codex_prep_assets_ready()
    run_assist_pipeline()


def load_qa_result(qa_path: Path) -> str:
    if not qa_path.exists():
        return "缺失"
    text = qa_path.read_text(encoding="utf-8", errors="ignore")
    m = re.search(r"【最终判定】[\s\S]*?结果：([^\n\r]+)", text)
    return m.group(1).strip() if m else "未知"


def run_batch(
    data_dir: str | Path = "data",
    template: str | None = "template.docx",
    out_base: str | None = "autofile",
    lite_output: bool = False,
) -> None:
    root = Path(data_dir)
    if not root.exists():
        raise FileNotFoundError(f"医学主题数据目录不存在：{root}")
    xlsx_files = sorted(root.glob("*.xlsx"), key=lambda p: p.name)
    if not xlsx_files:
        raise FileNotFoundError(f"未在 {root} 中找到任何 .xlsx 文件。")

    rows: List[Dict[str, str]] = []
    failures: List[str] = []
    for idx, xlsx_path in enumerate(xlsx_files, start=1):
        topic = xlsx_path.stem
        print(f"\n===== 批量任务 {idx}/{len(xlsx_files)}：{topic} =====")
        status = "通过"
        detail = ""
        try:
            run(topic=topic, xlsx=str(xlsx_path), template=template, out_base=out_base, lite_output=lite_output)
            qa_result = load_qa_result((Path(out_base or "autofile") / topic / "qa_check.txt"))
            if qa_result != "通过":
                raise RuntimeError(f"qa_check 最终判定为：{qa_result}")
        except Exception as exc:
            status = "失败"
            detail = f"{type(exc).__name__}: {exc}"
            failures.append(f"{topic} -> {detail}")
        rows.append(
            {
                "topic": topic,
                "xlsx": str(xlsx_path),
                "out_dir": str(Path(out_base or "autofile") / topic),
                "status": status,
                "detail": detail,
            }
        )

    write_csv(Path(out_base or "autofile") / "batch_report_summary.csv", rows, ["topic", "xlsx", "out_dir", "status", "detail"])
    if failures:
        raise RuntimeError("批量生成仍有未通过主题：" + "； ".join(failures))


def run_stage1(disease: str, out_base: str | None = "autofile") -> None:
    configure_runtime(disease_name=disease, out_base=Path(out_base) if out_base else None)
    run_stage1_evidence()


def run_stage3(
    disease: str,
    xlsx: str | None = None,
    out_base: str | None = "autofile",
) -> None:
    configure_runtime(
        disease_name=disease,
        excel_path=Path(xlsx) if xlsx else None,
        out_base=Path(out_base) if out_base else None,
    )
    run_stage3_ch4_and_figures()


def run_stage4(
    disease: str,
    template: str | None = "template.docx",
    out_base: str | None = "autofile",
) -> None:
    configure_runtime(
        disease_name=disease,
        template_path=Path(template) if template else None,
        out_base=Path(out_base) if out_base else None,
    )
    run_stage4_assemble_docx()


def run_stage5(disease: str, out_base: str | None = "autofile") -> None:
    configure_runtime(disease_name=disease, out_base=Path(out_base) if out_base else None)
    run_stage5_qa()


def main() -> None:
    args = parse_args()
    if args.all_topics:
        run_batch(data_dir=args.data_dir, template=args.template, out_base=args.out_base, lite_output=args.lite_output)
        return
    topic_name = resolve_topic_name(
        topic=args.topic,
        disease=args.disease,
        from_readme=args.from_readme,
        readme_path=args.readme,
    )
    run(topic=topic_name, xlsx=args.xlsx, template=args.template, out_base=args.out_base, lite_output=args.lite_output)


if __name__ == "__main__":
    main()
